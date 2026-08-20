from typing import List, Dict, Any, Optional
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.document_parser import (
    parse_document,
    extract_paragraphs,
    extract_sections,
    extract_tables,
    get_heading_level,
    iter_paragraphs_with_context,
)
from app.services.citation_sensor import (
    run_citation_sensor,
    _find_references_start,
    APPENDIX_HEADER_PATTERN,
)
from app.services.role_eligibility import (
    is_font_eligible,
    is_heading_role,
    is_body_role,
    is_caption_role,
    is_reference_role,
    typography_skipped,
    body_prose_paragraph,
)
from app.config import settings


# Authoritative LayoutViolation lives in app.services.layout_violation.
# Re-exported here so existing imports (`from app.services.layout_engine
# import LayoutViolation`) keep working. Avoid re-defining the class here
# or you reintroduce the layout_engine <-> citation_sensor import cycle.
LayoutViolation = __import__(
    "app.services.layout_violation", fromlist=["LayoutViolation"]
).LayoutViolation


def _heading_style_level(style_name: Optional[str], role: Optional[str]) -> Optional[int]:
    """Heading level for font/typography: authoritative role wins when the
    paragraph is heading-like; otherwise the style-derived level (legacy
    fallback for callers without roles)."""
    if is_heading_role(role):
        if role in ("REFERENCES_HEADING", "APPENDIX_HEADING"):
            return 1
        return int(role.rsplit("_", 1)[-1])
    return get_heading_level(style_name)


def check_font_consistency(
    paragraphs: List[Dict],
    preset,
    roles: Optional[List[Optional[str]]] = None,
) -> List[LayoutViolation]:
    """FR-2.1: Font Consistency - scan runs, flag mismatched font families.

    Role-aware (Phase 2A): only font-eligible roles (BODY, APPENDIX_BODY,
    heading-like, LIST_ITEM) validate. Cover labels, Title/Subtitle, TOC
    entries, equations, figure hosts, empty/field-only, UNKNOWN, and
    reference entries are skipped — the authoritative role decides, not the
    style name. When `roles` is None (legacy callers), the previous
    style-heuristic behavior is preserved.

    Profile-aware (Build 4): when the config carries an allowed font-and-size
    combination set, each run validates (family, size) TOGETHER against the
    allowed pairs — a valid family with a size from another family is
    rejected. Messages reference the selected profile.
    """
    violations = []
    expected_font = preset.FONT_FAMILY
    profile_label = getattr(preset, "profile_label", None)
    has_combos = bool(getattr(preset, "BODY_ALLOWED_FONT_COMBOS", ()))

    for i, para in enumerate(paragraphs):
        style = para.get("style_name", "")
        role = roles[i] if roles is not None else None
        if role is not None:
            if not is_font_eligible(role):
                continue
        elif not style or style.lower().startswith("heading"):
            continue  # Headings checked separately (legacy fallback)

        # Headings use the heading family; body/list use the body family.
        level = _heading_style_level(style, role)
        heading = level is not None
        expected_family = (
            getattr(preset, "HEADING_FONT_FAMILY", None) if heading else expected_font
        )

        for run in para.get("runs", []):
            font_name = run.get("font_name")
            if not font_name:
                continue
            font_size = run.get("font_size")
            if has_combos:
                # Allowed-pair profile (APA): validate family+size together.
                if not preset.is_font_pair_allowed(font_name, font_size, heading=heading):
                    if profile_label:
                        msg = (
                            f"This text uses {font_name} {font_size or '?'}pt. "
                            f"The selected {profile_label} profile requires an "
                            f"allowed font-and-size combination."
                        )
                    else:
                        msg = f"Font '{font_name}' does not match required '{expected_font}'"
                    violations.append(LayoutViolation(
                        rule_code="FONT_CONSISTENCY",
                        severity="MINOR",
                        location={"paragraph_index": para["index"], "run_index": run["index"]},
                        message=msg,
                        expected_value=", ".join(f"{f} {s:g}pt" for f, s in preset.BODY_ALLOWED_FONT_COMBOS) or expected_font,
                        actual_value=f"{font_name} {font_size or ''}pt",
                    ))
                continue
            # Exact-family profile.
            if expected_family and font_name.lower() != expected_family.lower():
                if profile_label:
                    msg = (
                        f"This text uses {font_name}. The selected {profile_label} "
                        f"profile requires {expected_family}."
                    )
                else:
                    msg = f"Font '{font_name}' does not match required '{expected_family}'"
                violations.append(LayoutViolation(
                    rule_code="FONT_CONSISTENCY",
                    severity="MINOR",
                    location={"paragraph_index": para["index"], "run_index": run["index"]},
                    message=msg,
                    expected_value=expected_family,
                    actual_value=font_name,
                ))
    return violations


def check_font_size(
    paragraphs: List[Dict],
    preset,
    roles: Optional[List[Optional[str]]] = None,
) -> List[LayoutViolation]:
    """FR-2.2: Font Size Alignment - verify heading/body sizes against presets.

    Role-aware (Phase 2A): only font-eligible roles validate. The expected
    size comes from the authoritative heading role when present, else the
    style-derived heading level (legacy fallback). Cover/TOC/equation/
    figure-host/empty/field-only/UNKNOWN/reference-entry paragraphs are
    skipped.

    Profile-aware (Build 4): heading sizes come from the resolved snapshot
    (per-level; APA inherits the body pair — no fixed 16/14/12 pt descent).
    When the profile uses an allowed font-and-size combination set, a run is
    only flagged when its (family, size) is NOT a valid pair; exact-size
    profiles keep the size-only check. Messages reference the profile.
    """
    violations = []
    profile_label = getattr(preset, "profile_label", None)
    has_combos = bool(getattr(preset, "BODY_ALLOWED_FONT_COMBOS", ()))

    for i, para in enumerate(paragraphs):
        style = para.get("style_name", "")
        role = roles[i] if roles is not None else None
        level = _heading_style_level(style, role)

        expected_size = None
        if level in (1, 2, 3):
            expected_size = preset.heading_expected_size(level) if hasattr(preset, "heading_expected_size") else {
                1: getattr(preset, "FONT_SIZE_H1", None),
                2: getattr(preset, "FONT_SIZE_H2", None),
                3: getattr(preset, "FONT_SIZE_H3", None),
            }.get(level)
        elif not level:
            if role is not None:
                # Role-aware: body-like roles get body size; everything else
                # (cover, TOC, equations, UNKNOWN...) is skipped.
                if is_body_role(role):
                    expected_size = preset.FONT_SIZE_BODY
            else:
                expected_size = preset.FONT_SIZE_BODY

        if expected_size is None:
            continue

        for run in para.get("runs", []):
            font_size = run.get("font_size")
            if not font_size:
                continue
            font_name = run.get("font_name")
            if has_combos:
                # Allowed-pair profile: flag only when the pair is invalid.
                if preset.is_font_pair_allowed(font_name, font_size, heading=level is not None):
                    continue
                severity = "MAJOR" if level else "MINOR"
                if profile_label:
                    msg = (
                        f"This {'heading' if level else 'body'} text uses "
                        f"{font_name} {font_size:g}pt. The selected {profile_label} "
                        f"profile requires an allowed font-and-size combination."
                    )
                else:
                    msg = f"{style or 'Body'} font size {font_size}pt is not an allowed combination"
                violations.append(LayoutViolation(
                    rule_code="FONT_SIZE",
                    severity=severity,
                    location={"paragraph_index": para["index"], "run_index": run["index"]},
                    message=msg,
                    expected_value=", ".join(f"{f} {s:g}pt" for f, s in (
                        getattr(preset, "HEADING_ALLOWED_FONT_COMBOS", ()) if level
                        else getattr(preset, "BODY_ALLOWED_FONT_COMBOS", ())
                    )) or f"{expected_size}pt",
                    actual_value=f"{font_name} {font_size:g}pt",
                ))
                continue
            if abs(font_size - expected_size) > 0.5:  # Allow 0.5pt tolerance
                severity = "MAJOR" if level else "MINOR"
                if profile_label:
                    what = "Heading" if level else "Body"
                    msg = (
                        f"This {what} text uses {font_size:g} pt. The selected "
                        f"{profile_label} profile requires {expected_size:g} pt."
                    )
                else:
                    msg = f"{style or 'Body'} font size {font_size}pt != required {expected_size}pt"
                violations.append(LayoutViolation(
                    rule_code="FONT_SIZE",
                    severity=severity,
                    location={"paragraph_index": para["index"], "run_index": run["index"]},
                    message=msg,
                    expected_value=f"{expected_size}pt",
                    actual_value=f"{font_size}pt",
                ))
    return violations


def _is_list_item(para_dict: Dict, paragraph_obj=None) -> bool:
    """Detect whether a paragraph is a numbered/bulleted list item.

    Two sources are unified so direct-formatting and style-inherited lists
    behave identically:
      - the `is_list_item` flag set by extract_paragraphs() from a DIRECT
        <w:numPr> on the paragraph, and
      - numbering inherited through the style chain (e.g. Word's List
        Bullet / List Number styles), detected via _style_has_numbering().

    List items are exempt from the body-alignment rule (they're legitimately
    left-aligned even when body is justified) and, when the preset is silent
    on list spacing, from the SPACE_BEFORE/SPACE_AFTER checks.
    """
    if para_dict.get("is_list_item", False):
        return True
    if paragraph_obj is not None:
        try:
            style = paragraph_obj.style
        except Exception:
            style = None
        if style is not None and _style_has_numbering(style):
            return True
    return False


def _paragraph_visible_text(para: Dict) -> str:
    """Visible text from actual text runs, trimmed.

    Drawing/image runs, field-instruction runs, and empty runs contribute
    nothing — a paragraph that only hosts an image (or only carries an
    uncached field) has no visible text and is not body text.
    """
    return "".join((r.get("text") or "") for r in para.get("runs", [])).strip()


def _is_caption_paragraph(para: Dict, paragraph_obj=None, preset=None) -> bool:
    """True when the paragraph is a caption — semantic or manual.

    Shared eligibility classifier for ALIGNMENT and SPACE_BEFORE/SPACE_AFTER
    so both rules agree on what counts as a caption:
      - semantic: Caption paragraph style (name contains 'caption') — the
        style covers SEQ-field captions whose visible text drops the number
        ("Figure : Semantically captioned…");
      - manual: visible text matches the preset's caption label patterns
        ("Table 1: …", "Figure 2 …", multilingual labels).
    Headings and list items whose text merely resembles a caption are NOT
    captions for eligibility — they are excluded by the style/pattern tests
    (heading/list styles don't contain 'caption' and their text only matches
    if the preset pattern matches, which the caller rules out first).
    """
    style_name = (para.get("style_name") or "").lower()
    if "caption" in style_name:
        return True
    if preset is not None and preset.is_caption_text(para.get("text") or ""):
        return True
    return False


def _style_has_numbering(style) -> bool:
    """True when a style (or any of its base styles) carries list numbering.

    Numbering inherited through the style chain (e.g. Word's List Bullet /
    List Number styles) lives on the STYLE's <w:numPr>, not on the
    paragraph — direct <w:numPr> detection alone misses those.
    """
    from docx.oxml.ns import qn

    seen = set()
    current = style
    while current is not None and getattr(current, "style_id", None) not in seen:
        seen.add(current.style_id)
        pPr = current.element.pPr
        if pPr is not None and pPr.find(qn("w:numPr")) is not None:
            return True
        current = current.base_style
    return False


def _references_span(paragraphs: List[Dict]):
    """Index span of the References section, or None when undetected.

    Reuses the citation sensor's References-header detection; the span ends
    at the first Appendix heading (bibliography ends there), matching the
    sensor's own boundary logic.
    """
    start = _find_references_start(paragraphs)
    if start is None:
        return None
    end = len(paragraphs)
    for para in paragraphs[start:]:
        text = (para.get("text") or "").strip()
        if text and APPENDIX_HEADER_PATTERN.match(text):
            end = para.get("index", end)
            break
    return range(start, end)


def _alignment_eligible(
    para: Dict,
    paragraph_obj,
    preset,
    references_span,
) -> bool:
    """ALIGNMENT applies only to eligible visible-text paragraphs.

    Skipped (alignment is a design choice, not a body-text requirement):
      - empty / field-only / image-only paragraphs (no visible text);
      - semantic Caption-style paragraphs;
      - manual caption text classified by the caption logic;
      - Title and Subtitle styles;
      - list items, including numbering inherited through the style chain;
      - Reference-list entries inside the detected References section.

    Headings and mixed text-and-image paragraphs remain eligible — headings
    use their configured preset alignment, and mixed paragraphs validate
    whenever visible text exists.
    """
    # 1/2/3. Empty, field-only, image-only: no visible text.
    if not _paragraph_visible_text(para):
        return False

    style_name = (para.get("style_name") or "").lower()

    # 6. Title / Subtitle styles.
    if style_name.startswith(("title", "subtitle")):
        return False

    # 4/5. Captions — semantic (Caption style) and manual (caption text).
    if _is_caption_paragraph(para, paragraph_obj, preset):
        return False

    # 7. Numbering inherited through the style chain.
    if paragraph_obj is not None:
        try:
            style = paragraph_obj.style
        except Exception:
            style = None
        if style is not None and _style_has_numbering(style):
            return False

    # 8. Reference-list entries inside the detected References section.
    if references_span is not None and para.get("index", -1) in references_span:
        return False

    return True


def check_paragraph_typography(
    paragraphs: List[Dict],
    preset,
    doc: Optional[Document] = None,
    roles: Optional[List[Optional[str]]] = None,
) -> List[LayoutViolation]:
    """FR-2.3: Paragraph typography — line spacing, spacing before/after, alignment.

    Role-aware (Phase 2A): the authoritative paragraph role decides which
    requirements apply.

      BODY / APPENDIX_BODY(prose):
        body line spacing, body space-before/after, body alignment.
      HEADING_1/2/3 / REFERENCES_HEADING / APPENDIX_HEADING:
        heading line spacing, heading space-before/after, heading alignment.
      LIST_ITEM:
        visible-text font checks only — list-specific spacing policy is
        preserved (LIST_SPACE_AFTER=None ⇒ no deterministic list-spacing
        finding; configured ⇒ SPACE_AFTER validates against it).
      CAPTION_TABLE / CAPTION_FIGURE:
        NO body font/alignment/line-spacing; only CAPTION_SPACE_BEFORE /
        CAPTION_SPACE_AFTER when explicitly configured (None ⇒ no finding).
      REFERENCE_ENTRY:
        reference-specific formatting only — REFERENCES_LINE_SPACING
        (default 2.0). No BODY line-spacing, alignment, or
        paragraph-spacing findings.
      COVER, TITLE, SUBTITLE, TOC_*, DISPLAYED_EQUATION, FIGURE_HOST,
      EMPTY, FIELD_ONLY, UNKNOWN:
        no typography findings; no invented cover requirements; UNKNOWN is
        never converted to BODY.

    When `roles` is None (legacy callers), the previous style/text-heuristic
    behavior is preserved exactly.
    """
    violations = []
    references_span = _references_span(paragraphs)
    profile_label = getattr(preset, "profile_label", None)

    def _msg(what: str, actual, expected, unit: str) -> str:
        """Profile-aware message: 'This <what> is <actual><unit>. The
        selected "<profile>" profile requires <expected><unit>.' Falls back
        to the legacy format when no profile label exists."""
        if profile_label:
            if unit:
                return (
                    f"This {what} is {actual:g} {unit}. The selected {profile_label} "
                    f"profile requires {expected:g} {unit}."
                )
            return (
                f"This {what} is {actual:g}. The selected {profile_label} "
                f"profile requires {expected:g}."
            )
        return f"{what} {actual} != required {expected}"

    def _line_spacing_msg(actual, expected) -> str:
        """Line-spacing message — a multiplier, no unit suffix."""
        if profile_label:
            return (
                f"This line spacing is {actual:g}. The selected {profile_label} "
                f"profile requires {expected:g}."
            )
        return f"Line spacing {actual} != required {expected}"

    # doc.paragraphs aligns 1:1 with extract_paragraphs output — pairing
    # gives access to the paragraph style chain for inherited numbering.
    paired = zip(paragraphs, doc.paragraphs) if doc is not None else ((p, None) for p in paragraphs)

    for i, (para, paragraph_obj) in enumerate(paired):
        style = para.get("style_name", "")
        role = roles[i] if roles is not None else None
        level = _heading_style_level(style, role)
        is_heading = level is not None
        is_list = _is_list_item(para, paragraph_obj)
        is_caption = _is_caption_paragraph(para, paragraph_obj, preset)

        # ---- Role-gated early exit: skip non-typography roles entirely.
        # Cover/TOC/equation/figure-host/empty/field-only/UNKNOWN get no
        # BODY typography findings; captions and reference entries are
        # handled by their dedicated branches below.
        if role is not None:
            if typography_skipped(role):
                continue
            if is_caption_role(role):
                # Caption-specific spacing only — no line-spacing/alignment.
                expected_before = preset.CAPTION_SPACE_BEFORE
                expected_after = preset.CAPTION_SPACE_AFTER
                actual_before = para.get("space_before")
                actual_after = para.get("space_after")
                if expected_before is not None and actual_before is not None and abs(actual_before - expected_before) > 1:
                    violations.append(LayoutViolation(
                        rule_code="SPACE_BEFORE",
                        severity="MINOR",
                        location={"paragraph_index": para["index"]},
                        message=_msg("space before this caption", actual_before, expected_before, "pt"),
                        expected_value=f"{expected_before}pt",
                        actual_value=f"{actual_before}pt",
                    ))
                if expected_after is not None and actual_after is not None and abs(actual_after - expected_after) > 1:
                    violations.append(LayoutViolation(
                        rule_code="SPACE_AFTER",
                        severity="MINOR",
                        location={"paragraph_index": para["index"]},
                        message=_msg("space after this caption", actual_after, expected_after, "pt"),
                        expected_value=f"{expected_after}pt",
                        actual_value=f"{actual_after}pt",
                    ))
                continue
            if is_reference_role(role):
                # Reference-specific formatting only: REFERENCES_LINE_SPACING.
                # No BODY line-spacing, alignment, or paragraph-spacing.
                if preset.REFERENCES_LINE_SPACING is not None:
                    expected = preset.REFERENCES_LINE_SPACING
                    actual = para.get("line_spacing")
                    if actual is not None and abs(actual - expected) > 0.1:
                        violations.append(LayoutViolation(
                            rule_code="LINE_SPACING",
                            severity="MINOR",
                            location={"paragraph_index": para["index"]},
                            message=(
                                f"This reference entry uses {actual:g} line spacing. "
                                f"The selected {profile_label or 'profile'} profile "
                                f"requires {expected:g}."
                            ) if profile_label else
                            f"References line spacing {actual} != required {expected}",
                            expected_value=str(expected),
                            actual_value=str(actual),
                        ))
                continue
            # APPENDIX_BODY is BODY-eligible only when it contains ordinary
            # academic prose — rubric/form content must never become BODY.
            if not body_prose_paragraph(role, para.get("text") or ""):
                continue
            # BODY/APPENDIX_BODY/LIST_ITEM/heading-like: fall through to the
            # common typography path (role-aware heading/body split below).

        # Line spacing — nullable requirement (None) skips the check.
        expected_line_spacing = preset.LINE_SPACING_HEADING if is_heading else preset.LINE_SPACING_BODY
        actual_line_spacing = para.get("line_spacing")
        if expected_line_spacing is not None and actual_line_spacing and abs(actual_line_spacing - expected_line_spacing) > 0.1:
            violations.append(LayoutViolation(
                rule_code="LINE_SPACING",
                severity="MINOR",
                location={"paragraph_index": para["index"]},
                message=_line_spacing_msg(actual_line_spacing, expected_line_spacing),
                expected_value=str(expected_line_spacing),
                actual_value=str(actual_line_spacing),
            ))

        # Space before/after. List items follow list-specific configuration:
        # when the preset is silent (LIST_SPACE_AFTER is None) no spacing
        # requirement is invented and list items are exempt from both checks;
        # when configured, only SPACE_AFTER is validated (there is no
        # list-before configuration). Caption paragraphs (semantic OR manual)
        # follow CAPTION_SPACE_BEFORE / CAPTION_SPACE_AFTER — None (default)
        # means no deterministic caption spacing requirement and both checks
        # are skipped; when configured, the caption validates against that
        # explicit value per side. Ordinary body/heading paragraphs are
        # unchanged.
        if is_list:
            if preset.LIST_SPACE_AFTER is not None:
                actual_after = para.get("space_after")
                if actual_after is not None and abs(actual_after - preset.LIST_SPACE_AFTER) > 1:
                    violations.append(LayoutViolation(
                        rule_code="SPACE_AFTER",
                        severity="MINOR",
                        location={"paragraph_index": para["index"]},
                        message=_msg("space after this list item", actual_after, preset.LIST_SPACE_AFTER, "pt"),
                        expected_value=f"{preset.LIST_SPACE_AFTER}pt",
                        actual_value=f"{actual_after}pt",
                    ))
        elif is_caption:
            expected_before = preset.CAPTION_SPACE_BEFORE
            expected_after = preset.CAPTION_SPACE_AFTER
            actual_before = para.get("space_before")
            actual_after = para.get("space_after")
            # None = the preset is silent on caption spacing — no
            # deterministic requirement, never a finding.
            if expected_before is not None and actual_before is not None and abs(actual_before - expected_before) > 1:
                violations.append(LayoutViolation(
                    rule_code="SPACE_BEFORE",
                    severity="MINOR",
                    location={"paragraph_index": para["index"]},
                    message=_msg("space before this caption", actual_before, expected_before, "pt"),
                    expected_value=f"{expected_before}pt",
                    actual_value=f"{actual_before}pt",
                ))
            if expected_after is not None and actual_after is not None and abs(actual_after - expected_after) > 1:
                violations.append(LayoutViolation(
                    rule_code="SPACE_AFTER",
                    severity="MINOR",
                    location={"paragraph_index": para["index"]},
                    message=_msg("space after this caption", actual_after, expected_after, "pt"),
                    expected_value=f"{expected_after}pt",
                    actual_value=f"{actual_after}pt",
                ))
        else:
            expected_before = preset.SPACE_BEFORE_HEADING if is_heading else preset.SPACE_BEFORE_BODY
            expected_after = preset.SPACE_AFTER_HEADING if is_heading else preset.SPACE_AFTER_BODY

            actual_before = para.get("space_before")
            actual_after = para.get("space_after")

            # Nullable requirement (None) → skip that deterministic check.
            if expected_before is not None and actual_before is not None and abs(actual_before - expected_before) > 1:
                violations.append(LayoutViolation(
                    rule_code="SPACE_BEFORE",
                    severity="MINOR",
                    location={"paragraph_index": para["index"]},
                    message=_msg("space before this paragraph", actual_before, expected_before, "pt"),
                    expected_value=f"{expected_before}pt",
                    actual_value=f"{actual_before}pt",
                ))

            if expected_after is not None and actual_after is not None and abs(actual_after - expected_after) > 1:
                violations.append(LayoutViolation(
                    rule_code="SPACE_AFTER",
                    severity="MINOR",
                    location={"paragraph_index": para["index"]},
                    message=_msg("space after this paragraph", actual_after, expected_after, "pt"),
                    expected_value=f"{expected_after}pt",
                    actual_value=f"{actual_after}pt",
                ))

        # Alignment — only eligible visible-text paragraphs are validated:
        # empty/image-only/caption/title/list/References paragraphs skip.
        if is_list:
            continue
        if not _alignment_eligible(para, paragraph_obj, preset, references_span):
            continue

        expected_align = preset.ALIGNMENT_HEADING if is_heading else preset.ALIGNMENT_BODY
        actual_align = para.get("alignment")
        align_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        }
        actual_align_str = align_map.get(actual_align, "unknown")
        # Nullable alignment (None) → no deterministic requirement.
        if expected_align is None:
            continue
        if actual_align_str != "unknown" and actual_align_str != expected_align:
            if profile_label:
                message = (
                    f"This text is {actual_align_str}-aligned. The selected "
                    f"{profile_label} profile requires {expected_align} alignment."
                )
            else:
                message = f"Alignment '{actual_align_str}' != required '{expected_align}'"
            violations.append(LayoutViolation(
                rule_code="ALIGNMENT",
                severity="MINOR",
                location={"paragraph_index": para["index"]},
                message=message,
                expected_value=expected_align,
                actual_value=actual_align_str,
            ))
    return violations


def check_page_margins(sections: List[Dict], preset) -> List[LayoutViolation]:
    """FR-2.4: Page Margins - measure physical page boundaries.

    Profile-aware (Build 4): nullable margin in the snapshot means no
    deterministic requirement — the check is skipped, never scored. Messages
    reference the selected profile when available.
    """
    violations = []
    profile_label = getattr(preset, "profile_label", None)
    margin_checks = [
        ("margin_left", "MARGIN_LEFT", preset.MARGIN_LEFT),
        ("margin_right", "MARGIN_RIGHT", preset.MARGIN_RIGHT),
        ("margin_top", "MARGIN_TOP", preset.MARGIN_TOP),
        ("margin_bottom", "MARGIN_BOTTOM", preset.MARGIN_BOTTOM),
    ]

    for section_idx, section in enumerate(sections):
        for key, rule_code, expected in margin_checks:
            # Nullable requirement → skip this deterministic check.
            if expected is None:
                continue
            actual = section.get(key)
            if actual is not None and abs(actual - expected) > 0.05:  # 0.05" tolerance
                name = key.replace("margin_", "")
                if profile_label:
                    message = (
                        f"The {name} margin is {actual:.2f} in. The selected "
                        f"{profile_label} profile requires {expected:.2f} in."
                    )
                else:
                    message = f"Page margin {name} {actual:.2f}in != required {expected}in"
                violations.append(LayoutViolation(
                    rule_code=rule_code,
                    severity="MAJOR",
                    location={"section_index": section_idx},
                    message=message,
                    expected_value=f"{expected}in",
                    actual_value=f"{actual:.2f}in",
                ))
    return violations


def check_heading_hierarchy(paragraphs: List[Dict]) -> List[LayoutViolation]:
    """FR-2.5: Heading hierarchy — detect skipped levels AND orphan first headings.

    Two checks:
    1. Skip detection: H1 -> H3 (missing H2) is a MAJOR violation.
    2. Orphan detection: the very first heading in the document must be H1.
       Starting at H2 or H3 breaks Word's outline view and any auto-TOC.
    """
    violations = []
    last_heading_level = 0
    first_heading_seen = False

    for para in paragraphs:
        level = get_heading_level(para.get("style_name"))
        if level is None:
            continue

        # Orphan check: first heading must be H1
        if not first_heading_seen:
            first_heading_seen = True
            if level != 1:
                violations.append(LayoutViolation(
                    rule_code="HEADING_HIERARCHY",
                    severity="MAJOR",
                    location={"paragraph_index": para["index"]},
                    message=f"First heading is H{level} (should be H1). The outline tree must start at level 1.",
                    expected_value="H1",
                    actual_value=f"H{level}",
                ))
            last_heading_level = level
            continue

        # Skip detection (existing logic)
        if last_heading_level > 0 and level > last_heading_level + 1:
            violations.append(LayoutViolation(
                rule_code="HEADING_HIERARCHY",
                severity="MAJOR",
                location={"paragraph_index": para["index"]},
                message=f"Heading level skipped: H{last_heading_level} -> H{level} (missing H{last_heading_level + 1})",
                expected_value=f"H{last_heading_level + 1}",
                actual_value=f"H{level}",
            ))
        last_heading_level = level

    return violations


def _extract_paragraph_text(p_elem) -> str:
    """Extract concatenated text from a <w:p> XML element.

    Uses .iter() to walk ALL descendants because <w:t> text elements are
    nested inside <w:r> runs — direct children of <w:p> are <w:pPr> and
    <w:r>, not <w:t>. Field instruction text (<w:instrText>) is NOT <w:t>,
    so SEQ field instructions never leak into the visible text.
    """
    text = ""
    for descendant in p_elem.iter():
        if descendant.tag.endswith("}t"):
            text += descendant.text or ""
    return text


# --- Semantic caption detection (FR-2.6) ------------------------------------
# A caption is VALID only when it carries Word semantics (Caption paragraph
# style and/or a SEQ Table/Figure field). Plain text that merely looks like
# "Table 2: ..." is a MANUAL caption. Labels are split by object type so a
# table never accepts a figure label and vice versa.

_CAPTION_STYLE_ID = "caption"

_TABLE_LABEL_RE = re.compile(r"^\s*(?:table|tab\.|jadual|表)\s*\d+", re.IGNORECASE)
_FIGURE_LABEL_RE = re.compile(
    r"^\s*(?:figure|fig\.|chart|gambar|rajah|graf|图|图表)\s*\d+", re.IGNORECASE
)

MANUAL_CAPTION_ACTION = (
    "Use Word's References → Insert Caption feature so the label and "
    "numbering remain consistent."
)


def _paragraph_caption_semantics(p_elem, style_names) -> dict:
    """Semantic caption signals of a <w:p>: Caption style, SEQ labels, flags.

    `style_names`: {style_id: lowercase style name} resolved from the doc.
    SEQ labels are collected from ALL <w:instrText> runs — Word splits field
    instructions across multiple runs, e.g. " SEQ " + "Table \\* ARABIC ".
    """
    from docx.oxml.ns import qn

    info = {"caption_style": False, "seq": set(), "heading": False, "list": False}
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        if pPr.find(qn("w:numPr")) is not None:
            info["list"] = True
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is not None:
            style_id = (pStyle.get(qn("w:val")) or "").lower()
            style_name = style_names.get(style_id, "")
            if style_id == _CAPTION_STYLE_ID or style_name == _CAPTION_STYLE_ID:
                info["caption_style"] = True
            if style_id.startswith("heading") or style_name.startswith("heading"):
                info["heading"] = True
    joined_instr = "".join(
        (t.text or "") for t in p_elem.iter(qn("w:instrText"))
    )
    for m in re.finditer(r"\bSEQ\s+(\w+)", joined_instr, re.IGNORECASE):
        info["seq"].add(m.group(1).lower())
    return info


def _caption_status(p_elem, label_re, seq_label, style_names) -> str:
    """Classify an adjacent paragraph as 'valid', 'manual', or 'none'.

    - valid:   label matches the object type AND (Caption style or matching
               SEQ field). Headings and list items are never valid.
    - manual:  text matches the label pattern but carries no Word semantics.
    - none:    not a caption of this object type (e.g. wrong label, prose
               without a number, or the object type's label is absent).
    """
    text = _extract_paragraph_text(p_elem)
    if not text or not label_re.search(text):
        return "none"
    info = _paragraph_caption_semantics(p_elem, style_names)
    if info["heading"] or info["list"]:
        return "manual"
    if info["caption_style"] or seq_label in info["seq"]:
        return "valid"
    return "manual"


def _adjacent_caption_status(element, label_re, seq_label, style_names, paragraphs=None):
    """Caption status from an element's sibling paragraphs (both directions).

    VALID beats MANUAL; a valid caption on either side satisfies the rule.
    When `paragraphs` (the doc-level paragraph list) is given and a MANUAL
    caption sibling is found, the second return value is that caption
    paragraph's zero-based index — the authoritative association the
    frontend uses to locate the caption block (the typed number may
    legitimately differ from the object's ordinal, so it is never required
    to match).
    """
    status = "missing"
    caption_index = None
    for direction in (element.getprevious, element.getnext):
        sibling = direction()
        if sibling is None or not sibling.tag.endswith("}p"):
            continue
        s = _caption_status(sibling, label_re, seq_label, style_names)
        if s == "valid":
            return "valid", caption_index
        if s == "manual":
            status = "manual"
            if paragraphs is not None:
                caption_index = next(
                    (i for i, p in enumerate(paragraphs) if p._p is sibling),
                    None,
                )
    return status, caption_index


def _find_image_context(doc: Document, image_rel) -> tuple:
    """Find (host_paragraph_index, prev_elem, next_elem) for an image.

    Walks the document body in order; when a paragraph contains a drawing
    whose r:embed matches the target image's rId, returns that paragraph's
    index plus the XML elements of its neighbours (None if absent/missing).
    Returns (-1, None, None) if not found.
    """
    from docx.oxml.ns import qn

    paragraphs = doc.paragraphs
    for idx, para in enumerate(paragraphs):
        # Look for <w:drawing> inside this paragraph's runs
        drawings = para._p.findall(".//" + qn("w:drawing"))
        if not drawings:
            # Fallback: try unprefixed
            drawings = para._p.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing")
        for drawing in drawings:
            # r:embed attribute on <a:blip> tells us which image rel this is
            blips = drawing.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
            for blip in blips:
                embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                if embed and embed in doc.part.rels:
                    if doc.part.rels[embed].target_ref == image_rel.target_ref:
                        return idx, para._p.getprevious(), para._p.getnext()
    return -1, None, None


def _find_image_drawing(doc: Document, image_rel):
    """Return the <w:drawing> element embedding this specific image.

    Matches the r:embed rId on <a:blip> against the relationship id, so
    each image resolves to its own drawing — never another image's. Works
    for inline and anchored drawings (both live inside <w:drawing>) and
    for multiple images in one paragraph (each matched by its own rId).
    Returns None if the image is not found in the body.
    """
    from docx.oxml.ns import qn

    for drawing in doc.element.body.iter(qn("w:drawing")):
        for blip in drawing.findall(".//" + qn("a:blip")):
            embed = blip.get(qn("r:embed"))
            if embed and embed == image_rel.rId:
                return drawing
    return None


def _extract_image_alt_text(doc: Document, image_rel) -> str:
    """Extract the docPr@descr (alt-text) for a specific image.

    Resolves alt-text from the image's own drawing properties only — never
    from another image's docPr. Returns "" when the image has no drawing or
    no alt-text (descr/title empty or absent).
    """
    ns = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
    drawing = _find_image_drawing(doc, image_rel)
    if drawing is None:
        return ""
    for docpr in drawing.findall(".//" + ns + "docPr"):
        descr = docpr.get("descr") or docpr.get("title") or ""
        if descr:
            return descr
    return ""


def check_media_captions(
    doc: Document,
    paragraphs: List[Dict],
    preset,
    table_roles: Optional[List[str]] = None,
    paragraph_roles: Optional[List[str]] = None,
    cover_end: Optional[int] = None,
) -> List[LayoutViolation]:
    """FR-2.6: Media captions — scholarly Tables and academic Figures.

    Role-gated (Phase 2B): Caption requirements apply ONLY to proven
    scholarly Tables (`SCHOLARLY_TABLE`) and academic Figures (an image
    hosted in an academic BODY-region paragraph). Administrative cover
    Tables, assignment/assessment/rubric Tables, layout Tables, UNKNOWN
    Tables, cover logos, repeated header images, and decorative images are
    EXEMPT from Caption findings — no scored finding, no manual-review
    deduction.

    Adjacent captions are classified semantically:
    - VALID:   Caption paragraph style and/or a matching SEQ Table/Figure
               field, with a label matching the object type.
    - MANUAL:  text matches 'Table N'/'Figure N' but no Word semantics —
               flagged with MANUAL_CAPTION (MINOR).
    - MISSING: no matching adjacent caption — TABLE_CAPTION_MISSING /
               IMAGE_CAPTION_MISSING.

    Multilingual labels are preserved and split by object type. Headings,
    list items, and table-cell text are never accepted as captions.

    Alt-text: kept independent of Caption eligibility — a Caption exemption
    never automatically exempts Alt Text. (Logo/decorative Alt Text policy
    is unresolved in this Build; see warnings.)
    """
    violations = []
    style_names = {s.style_id: (s.name or "").lower() for s in doc.styles}

    # ---- Academic-region determination for images ----
    # An image host paragraph is FIGURE_HOST regardless of region (the
    # classifier's drawing check fires before the cover check). Academic
    # Figure eligibility needs the REGION: an image hosted at/after the
    # cover end is an academic Figure; images before it (cover logos) are
    # exempt from Caption findings. Body-first documents have cover_end=0 →
    # everything is academic. cover_end=None (legacy callers) falls back to
    # the first academic Heading as the boundary.
    effective_cover_end = cover_end
    if effective_cover_end is None:
        first_heading_idx = None
        if paragraph_roles is not None:
            for i, r in enumerate(paragraph_roles):
                if r in ("HEADING_1", "HEADING_2", "HEADING_3"):
                    first_heading_idx = i
                    break
        effective_cover_end = first_heading_idx if first_heading_idx is not None else 0

    # ---- Tables (role-gated) ----
    for table_idx, table in enumerate(doc.tables):
        role = table_roles[table_idx] if table_roles is not None and table_idx < len(table_roles) else None
        # TABLE_CAPTION_MISSING / MANUAL_CAPTION apply ONLY to proven
        # scholarly Tables. Administrative/rubric/layout/UNKNOWN tables get
        # NO caption finding — a Word table is not automatically scholarly.
        if role != "SCHOLARLY_TABLE":
            continue
        status, caption_idx = _adjacent_caption_status(
            table._tbl, _TABLE_LABEL_RE, "table", style_names, doc.paragraphs
        )
        if status == "valid":
            continue
        if status == "manual":
            location = {"table_index": table_idx}
            if caption_idx is not None:
                location["paragraph_index"] = caption_idx
            violations.append(LayoutViolation(
                rule_code="MANUAL_CAPTION",
                severity="MINOR",
                location=location,
                message=(
                    f"Table {table_idx + 1} has a manually typed caption. "
                    f"{MANUAL_CAPTION_ACTION}"
                ),
                expected_value="Word caption (References → Insert Caption)",
                actual_value="Manual 'Table N' text without Caption style or SEQ field",
            ))
        else:
            violations.append(LayoutViolation(
                rule_code="TABLE_CAPTION_MISSING",
                severity="MINOR",
                location={"table_index": table_idx},
                message=(
                    f"Table {table_idx + 1} has no caption. Add a paragraph "
                    f"above or below it starting with 'Table {table_idx + 1}: ' "
                    f"(or 'Jadual {table_idx + 1}: ' / '表{table_idx + 1}')."
                ),
                expected_value=f"Table {table_idx + 1}: <description>",
                actual_value="No caption found",
            ))

    # ---- Images — caption + alt-text checks ----
    image_idx = 0
    for rel in doc.part.rels.values():
        if "image" not in rel.target_ref:
            continue
        image_idx += 1

        # Walk the document body to find the paragraph hosting this image.
        host_para_idx, prev_elem, next_elem = _find_image_context(doc, rel)
        # Academic Figure = an image hosted at/after the cover end. Images
        # before it (cover logos, cover branding) are exempt from Figure
        # Caption findings.
        academic_figure = (
            host_para_idx >= 0
            and host_para_idx >= effective_cover_end
        )
        if academic_figure:
            status = "missing"
            for elem in (prev_elem, next_elem):
                if elem is None or not elem.tag.endswith("}p"):
                    continue
                s = _caption_status(elem, _FIGURE_LABEL_RE, "figure", style_names)
                if s == "valid":
                    status = "valid"
                    break
                if s == "manual":
                    status = "manual"
            if status == "manual":
                violations.append(LayoutViolation(
                    rule_code="MANUAL_CAPTION",
                    severity="MINOR",
                    location={"image_index": image_idx - 1, "paragraph_index": host_para_idx},
                    message=(
                        f"Image {image_idx} has a manually typed caption. "
                        f"{MANUAL_CAPTION_ACTION}"
                    ),
                    expected_value="Word caption (References → Insert Caption)",
                    actual_value="Manual 'Figure N' text without Caption style or SEQ field",
                ))
            elif status == "missing":
                violations.append(LayoutViolation(
                    rule_code="IMAGE_CAPTION_MISSING",
                    severity="MINOR",
                    location={"image_index": image_idx - 1, "paragraph_index": host_para_idx},
                    message=(
                        f"Image {image_idx} has no caption. Add a paragraph "
                        f"below it starting with 'Figure {image_idx}: ' "
                        f"(or 'Gambar {image_idx}: ' / '图{image_idx}')."
                    ),
                    expected_value=f"Figure {image_idx}: <description>",
                    actual_value="No caption found",
                ))

        # Alt-text check — walk the inline drawing's docPr element. Kept
        # INDEPENDENT of Caption eligibility: a Caption exemption never
        # automatically exempts Alt Text for academic-region images.
        #
        # Cover-region images (cover logo / institutional branding) have NO
        # explicit logo/decorative Alt Text policy in the selected profile —
        # the policy is UNRESOLVED. Per the Phase 2B contract, an unresolved
        # policy must not silently apply a universal rule: no scored finding
        # is created for cover-region images, and the unresolved policy is
        # reported as a warning at the call site.
        alt_text = _extract_image_alt_text(doc, rel)
        in_academic_region = host_para_idx >= 0 and host_para_idx >= effective_cover_end
        if in_academic_region and (not alt_text or not alt_text.strip()):
            violations.append(LayoutViolation(
                rule_code="IMAGE_ALT_TEXT_MISSING",
                severity="MINOR",
                location={"image_index": image_idx - 1, "paragraph_index": host_para_idx},
                message=(
                    f"Image {image_idx} has no alt-text. Right-click the image "
                    f"-> Alt Text -> enter a concise description for screen readers."
                ),
                expected_value="Non-empty alt-text description",
                actual_value="Empty or missing docPr@descr",
            ))

    return violations


def run_static_rules_engine(
    file_bytes: bytes,
    config=None,
) -> List[LayoutViolation]:
    """Main entry point - runs all static layout checks.

    `config` (Build 4): an immutable snapshot-derived rule config
    (EffectiveProfileConfig). Production audit execution ALWAYS passes it —
    the resolved EffectiveProfileSnapshot drives every supported rule, and
    global PresetConfig defaults are never read during a production audit.
    When `config` is None (legacy callers / tests), the global PresetConfig
    is used as the compatibility default.
    """
    doc = parse_document(file_bytes)
    preset = config if config is not None else settings.PRESET

    paragraphs = extract_paragraphs(doc)
    sections = extract_sections(doc)

    # Phase 2A: the authoritative paragraph roles (the same classifier that
    # populates document_blocks.role) gate Font + Paragraph Typography
    # eligibility. Roles are computed once and shared by the rules — rules
    # never reclassify paragraphs themselves.
    from app.services.role_classifier import classify_paragraphs, classify_table_roles, cover_region_end
    roles = classify_paragraphs(doc, paragraphs)
    table_roles = classify_table_roles(doc)
    cover_end = cover_region_end(doc, paragraphs)

    all_violations = []
    all_violations.extend(check_font_consistency(paragraphs, preset, roles=roles))
    all_violations.extend(check_font_size(paragraphs, preset, roles=roles))
    all_violations.extend(check_paragraph_typography(paragraphs, preset, doc=doc, roles=roles))
    all_violations.extend(check_page_margins(sections, preset))
    all_violations.extend(check_heading_hierarchy(paragraphs))
    all_violations.extend(check_media_captions(
        doc, paragraphs, preset,
        table_roles=table_roles,
        paragraph_roles=roles,
        cover_end=cover_end,
    ))
    all_violations.extend(run_citation_sensor(doc, paragraphs))

    return all_violations
