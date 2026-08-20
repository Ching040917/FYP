"""Document Role Classification — Phase 1 PoC (read-only classifier).

One authoritative classification of every DOCX paragraph (and table) into a
structural ROLE, computed once during parsing and persisted on
document_blocks. Rule engines (Font/Spacing/Alignment/Caption/mapping) must
consume these roles via eligibility helpers instead of re-deriving their own
heuristics.

Phase 1 scope: classify and REPORT only. No rule migration, no scoring
change, no finding suppression. Historical audits (no role metadata) are
untouched.

Roles are conservative: structural OOXML evidence > style > visible text.
UNKNOWN means "no proven role" — never a deduction, never a guess.

Paragraph roles:
  COVER, TITLE, SUBTITLE, TABLE_OF_CONTENTS_HEADING, TABLE_OF_CONTENTS_ENTRY,
  HEADING_1, HEADING_2, HEADING_3, BODY, LIST_ITEM, DISPLAYED_EQUATION,
  CAPTION_TABLE, CAPTION_FIGURE, FIGURE_HOST, REFERENCES_HEADING,
  REFERENCE_ENTRY, APPENDIX_HEADING, APPENDIX_BODY, EMPTY, FIELD_ONLY, UNKNOWN.

Table roles (computed, not persisted — no storage structure exists):
  ADMINISTRATIVE_TABLE, SCHOLARLY_TABLE, RUBRIC_TABLE, LAYOUT_TABLE, UNKNOWN.
Table-cell content is CONTEXT metadata (table_index/row_index/cell_index),
never a paragraph role.

Table role policy (one consistent classification):
  - ADMINISTRATIVE_TABLE — any table inside the proven cover region: cover
    information tables, assignment-detail tables, and the MARK/Comments
    assessment table alike (cover-region position decides, before content).
  - RUBRIC_TABLE — rubric-shaped tables (header keywords) in back-matter
    (References or Appendix regions), e.g. the final marking-criteria table.
  - SCHOLARLY_TABLE — manuscript-body tables ONLY with an adjacent numbered
    caption; never inferred from row/column structure alone.
  - UNKNOWN — ambiguous tables and non-caption-target back-matter tables.
Administrative and rubric tables remain ineligible for future
TABLE_CAPTION deductions. No scholarly Caption eligibility is inferred from
table classification.

Cover model:
  Cover may begin at document start only when at least one RELIABLE signal
  exists in the region before the first boundary: Title/Subtitle style, a
  cover-position image/logo (body, header/footer, or table cell — a logo may
  be invisible to `doc.paragraphs`), an administrative/form/assessment table,
  or multiple short structural labels associated with administrative tables.
  Labels alone prove nothing; a short first paragraph alone is insufficient;
  visible words such as `Detail`, `Assessment`, or `Programme` are not
  sufficient without structural evidence.

  Cover ends immediately before the first independently proven boundary in
  OOXML body order: a Table of Contents boundary, an academic Heading 1/2/3,
  a References heading, or an ordinary academic BODY paragraph following a
  structural transition. A document starting directly with a Heading, or with
  a normal multi-sentence paragraph, never becomes COVER; the entire first
  section is never blanket-classified as COVER.

  TOC limitation (documented, not worked around in this Build):
  real-document TOC content may be stored in fields, text boxes, or
  structures not represented by `doc.paragraphs`. Absence from paragraph
  blocks never extends COVER past the first academic Heading — the heading is
  an independent cover end. No synthetic TOC blocks are created.

  Authoritative order model: paragraph and table positions are derived from
  the OOXML body children (paragraphs + tables + drawings in document order),
  never from `doc.paragraphs` order alone when tables interrupt the body
  flow.
"""
import re
from typing import Dict, List, Optional, Set

from docx import Document
from docx.oxml.ns import qn

from app.services.document_parser import get_heading_level, extract_paragraphs
from app.services.citation_sensor import _find_references_start

# ---------------------------------------------------------------------------
# Role enums (plain strings — no DB/API schema dependency)
# ---------------------------------------------------------------------------

PARAGRAPH_ROLES = frozenset({
    "COVER",
    "TITLE",
    "SUBTITLE",
    "TABLE_OF_CONTENTS_HEADING",
    "TABLE_OF_CONTENTS_ENTRY",
    "HEADING_1",
    "HEADING_2",
    "HEADING_3",
    "BODY",
    "LIST_ITEM",
    "DISPLAYED_EQUATION",
    "CAPTION_TABLE",
    "CAPTION_FIGURE",
    "FIGURE_HOST",
    "REFERENCES_HEADING",
    "REFERENCE_ENTRY",
    "APPENDIX_HEADING",
    "APPENDIX_BODY",
    "EMPTY",
    "FIELD_ONLY",
    "UNKNOWN",
})

TABLE_ROLES = frozenset({
    "ADMINISTRATIVE_TABLE",
    "SCHOLARLY_TABLE",
    "RUBRIC_TABLE",
    "LAYOUT_TABLE",
    "UNKNOWN",
})

# Field instruction labels that make a paragraph structural (not body).
_FIELD_LABELS = frozenset({"TOC", "SEQ", "PAGE", "NUMPAGES", "REF", "STYLEREF", "HYPERLINK"})

# References-section headers (mirror of the citation sensor's set).
_REFERENCES_HEADERS = frozenset({"references", "bibliography", "works cited", "reference list"})

_NUMERIC_SECTION_PREFIX = re.compile(r"^\d+(?:\.\d+)*\.?\s+")

# Rubric headings terminate the reference span (back-matter region).
_RUBRIC_HEADING_RE = re.compile(r"\brubric", re.IGNORECASE)

# Header keywords that make a small cover table recognisably administrative
# (cover info / assignment detail / assessment). Evaluated ONLY inside the
# cover region — scholarly body tables never reach this predicate.
_ADMIN_HEADER_KEYWORDS = (
    "name", "student", "id", "matric", "programme", "program", "course",
    "module", "subject", "semester", "intake", "session", "faculty",
    "lecturer", "tutor", "supervisor", "assignment", "submission", "date",
    "assessment", "mark", "comments", "detail", "class", "group",
)


# ---------------------------------------------------------------------------
# Structural evidence extraction (per <w:p>)
# ---------------------------------------------------------------------------


def _para_fields(p_elem) -> Set[str]:
    """Field instruction labels in a paragraph (TOC/SEQ/PAGE/REF...)."""
    if p_elem is None:
        return set()
    labels: Set[str] = set()
    for t in p_elem.iter(qn("w:instrText")):
        text = (t.text or "").upper()
        for word in re.findall(r"[A-Z][A-Z0-9_]*", text):
            if word in _FIELD_LABELS:
                labels.add(word)
    return labels


def _has_drawing(p_elem) -> bool:
    if p_elem is None:
        return False
    return len(p_elem.findall(".//" + qn("w:drawing"))) > 0


def _has_o_math(p_elem) -> bool:
    if p_elem is None:
        return False
    return len(p_elem.findall(".//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath")) > 0


def _style_chain_has_numbering(paragraph_obj) -> bool:
    """Numbering inherited through the style chain (list styles)."""
    seen = set()
    current = paragraph_obj.style if paragraph_obj is not None else None
    while current is not None and getattr(current, "style_id", None) not in seen:
        seen.add(current.style_id)
        pPr = current.element.pPr
        if pPr is not None and pPr.find(qn("w:numPr")) is not None:
            return True
        current = current.base_style
    return False


def _visible_text(para: Dict) -> str:
    return "".join((r.get("text") or "") for r in para.get("runs", [])).strip()


# ---------------------------------------------------------------------------
# Authoritative body-order model (paragraphs + tables in document order)
# ---------------------------------------------------------------------------


def _body_sequence(doc: Document) -> List[Dict]:
    """Top-level OOXML body children in document order.

    Returns one entry per body-level element:
      {'kind': 'p',   'para_index': int, 'has_drawing': bool}
      {'kind': 'tbl', 'table_index': int, 'before_para': Optional[int]}
    `para_index` equals the python-docx `doc.paragraphs` index; `before_para`
    is the index of the last body paragraph before the table (None when the
    table is the first child). Tables interrupt the paragraph flow — never
    infer a table's neighbours from `doc.paragraphs` order alone.
    """
    seq: List[Dict] = []
    para_count = 0
    table_count = 0
    last_para: Optional[int] = None
    for child in doc.element.body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            seq.append({
                "kind": "p",
                "para_index": para_count,
                "has_drawing": len(child.findall(".//" + qn("w:drawing"))) > 0,
            })
            last_para = para_count
            para_count += 1
        elif tag == qn("w:tbl"):
            seq.append({
                "kind": "tbl",
                "table_index": table_count,
                "before_para": last_para,
            })
            table_count += 1
    return seq


# ---------------------------------------------------------------------------
# Region boundaries
# ---------------------------------------------------------------------------


def _is_references_heading_text(text: str) -> bool:
    candidate = _NUMERIC_SECTION_PREFIX.sub("", (text or "").strip())
    return candidate.lower() in _REFERENCES_HEADERS


def _is_appendix_heading(para: Dict, text: str) -> bool:
    """Appendix heading: heading-styled OR a short exact 'Appendix' label.
    Body text merely starting with 'Appendix' (e.g. 'Appendix content …') is
    never a heading."""
    if get_heading_level(para.get("style_name")) is not None:
        return True
    stripped = text.strip()
    return bool(re.match(r"^appendix(?:\s+[A-Za-z0-9]+)?$", stripped, re.IGNORECASE))


def _is_rubric_heading(para: Dict) -> bool:
    style = (para.get("style_name") or "").lower()
    text = para.get("text") or ""
    return get_heading_level(para.get("style_name")) is not None and bool(_RUBRIC_HEADING_RE.search(text))


def _references_span_of(paragraphs: List[Dict], ref_start: Optional[int]):
    if ref_start is None:
        return None
    end = len(paragraphs)
    for para in paragraphs[ref_start:]:
        text = (para.get("text") or "").strip()
        if re.match(r"^appendix\b", text, re.IGNORECASE) or _is_rubric_heading(para):
            end = para["index"]
            break
    return range(ref_start, end)


def _appendix_span_of(paragraphs: List[Dict]):
    start = None
    for para in paragraphs:
        if re.match(r"^appendix\b", (para.get("text") or "").strip(), re.IGNORECASE):
            start = para["index"]
            break
    if start is None:
        return None
    return range(start, len(paragraphs))


def _looks_like_body(text: str) -> bool:
    """Ordinary prose heuristic: a sentence-like paragraph, NOT a cover
    label line ('Detail:', 'Assessment:', a short title phrase)."""
    s = text.strip()
    if not s:
        return False
    if len(s) >= 40:
        return True
    if s.endswith(":"):
        return False  # label lines
    if s.endswith((".", "!", "?")) and len(s) >= 20:
        return True  # short sentence with a period
    if re.search(r"[.!?]\s+[A-Z]", s):
        return True  # multi-sentence
    return False


def _cover_boundary(paragraphs_by_index: Dict[int, Dict], body_seq: List[Dict]) -> Optional[int]:
    """First independently proven cover end, in OOXML body order.

    Boundaries: academic Heading 1/2/3, a TOC boundary, a References heading,
    or the first ordinary academic BODY paragraph following a structural
    transition. Returns None when no boundary exists (cover-only document).
    A document starting directly with a Heading returns 0 — the cover signal
    check then fails on `boundary <= 0` and the document stays academic.
    """
    seen_structural = False
    for e in body_seq:
        if e["kind"] == "tbl":
            # A table before the boundary is structural context — it sits
            # between paragraphs, so doc.paragraphs order alone cannot see it.
            seen_structural = True
            continue
        para = paragraphs_by_index.get(e["para_index"])
        if para is None:
            continue
        text = (para.get("text") or "").strip()
        style = (para.get("style_name") or "").lower()
        level = get_heading_level(para.get("style_name"))
        if level in (1, 2, 3):
            return para["index"]
        if ("toc" in style or "table of contents" in text.lower()
                or _is_references_heading_text(text)):
            return para["index"]
        if seen_structural and _looks_like_body(text):
            # Ordinary academic body paragraph following a structural
            # transition ends the cover immediately before it.
            return para["index"]
        if style.startswith(("title", "subtitle")) or e["has_drawing"]:
            seen_structural = True
    return None


# ---------------------------------------------------------------------------
# Cover signals (reliable structural evidence)
# ---------------------------------------------------------------------------


def _is_admin_table(table) -> bool:
    """Administrative cover table: key-value form (2 columns of short label
    cells) OR a small table with an administrative keyword header row.

    Evaluated ONLY in cover context — scholarly body tables never reach this
    predicate, so row/column structure alone never classifies a scholarly
    table.
    """
    try:
        rows = [[c.text.strip() for c in row.cells] for row in table.rows[:3]]
    except Exception:
        return False
    if not rows:
        return False
    # Key-value form: 2 columns, first cell of each row a short label with no
    # sentence punctuation — the classic cover information / assignment
    # detail / MARK table.
    if len(rows[0]) == 2 and len(rows) >= 2:
        if all(
            row[0] and len(row[0]) <= 40 and not re.search(r"[.!?]\s*$", row[0])
            for row in rows
        ):
            return True
    # Header-keyword form: >= 2 admin keywords in the header row, small table.
    joined = " ".join(rows[0]).lower()
    hits = sum(1 for k in _ADMIN_HEADER_KEYWORDS if k in joined)
    try:
        small = len(table.rows) <= 12
    except Exception:
        small = True
    return hits >= 2 and small


def _header_footer_has_drawing(doc: Document) -> bool:
    """A logo in a header/footer part — invisible to `doc.paragraphs`."""
    for sec in doc.sections:
        sectPr = sec._sectPr
        if sectPr is None:
            continue
        for ref in (*sectPr.findall(qn("w:headerReference")),
                    *sectPr.findall(qn("w:footerReference"))):
            rid = ref.get(qn("r:id"))
            if not rid:
                continue
            try:
                part = doc.part.rels[rid].target_part
            except (KeyError, AttributeError):
                continue
            if part is not None and len(part.element.findall(".//" + qn("w:drawing"))) > 0:
                return True
    return False


def _region_table_has_drawing(doc: Document, body_seq: List[Dict], boundary: int) -> bool:
    """A drawing inside a cover-region table cell (logo embedded in a table)."""
    for e in body_seq:
        if e["kind"] != "tbl":
            continue
        pos = e["before_para"] if e["before_para"] is not None else -1
        if pos >= boundary:
            continue
        try:
            tbl = doc.tables[e["table_index"]]
        except IndexError:
            continue
        if len(tbl._tbl.findall(".//" + qn("w:drawing"))) > 0:
            return True
    return False


def _region_has_drawing(doc: Document, body_seq: List[Dict], boundary: int) -> bool:
    """Cover-position drawing/logo: body paragraph, header/footer part, or
    table cell — a logo may not be visible through `doc.paragraphs`."""
    for e in body_seq:
        if e["kind"] == "p" and e["para_index"] < boundary and e["has_drawing"]:
            return True
    return _header_footer_has_drawing(doc) or _region_table_has_drawing(doc, body_seq, boundary)


def _has_multiple_labels(region_paras: List[Dict], paragraphs_by_index: Dict[int, Dict]) -> bool:
    """At least two short non-prose lines in the region (structural labels
    around cover tables)."""
    labels = 0
    for e in region_paras:
        para = paragraphs_by_index.get(e["para_index"])
        if para is None:
            continue
        text = (para.get("text") or "").strip()
        if text and not _looks_like_body(text):
            labels += 1
    return labels >= 2


def _cover_signal_exists(
    doc: Document,
    paragraphs_by_index: Dict[int, Dict],
    body_seq: List[Dict],
    boundary: Optional[int],
) -> bool:
    """At least one RELIABLE cover signal before `boundary` (index > 0).

    Signals (any one qualifies):
      - Title or Subtitle style;
      - cover-position image or logo (body / header / footer / table cell);
      - administrative / form / assessment table;
      - multiple short structural labels associated with administrative
        tables.
    A body-first document (region contains prose) or labels without any
    structural signal never qualifies — no cover is ever guessed.
    """
    if boundary is None or boundary <= 0:
        return False
    region_paras = [e for e in body_seq if e["kind"] == "p" and e["para_index"] < boundary]
    region_tables = [
        e for e in body_seq if e["kind"] == "tbl"
        and (e["before_para"] is None or e["before_para"] < boundary)
    ]

    # 1. Title / Subtitle style anywhere in the region.
    for e in region_paras:
        para = paragraphs_by_index.get(e["para_index"])
        if para is None:
            continue
        if (para.get("style_name") or "").lower().startswith(("title", "subtitle")):
            return True

    # 2. Cover-position drawing or logo (may be invisible to doc.paragraphs).
    if _region_has_drawing(doc, body_seq, boundary):
        return True

    # 3. Administrative / form / assessment table in the region.
    admin_tables = [
        e for e in region_tables
        if _is_admin_table(doc.tables[e["table_index"]])
    ]
    if admin_tables:
        return True

    # 4. Multiple short structural labels associated with admin tables
    #    (corroboration — labels alone still prove nothing).
    if admin_tables and _has_multiple_labels(region_paras, paragraphs_by_index):
        return True

    return False


# ---------------------------------------------------------------------------
# Paragraph classification
# ---------------------------------------------------------------------------


def cover_region_end(doc: Document, paragraphs: List[Dict]) -> int:
    """Zero-based index where the cover region ends (first paragraph AFTER
    cover). 0 = no cover (body-first). Cover-only documents return
    len(paragraphs). Shared derivation with classify_paragraphs so region
    membership always agrees."""
    body_seq = _body_sequence(doc)
    by_index = {p["index"]: p for p in paragraphs}
    boundary = _cover_boundary(by_index, body_seq)
    if _cover_signal_exists(doc, by_index, body_seq, boundary):
        return boundary if boundary is not None else len(paragraphs)
    return 0


def classify_paragraphs(
    doc: Document,
    paragraphs: List[Dict],
    cover_end: Optional[int] = None,
    references_span: Optional[range] = None,
    appendix_span: Optional[range] = None,
) -> List[str]:
    """Classify each paragraph dict (index-aligned with extract_paragraphs).

    Optional boundary hints are derived structurally when absent.
    """
    if references_span is None:
        ref_start = _find_references_start(paragraphs)
        references_span = _references_span_of(paragraphs, ref_start)
    if appendix_span is None:
        appendix_span = _appendix_span_of(paragraphs)

    # Rubric/back-matter region starts at the first rubric heading (a
    # heading whose text mentions rubric) and extends to the end.
    rubric_region = None
    for p in paragraphs:
        if _is_rubric_heading(p):
            rubric_region = p["index"]
            break

    body_seq = _body_sequence(doc)
    by_index = {p["index"]: p for p in paragraphs}
    if cover_end is None:
        boundary = _cover_boundary(by_index, body_seq)
        if _cover_signal_exists(doc, by_index, body_seq, boundary):
            # Cover-only document (signal but no boundary): cover runs to the
            # end. Otherwise the boundary index is the first paragraph AFTER
            # the cover region.
            cover_end = boundary if boundary is not None else len(paragraphs)
        else:
            cover_end = 0
    paired = (
        zip(paragraphs, doc.paragraphs)
        if len(doc.paragraphs) == len(paragraphs)
        else ((p, None) for p in paragraphs)
    )

    roles: List[str] = []
    for para, paragraph_obj in paired:
        roles.append(_classify_one(para, paragraph_obj, cover_end, references_span, appendix_span, rubric_region))
    return roles


def _style_ids(p_elem) -> Set[str]:
    """Raw w:pStyle val(s) on the paragraph (python-docx may not resolve an
    unknown style id into a name — the raw id is the authoritative TOC cue)."""
    if p_elem is None:
        return set()
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        return set()
    ids = set()
    for pStyle in pPr.findall(qn("w:pStyle")):
        val = (pStyle.get(qn("w:val")) or "").lower()
        if val:
            ids.add(val)
    return ids


def _classify_one(
    para: Dict,
    paragraph_obj,
    cover_end: int,
    references_span: Optional[range],
    appendix_span: Optional[range],
    rubric_region: Optional[int],
) -> str:
    index = para["index"]
    text = (para.get("text") or "").strip()
    style = (para.get("style_name") or "").lower()
    p_elem = paragraph_obj._p if paragraph_obj is not None else None
    style_ids = _style_ids(p_elem)
    visible = _visible_text(para)

    # 1. Structural: figure host (drawing) and displayed equation (OMML)
    #    BEFORE empty — neither has visible run text but both have roles.
    if p_elem is not None and _has_drawing(p_elem):
        return "FIGURE_HOST"
    if p_elem is not None and _has_o_math(p_elem):
        return "DISPLAYED_EQUATION"

    # 2. Empty / field-only (no visible text, no run findings).
    if not visible:
        if p_elem is not None and _para_fields(p_elem):
            return "FIELD_ONLY"
        return "EMPTY"

    # 3. TOC entries (structural: TOC field + style/hyperlink). Raw style id
    #    (TOC1..TOC9) and resolved name both count.
    fields = _para_fields(p_elem) if p_elem is not None else set()
    if "TOC" in fields:
        return "TABLE_OF_CONTENTS_ENTRY"
    if any(s.startswith("toc") and s != "toc heading" for s in style_ids):
        return "TABLE_OF_CONTENTS_ENTRY"
    if "toc" in style and "toc heading" not in style:
        return "TABLE_OF_CONTENTS_ENTRY"
    if _is_toc_heading(para):
        return "TABLE_OF_CONTENTS_HEADING"

    # 4. Displayed equation (Equation style — non-OMML fallback).
    if "equation" in style:
        return "DISPLAYED_EQUATION"

    # 5. Captions — semantic (Caption style / SEQ) and manual labels.
    if "caption" in style or "SEQ" in fields:
        return _caption_role(para, p_elem, style)
    if _is_manual_caption_text(text):
        return _caption_role(para, p_elem, style)

    # 6. Title / Subtitle styles.
    if style.startswith("title"):
        return "TITLE"
    if style.startswith("subtitle"):
        return "SUBTITLE"

    # 7. Structural boundaries FIRST (they outrank heading style): References,
    #    Appendix, and rubric/back-matter regions.
    if references_span is not None and index in references_span:
        if _is_references_heading_text(text):
            return "REFERENCES_HEADING"
        return "REFERENCE_ENTRY"
    if appendix_span is not None and index in appendix_span:
        if _is_appendix_heading(para, text):
            return "APPENDIX_HEADING"
        return "APPENDIX_BODY"
    if rubric_region is not None and index >= rubric_region:
        return "UNKNOWN"  # rubric heading + content — back-matter, never BODY

    # 8. Headings (after structural boundaries).
    level = get_heading_level(para.get("style_name"))
    if level == 1:
        return "HEADING_1"
    if level == 2:
        return "HEADING_2"
    if level == 3:
        return "HEADING_3"
    if level is not None:  # level 4+ — outside the model, never guessed
        return "UNKNOWN"

    # 9. Cover region (only when proven — body-first docs stay BODY).
    if cover_end > 0 and index < cover_end:
        return "COVER"

    # 10. Lists.
    if para.get("is_list_item", False) or (paragraph_obj is not None and _style_chain_has_numbering(paragraph_obj)):
        return "LIST_ITEM"

    # 11. Ordinary body.
    if visible:
        return "BODY"
    return "UNKNOWN"


def _is_toc_heading(para: Dict) -> bool:
    style = (para.get("style_name") or "").lower()
    text = (para.get("text") or "").strip().lower()
    return "toc heading" in style or "table of contents" in text


def _is_manual_caption_text(text: str) -> bool:
    from app.config import settings

    return settings.PRESET.is_caption_text(text)


def _caption_role(para: Dict, p_elem, style: str) -> str:
    """CAPTION_TABLE vs CAPTION_FIGURE by label/SEQ object type."""
    text = para.get("text") or ""
    if re.search(r"^\s*(?:table|tab\.|jadual|表)\s*\d+", text, re.IGNORECASE):
        return "CAPTION_TABLE"
    if re.search(r"^\s*(?:figure|fig\.|chart|gambar|rajah|graf|图|图表)\s*\d+", text, re.IGNORECASE):
        return "CAPTION_FIGURE"
    # SEQ field label decides when visible text has no number (Word fields).
    if p_elem is not None:
        joined = "".join((t.text or "") for t in p_elem.iter(qn("w:instrText")))
        if re.search(r"\bSEQ\s+table\b", joined, re.IGNORECASE):
            return "CAPTION_TABLE"
        if re.search(r"\bSEQ\s+figure\b", joined, re.IGNORECASE):
            return "CAPTION_FIGURE"
    return "UNKNOWN"


# ---------------------------------------------------------------------------
# Table role classification (conservative, body-order aware)
# ---------------------------------------------------------------------------


def classify_table_roles(
    doc: Document,
    cover_end: Optional[int] = None,
    refs_start: Optional[int] = None,
    appendix_start: Optional[int] = None,
) -> List[str]:
    """Classify every doc.tables[i] into a table role using BODY ORDER.

    A table's position = the last body paragraph before it (in OOXML body
    order), so tables that interrupt the paragraph flow are placed exactly.
    Region membership is decided by that paragraph's index:
      - before the cover boundary → ADMINISTRATIVE_TABLE (cover info,
        assignment detail, and MARK/Comments assessment tables alike);
      - in References/Appendix back-matter → RUBRIC_TABLE when rubric-shaped
        (header keywords), else UNKNOWN;
      - manuscript body → SCHOLARLY_TABLE only with an adjacent numbered
        caption; never from structure alone. Ambiguity returns UNKNOWN.

    Boundary hints default to the SAME structural derivation as
    classify_paragraphs (cover signals, References span, Appendix span), so
    table and paragraph roles always agree on region membership.
    """
    paragraphs = extract_paragraphs(doc)
    body_seq = _body_sequence(doc)
    by_index = {p["index"]: p for p in paragraphs}

    if cover_end is None:
        boundary = _cover_boundary(by_index, body_seq)
        if _cover_signal_exists(doc, by_index, body_seq, boundary):
            cover_end = boundary if boundary is not None else len(paragraphs)
        else:
            cover_end = 0
    if refs_start is None:
        refs_start = _find_references_start(paragraphs)
    if appendix_start is None:
        appendix_span = _appendix_span_of(paragraphs)
        appendix_start = appendix_span.start if appendix_span is not None else None

    table_positions: List[int] = []
    for e in body_seq:
        if e["kind"] == "tbl":
            table_positions.append(e["before_para"] if e["before_para"] is not None else -1)

    roles: List[str] = []
    for idx, pos in enumerate(table_positions):
        # `pos` = index of the LAST body paragraph BEFORE the table; the
        # table sits after that paragraph, so region membership is decided
        # by pos (< boundary ⇒ before the boundary).
        if cover_end > 0 and pos < cover_end:
            roles.append("ADMINISTRATIVE_TABLE")
            continue
        if refs_start is not None and pos >= refs_start:
            roles.append(_rubric_or_unknown(doc, idx))
            continue
        if appendix_start is not None and pos >= appendix_start:
            roles.append(_rubric_or_unknown(doc, idx))
            continue
        # Manuscript-body region: scholarly ONLY with surrounding academic
        # evidence (adjacent numbered caption) — see _is_scholarly_table.
        roles.append(_scholarly_or_unknown(doc, idx))
    return roles


def _rubric_or_unknown(doc: Document, table_idx: int) -> str:
    """Back-matter table: RUBRIC only when rubric-shaped (header keywords) —
    conservative; otherwise UNKNOWN. Never a scholarly caption target."""
    table = doc.tables[table_idx]
    header = _first_row_text(table)
    if header and any(k in header.lower() for k in ("marks", "criteria", "rubric", "score", "grade")):
        return "RUBRIC_TABLE"
    return "UNKNOWN"


def _scholarly_or_unknown(doc: Document, table_idx: int) -> str:
    """Manuscript-body table: SCHOLARLY_TABLE only with reliable academic
    evidence — an adjacent numbered Table Caption OR surrounding academic
    body prose that refers to the Table ("as shown in Table 2", "the table
    below"). Structure alone never qualifies; ambiguity returns UNKNOWN."""
    if _has_adjacent_numbered_caption(doc, table_idx):
        return "SCHOLARLY_TABLE"
    if _has_in_text_table_reference(doc, table_idx):
        return "SCHOLARLY_TABLE"
    return "UNKNOWN"


def _has_in_text_table_reference(doc: Document, table_idx: int) -> bool:
    """True when nearby body paragraphs refer to the table in academic
    prose — a numbered in-text reference ("Table 2 shows", "see Table 2")
    or an explicit deictic reference ("the table below/above"). Evaluated
    only in the manuscript-body region (never cover/back-matter)."""
    children = list(doc.element.body.iterchildren())
    tbl_positions = [i for i, c in enumerate(children) if c.tag == qn("w:tbl")]
    if table_idx >= len(tbl_positions):
        return False
    pos = tbl_positions[table_idx]
    # Window of surrounding paragraphs (2 before / 3 after — a caption or
    # prose reference typically sits adjacent).
    window = []
    for delta in range(-2, 4):
        n = pos + delta
        if 0 <= n < len(children) and children[n].tag == qn("w:p"):
            text = "".join((t.text or "") for t in children[n].iter(qn("w:t")))
            window.append(text)
    for text in window:
        lowered = text.strip().lower()
        if not lowered:
            continue
        # Numbered in-text reference (not a caption line — captions are
        # already handled by _has_adjacent_numbered_caption).
        if re.search(r"\b(?:see|shown in|as (?:seen|shown) in|in)\s+table\s*\d+", lowered):
            return True
        if re.search(r"\bthe\s+table\s+(?:below|above)\b", lowered):
            return True
    return False


def _first_row_text(table) -> str:
    try:
        row = table.rows[0]
        return " ".join(c.text for c in row.cells)
    except Exception:
        return ""


def _has_adjacent_numbered_caption(doc: Document, table_idx: int) -> bool:
    """True when a paragraph adjacent to this table (in body order) is a
    numbered 'Table N' caption (semantic or manual)."""
    children = list(doc.element.body.iterchildren())
    tbl_positions = [i for i, c in enumerate(children) if c.tag == qn("w:tbl")]
    if table_idx >= len(tbl_positions):
        return False
    pos = tbl_positions[table_idx]
    for delta in (-1, 1):
        n = pos + delta
        if 0 <= n < len(children) and children[n].tag == qn("w:p"):
            text = "".join((t.text or "") for t in children[n].iter(qn("w:t")))
            if re.search(r"^\s*(?:table|tab\.|jadual|表)\s*\d+", text, re.IGNORECASE):
                return True
    return False
