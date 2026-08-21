from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from typing import List, Dict, Any, Optional
import io


# ---------------------------------------------------------------------------
# Effective-font resolution (style-hierarchy aware)
# ---------------------------------------------------------------------------
#
# python-docx `run.font.size` / `run.font.name` read ONLY the run's own
# <w:rPr>. When a run inherits 12 pt from the Normal style (the common real-
# document case), both return None and the deterministic font checks would
# silently skip every BODY paragraph. These resolvers walk the full OOXML
# inheritance chain instead:
#
#   1. direct run properties      w:rPr/w:sz | w:rPr/w:rFonts
#   2. character style chain      run.style + base styles
#   3. paragraph style chain      para.style + base styles
#   4. document defaults          w:styles/w:docDefaults/w:rPrDefault/w:rPr
#
# Resolution is cycle-safe (visited set), read-only (never mutates
# python-docx objects), and never invents values it cannot prove.

_FONT_SOURCE_DIRECT = "direct"
_FONT_SOURCE_CHAR_STYLE = "character_style"
_FONT_SOURCE_PARA_STYLE = "paragraph_style"
_FONT_SOURCE_STYLE_CHAIN = "style_chain"
_FONT_SOURCE_DOC_DEFAULT = "document_default"
_FONT_SOURCE_UNRESOLVED = "unresolved"

_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _sz_val_to_pt(sz_elm) -> Optional[float]:
    """w:sz val (half-points) -> points, or None when absent/invalid."""
    if sz_elm is None:
        return None
    raw = sz_elm.get(_W_NS + "val")
    if raw is None:
        return None
    try:
        half_points = int(raw)
    except ValueError:
        return None
    if half_points <= 0:
        return None
    return half_points / 2.0


def _rpr_sz(rPr) -> Optional[float]:
    if rPr is None:
        return None
    return _sz_val_to_pt(rPr.find(_W_NS + "sz"))


def _rpr_family(rPr) -> Optional[str]:
    """Read an explicit ascii/high-ANSI family from an rPr element.

    Only proven declarations are returned — theme fonts (w:asciiTheme)
    cannot be resolved to a concrete family name without the theme part,
    so they are deliberately NOT invented here.
    """
    if rPr is None:
        return None
    rfonts = rPr.find(_W_NS + "rFonts")
    if rfonts is None:
        return None
    for attr in ("w:ascii", "w:hAnsi"):
        val = rfonts.get(_W_NS + attr.split(":")[1])
        if val:
            return val
    return None


def _style_rpr(style) -> Optional[Any]:
    try:
        element = style.element
    except AttributeError:
        return None
    if element is None:
        return None
    return element.find(_W_NS + "rPr")


def _walk_style_chain(base_style):
    """Yield (style, source) pairs along the basedOn chain, cycle-safe.

    First yield is the style itself; subsequent yields are its ancestors.
    Source is 'character_style' for the first hop when the anchor was a
    character style and 'paragraph_style' when it was a paragraph style;
    deeper hops are 'style_chain'.
    """
    seen = set()
    current = base_style
    is_first = True
    while current is not None:
        style_id = id(current)
        if style_id in seen:  # cycle guard
            break
        seen.add(style_id)
        if is_first:
            # Caller distinguishes char vs para by the anchor type.
            yield current, None
            is_first = False
        else:
            yield current, _FONT_SOURCE_STYLE_CHAIN
        try:
            current = current.base_style
        except Exception:
            break


def resolve_effective_size(
    run,
    para,
    styles_element=None,
):
    """Resolve one run's effective font size in points.

    Returns (size_pt_or_None, source) where source is one of:
    direct | character_style | paragraph_style | style_chain |
    document_default | unresolved.
    """
    # 1. Direct run properties.
    size = _rpr_sz(run._element.rPr) if run._element.rPr is not None else None
    if size is not None:
        return size, _FONT_SOURCE_DIRECT

    # 2. Run character style + base chain.
    try:
        char_style = run.style
    except Exception:
        char_style = None
    if char_style is not None and getattr(char_style, "type", None) == WD_STYLE_TYPE.CHARACTER:
        for style, source in _walk_style_chain(char_style):
            if source is None:
                source = _FONT_SOURCE_CHAR_STYLE
            size = _rpr_sz(_style_rpr(style))
            if size is not None:
                return size, source

    # 3. Paragraph style + base chain.
    try:
        para_style = para.style if para is not None else None
    except Exception:
        para_style = None
    if para_style is not None:
        for style, source in _walk_style_chain(para_style):
            if source is None:
                source = _FONT_SOURCE_PARA_STYLE
            size = _rpr_sz(_style_rpr(style))
            if size is not None:
                return size, source

    # 4. Document defaults.
    if styles_element is not None:
        defaults = styles_element.find(_W_NS + "docDefaults")
        if defaults is not None:
            rpr_default = defaults.find(_W_NS + "rPrDefault")
            if rpr_default is not None:
                size = _rpr_sz(rpr_default.find(_W_NS + "rPr"))
                if size is not None:
                    return size, _FONT_SOURCE_DOC_DEFAULT

    return None, _FONT_SOURCE_UNRESOLVED


def resolve_effective_family(
    run,
    para,
    styles_element=None,
):
    """Resolve one run's effective font family name (same precedence)."""
    # 1. Direct run properties.
    family = _rpr_family(run._element.rPr) if run._element.rPr is not None else None
    if family:
        return family, _FONT_SOURCE_DIRECT

    # 2. Character style chain.
    try:
        char_style = run.style
    except Exception:
        char_style = None
    if char_style is not None and getattr(char_style, "type", None) == WD_STYLE_TYPE.CHARACTER:
        for style, source in _walk_style_chain(char_style):
            if source is None:
                source = _FONT_SOURCE_CHAR_STYLE
            family = _rpr_family(_style_rpr(style))
            if family:
                return family, source

    # 3. Paragraph style chain.
    try:
        para_style = para.style if para is not None else None
    except Exception:
        para_style = None
    if para_style is not None:
        for style, source in _walk_style_chain(para_style):
            if source is None:
                source = _FONT_SOURCE_PARA_STYLE
            family = _rpr_family(_style_rpr(style))
            if family:
                return family, source

    # 4. Document defaults.
    if styles_element is not None:
        defaults = styles_element.find(_W_NS + "docDefaults")
        if defaults is not None:
            rpr_default = defaults.find(_W_NS + "rPrDefault")
            if rpr_default is not None:
                family = _rpr_family(rpr_default.find(_W_NS + "rPr"))
                if family:
                    return family, _FONT_SOURCE_DOC_DEFAULT

    return None, _FONT_SOURCE_UNRESOLVED


def parse_document(file_bytes: bytes) -> Document:
    """Load .docx from bytes into python-docx Document."""
    return Document(io.BytesIO(file_bytes))


def extract_paragraphs(doc: Document) -> List[Dict[str, Any]]:
    """Extract paragraph-level formatting info.

    Adds `is_list_item` flag — True when the paragraph carries a <w:numPr>
    element (i.e. it's a numbered or bulleted list item). Used by the
    alignment rule to exempt list items from the justify requirement.

    Each run additionally carries effective font resolution:
      - font_size / font_size_source
      - font_name / font_name_source
    resolved through the OOXML style hierarchy (direct → character style →
    paragraph style chain → document defaults). `None` means genuinely
    unresolved — never "matches the requirement".
    """
    from docx.oxml.ns import qn

    styles_element = doc.styles.element

    paragraphs = []
    for idx, para in enumerate(doc.paragraphs):
        runs_info = []
        for run_idx, run in enumerate(para.runs):
            size, size_source = resolve_effective_size(run, para, styles_element)
            family, family_source = resolve_effective_family(run, para, styles_element)
            runs_info.append({
                "index": run_idx,
                "text": run.text,
                # Keep the original direct-only fields for backward
                # compatibility with any legacy consumer.
                "font_name": run.font.name,
                "font_size": run.font.size.pt if run.font.size else None,
                "bold": run.font.bold,
                "italic": run.font.italic,
                "underline": run.font.underline,
                # Effective (style-hierarchy-resolved) values + provenance.
                "effective_font_size": size,
                "font_size_source": size_source,
                "effective_font_name": family,
                "font_name_source": family_source,
            })

        pf = para.paragraph_format

        # Detect list item — walk the paragraph's <w:pPr> for <w:numPr>
        is_list_item = False
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            if pPr.find(qn("w:numPr")) is not None:
                is_list_item = True

        paragraphs.append({
            "index": idx,
            "text": para.text,
            "style_name": para.style.name if para.style else None,
            "alignment": para.alignment,
            "line_spacing": pf.line_spacing,
            "space_before": pf.space_before.pt if pf.space_before else None,
            "space_after": pf.space_after.pt if pf.space_after else None,
            "is_list_item": is_list_item,
            "runs": runs_info,
        })
    return paragraphs


def extract_sections(doc: Document) -> List[Dict[str, Any]]:
    """Extract page margin info + OOXML boundary metadata from sections.

    Each section is described by the `w:sectPr` that CLOSES it:
      - a paragraph-level `w:pPr/w:sectPr` ends the current section at that
        paragraph (zero-based `end_paragraph_index`);
      - the final body-level `w:sectPr` closes the last section
        (`end_paragraph_index=None` — it runs to the end of the document);
      - the first paragraph-level `w:sectPr` opens section 0, so each
        section's `start_paragraph_index` is the paragraph right after the
        previous section's closing sectPr (0 for the first section).

    `section_index` is zero-based and follows document traversal order —
    never inferred from the section count alone. Boundaries are nullable
    (historical/odd documents may lack a paragraph-level sectPr).
    """
    from docx.oxml.ns import qn

    # --- build the ordered sectPr list (paragraph-level first, body last) ---
    sect_prs: List[Dict[str, Any]] = []
    for idx, para in enumerate(doc.paragraphs):
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            sect = pPr.find(qn("w:sectPr"))
            if sect is not None:
                sect_prs.append({"sectPr": sect, "para_index": idx})
    final_sect = doc.element.body.find(qn("w:sectPr"))
    if final_sect is not None:
        sect_prs.append({"sectPr": final_sect, "para_index": None})

    # --- derive per-section boundaries from traversal order ---
    # start index = paragraph after the previous sectPr (0 for the first).
    sections = []
    for i, entry in enumerate(sect_prs):
        sect = entry["sectPr"]
        prev_end = sect_prs[i - 1]["para_index"] if i > 0 else None
        start = None if prev_end is None else prev_end + 1
        if i == 0:
            start = 0
        sections.append({
            "section_index": i,
            "start_paragraph_index": start,
            "end_paragraph_index": entry["para_index"],
            "break_type": _sect_break_type(sect),
            "page_width": _emu_to_inches(sect, qn("w:pgSz"), "w:w"),
            "page_height": _emu_to_inches(sect, qn("w:pgSz"), "w:h"),
            "margin_left": _emu_to_inches(sect, qn("w:pgMar"), "w:left"),
            "margin_right": _emu_to_inches(sect, qn("w:pgMar"), "w:right"),
            "margin_top": _emu_to_inches(sect, qn("w:pgMar"), "w:top"),
            "margin_bottom": _emu_to_inches(sect, qn("w:pgMar"), "w:bottom"),
        })

    # Fallback for documents with no sectPr at all: a single section.
    if not sections:
        sections.append({
            "section_index": 0,
            "start_paragraph_index": 0,
            "end_paragraph_index": None,
            "break_type": "nextPage",
            "page_width": _emu_to_inches(doc.element.body.find(qn("w:sectPr")), qn("w:pgSz"), "w:w"),
            "page_height": _emu_to_inches(doc.element.body.find(qn("w:sectPr")), qn("w:pgSz"), "w:h"),
            "margin_left": None,
            "margin_right": None,
            "margin_top": None,
            "margin_bottom": None,
        })
    return sections


def _sect_break_type(sect) -> str:
    """w:type value normalized: nextPage | continuous | oddPage | evenPage."""
    from docx.oxml.ns import qn
    if sect is None:
        return "nextPage"
    typ = sect.find(qn("w:type"))
    val = typ.get(qn("w:val")) if typ is not None else None
    if not val:
        return "nextPage"  # OOXML default is nextPage
    return val


def _emu_to_inches(sect, parent_tag: str, attr: str) -> Optional[float]:
    """Read a twips dimension as inches, or None when absent.

    OOXML stores `w:pgSz` (page size) and `w:pgMar` (margins) in twips
    (twentieths of a point; 1 twip = 1/1440 inch). This matches the
    python-docx Section accessors (left_margin.inches etc.).
    """
    from docx.oxml.ns import qn
    if sect is None:
        return None
    parent = sect.find(parent_tag)
    if parent is None:
        return None
    val = parent.get(qn(attr))
    if val is None:
        return None
    try:
        return int(val) / 1440.0  # twips → inches
    except (TypeError, ValueError):
        return None


def extract_tables(doc: Document) -> List[Dict[str, Any]]:
    """Extract table info with surrounding context for caption detection."""
    tables = []
    for idx, table in enumerate(doc.tables):
        # Get preceding/following paragraphs for caption detection
        prev_para = None
        next_para = None
        # python-docx doesn't directly give table position, use element traversal
        tbl_element = table._tbl
        prev_elem = tbl_element.getprevious()
        next_elem = tbl_element.getnext()

        if prev_elem is not None and prev_elem.tag.endswith('}p'):
            from docx.oxml.ns import qn
            if prev_elem.find(qn('w:pPr')) is not None:
                pass  # Could extract text here

        tables.append({
            "index": idx,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "alignment": table.alignment,
        })
    return tables


def extract_inline_shapes(doc: Document) -> List[Dict[str, Any]]:
    """Extract inline shapes (images) info."""
    shapes = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            shapes.append({
                "type": "image",
                "content_type": rel.target_ref,
            })
    return shapes


def get_heading_level(style_name: Optional[str]) -> Optional[int]:
    """Determine heading level from style name."""
    if not style_name:
        return None
    style_lower = style_name.lower()
    if style_lower.startswith("heading"):
        try:
            return int(style_lower.replace("heading", "").strip())
        except ValueError:
            pass
    return None


def iter_paragraphs_with_context(doc: Document):
    """Yield (paragraph, prev_para, next_para) for caption detection."""
    paragraphs = doc.paragraphs
    for i, para in enumerate(paragraphs):
        prev = paragraphs[i - 1] if i > 0 else None
        nxt = paragraphs[i + 1] if i < len(paragraphs) - 1 else None
        yield para, prev, nxt


def extract_document_blocks(doc: Document) -> List[Dict[str, Any]]:
    """Ordered paragraph-only blocks for the Evidence-Linked Document Preview.

    Derived from extract_paragraphs so each block's `order`/`index` align
    exactly with the paragraph_index values findings already carry. Empty
    paragraphs are preserved for index fidelity. No run-level data, tables,
    images, page numbers, HTML, file bytes, or raw XML. This is a preview
    surface — document text must never be logged.

    Each block carries a `role` from the authoritative role classifier
    (Phase 1 PoC — additive, nullable; historical blocks lack it).
    """
    from app.services.role_classifier import classify_paragraphs

    paragraphs = extract_paragraphs(doc)
    roles = classify_paragraphs(doc, paragraphs)
    blocks = []
    for para, role in zip(paragraphs, roles):
        blocks.append({
            "order": para["index"],
            "type": "paragraph",
            "index": para["index"],
            "text": para["text"],
            "style_name": para["style_name"],
            "heading_level": get_heading_level(para["style_name"]),
            "role": role,
        })
    return blocks


def extract_document_stats(doc: Document) -> Dict[str, int]:
    """Compute document-level statistics for the dashboard hero panel.

    Replaces the frontend mammoth-based reparse (stats.ts). The backend
    already walks the doc during the rules engine pass — computing stats
    here is essentially free and keeps the frontend pure-render.
    """
    paragraphs = doc.paragraphs
    word_count = sum(
        len(p.text.split()) for p in paragraphs if p.text and p.text.strip()
    )
    heading_count = sum(
        1 for p in paragraphs
        if p.style and get_heading_level(p.style.name) is not None
    )
    image_count = sum(
        1 for rel in doc.part.rels.values() if "image" in rel.target_ref
    )
    return {
        "paragraphs": len(paragraphs),
        "headings": heading_count,
        "tables": len(doc.tables),
        "images": image_count,
        "sections": len(doc.sections),
        "words": word_count,
    }
