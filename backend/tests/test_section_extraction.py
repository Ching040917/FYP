"""Section boundary metadata extraction tests (PoC).

Validates the OOXML-derived boundaries in extract_sections:
- paragraph-level w:pPr/w:sectPr ends the current section at that paragraph;
- the final body-level w:sectPr closes the final section (end = None);
- zero-based section/paragraph identities preserved in traversal order;
- break_type normalized (nextPage/continuous/oddPage/evenPage);
- page size and margins read as inches from twips;
- boundaries never inferred from section count alone;
- no document text or paths exposed.
"""
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.services.document_parser import extract_sections


def _add_section_break(doc, para_index, break_type="nextPage"):
    """Insert a paragraph-level sectPr on the paragraph at para_index."""
    p = doc.paragraphs[para_index]
    sect = OxmlElement("w:sectPr")
    if break_type != "nextPage":
        typ = OxmlElement("w:type")
        typ.set(qn("w:val"), break_type)
        sect.append(typ)
    p._p.get_or_add_pPr().append(sect)
    return sect


def test_single_section_no_breaks():
    doc = Document()
    doc.add_paragraph("alpha")
    doc.add_paragraph("beta")
    sections = extract_sections(doc)
    assert len(sections) == 1
    s = sections[0]
    assert s["section_index"] == 0
    assert s["start_paragraph_index"] == 0
    assert s["end_paragraph_index"] is None  # final body sectPr
    assert s["break_type"] == "nextPage"
    # default python-docx template: 8.5x11 letter
    assert abs(s["page_width"] - 8.5) < 0.01
    assert abs(s["page_height"] - 11.0) < 0.01


def test_two_sections_paragraph_level_boundaries():
    doc = Document()
    doc.add_paragraph("alpha")     # paragraph 0 — section 0
    doc.add_section()              # inserts a paragraph-level sectPr (ends section 0)
    doc.add_paragraph("gamma")     # paragraph 2 — section 1 (final body sectPr)
    sections = extract_sections(doc)
    assert len(sections) == 2
    assert sections[0]["section_index"] == 0
    assert sections[0]["start_paragraph_index"] == 0
    assert sections[0]["end_paragraph_index"] == 1  # the empty break paragraph
    assert sections[1]["section_index"] == 1
    assert sections[1]["start_paragraph_index"] == 2
    assert sections[1]["end_paragraph_index"] is None


def test_zero_based_identities_preserved():
    doc = Document()
    doc.add_paragraph("a")
    doc.add_section()  # break para 1 — continuous
    doc.add_paragraph("b")
    doc.add_section()  # break para 3 — oddPage
    doc.add_paragraph("c")
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _Oxml
    sect1 = doc.paragraphs[1]._p.pPr.find(_qn("w:sectPr"))
    typ = _Oxml("w:type")
    typ.set(_qn("w:val"), "continuous")
    sect1.append(typ)
    sect2 = doc.paragraphs[3]._p.pPr.find(_qn("w:sectPr"))
    typ = _Oxml("w:type")
    typ.set(_qn("w:val"), "oddPage")
    sect2.append(typ)
    sections = extract_sections(doc)
    assert [s["section_index"] for s in sections] == [0, 1, 2]
    # break_type is the type of the sectPr that CLOSES the section:
    # section 0 ends at the continuous break (para 1), section 1 ends at
    # the oddPage break (para 3), section 2 is closed by the body sectPr.
    assert [s["break_type"] for s in sections] == ["continuous", "oddPage", "nextPage"]
    assert sections[1]["start_paragraph_index"] == 2
    assert sections[1]["end_paragraph_index"] == 3


def test_break_types_odd_even_continuous():
    doc = Document()
    doc.add_paragraph("a")
    doc.add_section()  # break para 1 — continuous
    doc.add_paragraph("b")
    doc.add_section()  # break para 3 — evenPage
    doc.add_paragraph("c")
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _Oxml
    sect1 = doc.paragraphs[1]._p.pPr.find(_qn("w:sectPr"))
    typ = _Oxml("w:type")
    typ.set(_qn("w:val"), "continuous")
    sect1.append(typ)
    sect2 = doc.paragraphs[3]._p.pPr.find(_qn("w:sectPr"))
    typ = _Oxml("w:type")
    typ.set(_qn("w:val"), "evenPage")
    sect2.append(typ)
    sections = extract_sections(doc)
    # break_type of the sectPr that closes each section
    assert [s["break_type"] for s in sections] == ["continuous", "evenPage", "nextPage"]


def test_page_size_and_margins_in_inches():
    doc = Document()
    sec0 = doc.sections[0]
    sec0.page_width = Inches(8.5)
    sec0.page_height = Inches(11.0)
    sec0.left_margin = Inches(2.0)
    sec0.right_margin = Inches(1.25)
    sec0.top_margin = Inches(1.0)
    sec0.bottom_margin = Inches(0.9)
    doc.add_paragraph("alpha")
    sec1 = doc.add_section()
    sec1.page_width = Inches(11.0)  # landscape
    sec1.page_height = Inches(8.5)
    sec1.left_margin = Inches(1.0)
    doc.add_paragraph("beta")
    s0, s1 = extract_sections(doc)
    assert abs(s0["page_width"] - 8.5) < 0.01
    assert abs(s0["page_height"] - 11.0) < 0.01
    assert abs(s0["margin_left"] - 2.0) < 0.01
    assert abs(s0["margin_right"] - 1.25) < 0.01
    assert abs(s0["margin_top"] - 1.0) < 0.01
    assert abs(s0["margin_bottom"] - 0.9) < 0.01
    assert abs(s1["page_width"] - 11.0) < 0.01
    assert abs(s1["page_height"] - 8.5) < 0.01
    assert abs(s1["margin_left"] - 1.0) < 0.01


def test_no_document_text_or_paths_exposed():
    doc = Document()
    doc.add_paragraph("SECRET TEXT alpha")
    _add_section_break(doc, 0)
    doc.add_section()
    doc.add_paragraph("SECRET TEXT beta")
    for s in extract_sections(doc):
        blob = str(s)
        assert "SECRET" not in blob
        assert ".docx" not in blob
        assert "C:" not in blob
        assert "tmp" not in blob


def test_boundaries_not_inferred_from_count():
    # A document with 2 sections but NO paragraph-level sectPr (only the
    # final body sectPr) must produce ONE section with end=None — never two
    # sections guessed from a section count.
    doc = Document()
    doc.add_paragraph("only paragraph")
    # remove any paragraph-level sectPr python-docx may have inserted
    for p in doc.paragraphs:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is not None:
            sect = pPr.find(qn("w:sectPr"))
            if sect is not None:
                pPr.remove(sect)
    sections = extract_sections(doc)
    assert len(sections) == 1
    assert sections[0]["end_paragraph_index"] is None
