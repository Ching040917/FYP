"""Engine tests — Font Size / Font Family deterministic checks with
effective (style-hierarchy-resolved) run values, LIST_ITEM body-size
eligibility, and paragraph-level grouped findings."""
import io

from docx import Document
from docx.shared import Pt

from app.services.document_parser import parse_document, extract_paragraphs
from app.services.layout_engine import check_font_size, check_font_consistency
from app.services.role_classifier import classify_paragraphs


def _build_doc(builder):
    doc = Document()
    builder(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return parse_document(buf.getvalue())


def _normal_12(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)


class ExactPreset:
    """Institution-style exact requirement: TNR 21.5 pt body + headings."""
    profile_label = '"ecfawf" custom profile'
    FONT_FAMILY = "Times New Roman"
    FONT_SIZE_BODY = 21.5
    FONT_SIZE_H1 = None
    FONT_SIZE_H2 = None
    FONT_SIZE_H3 = None
    HEADING_FONT_FAMILY = "Times New Roman"
    HEADING_FONT_SIZE_PT = 21.5
    BODY_ALLOWED_FONT_COMBOS = ()
    HEADING_ALLOWED_FONT_COMBOS = ()

    def heading_expected_size(self, level):
        return self.HEADING_FONT_SIZE_PT

    def is_font_pair_allowed(self, family, size, heading=False):
        if family is None or size is None:
            return True
        return family.lower() == "times new roman" and abs(size - 21.5) < 0.01


def _repro_doc():
    def build(doc):
        _normal_12(doc)
        h1 = doc.styles["Heading 1"]

        doc.add_paragraph("University Name")            # 0 cover-ish
        p = doc.add_paragraph()                          # 1 H1 direct 12pt
        r = p.add_run("CHAPTER 1  INTRODUCTION")
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        p.style = h1
        doc.add_paragraph("First inherited body paragraph.")   # 2 BODY
        doc.add_paragraph("Second inherited body paragraph.")  # 3 BODY
        lp = doc.add_paragraph("A bullet item", style="List Bullet")  # 4 LIST_ITEM

    return build


def test_body_inherited_size_creates_finding():
    """The confirmed defect: style-inherited BODY 12 pt vs required 21.5 pt."""
    doc = _build_doc(_repro_doc())
    paras = extract_paragraphs(doc)
    roles = classify_paragraphs(doc, paras)
    violations = check_font_size(paras, ExactPreset(), roles=roles)

    body_findings = [v for v in violations
                     if v.location.get("paragraph_index") in (2, 3)]
    assert len(body_findings) == 2
    for v in body_findings:
        assert v.rule_code == "FONT_SIZE"
        assert v.severity == "MINOR"
        assert v.actual_value == "12pt"
        assert "21.5" in v.expected_value


def test_list_item_inherited_size_uses_body_requirement():
    doc = _build_doc(_repro_doc())
    paras = extract_paragraphs(doc)
    roles = classify_paragraphs(doc, paras)
    violations = check_font_size(paras, ExactPreset(), roles=roles)

    list_findings = [v for v in violations
                     if v.location.get("paragraph_index") == 4]
    assert len(list_findings) == 1
    v = list_findings[0]
    assert v.rule_code == "FONT_SIZE"
    assert v.severity == "MINOR"          # list items are not headings
    assert v.actual_value == "12pt"


def test_heading_direct_findings_unchanged():
    doc = _build_doc(_repro_doc())
    paras = extract_paragraphs(doc)
    roles = classify_paragraphs(doc, paras)
    violations = check_font_size(paras, ExactPreset(), roles=roles)

    h1 = [v for v in violations if v.location.get("paragraph_index") == 1]
    assert len(h1) == 1
    assert h1[0].severity == "MAJOR"
    assert h1[0].actual_value == "12pt"


def test_duplicate_identical_runs_in_one_paragraph_create_one_finding():
    def build(doc):
        _normal_12(doc)
        p = doc.add_paragraph()
        for _ in range(4):  # four identical inherited runs
            p.add_run("same size ")
        return None

    doc = _build_doc(build)
    paras = extract_paragraphs(doc)
    roles = classify_paragraphs(doc, paras)
    violations = check_font_size(paras, ExactPreset(), roles=roles)
    assert len(violations) == 1
    loc = violations[0].location
    assert loc["paragraph_index"] == 0
    assert loc["run_indexes"] == [0, 1, 2, 3]
    assert loc["run_index"] == 0


def test_mixed_non_compliant_values_produce_distinct_findings():
    def build(doc):
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(21.5)  # compliant base
        p = doc.add_paragraph()
        a = p.add_run("twelve ")
        a.font.size = Pt(12)
        b = p.add_run("fourteen ")
        b.font.size = Pt(14)
        c = p.add_run("twelve again ")
        c.font.size = Pt(12)
        ok = p.add_run("compliant")
        # ok inherits the compliant 21.5

    doc = _build_doc(build)
    paras = extract_paragraphs(doc)
    roles = classify_paragraphs(doc, paras)
    violations = check_font_size(paras, ExactPreset(), roles=roles)
    assert len(violations) == 2  # distinct actual values only
    actuals = sorted(v.actual_value for v in violations)
    assert actuals == ["12pt", "14pt"]
    # Grouped runs preserved: the two 12pt runs share one finding.
    twelve = next(v for v in violations if v.actual_value == "12pt")
    assert sorted(twelve.location["run_indexes"]) == [0, 2]


def test_compliant_runs_never_flagged():
    def build(doc):
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(21.5)
        doc.add_paragraph("fully compliant body text")

    doc = _build_doc(build)
    paras = extract_paragraphs(doc)
    roles = classify_paragraphs(doc, paras)
    assert check_font_size(paras, ExactPreset(), roles=roles) == []
    assert check_font_consistency(paras, ExactPreset(), roles=roles) == []


def test_different_paragraphs_are_never_merged():
    def build(doc):
        _normal_12(doc)
        doc.add_paragraph("alpha body paragraph")
        doc.add_paragraph("beta body paragraph")

    doc = _build_doc(build)
    paras = extract_paragraphs(doc)
    roles = classify_paragraphs(doc, paras)
    violations = check_font_size(paras, ExactPreset(), roles=roles)
    indexes = sorted(v.location["paragraph_index"] for v in violations)
    assert indexes == [0, 1]


def test_unresolved_values_are_not_fabricated():
    from app.services.layout_engine import LayoutViolation
    from app.services.document_parser import resolve_effective_size

    # A hand-built run dict with unresolved effective size must be skipped,
    # never reported as matching or given an invented value.
    paras = [{
        "index": 0,
        "style_name": "Normal",
        "runs": [
            {"index": 0, "text": "ghost",
             "effective_font_size": None, "font_size_source": "unresolved",
             "effective_font_name": None, "font_name_source": "unresolved"},
        ],
    }]
    roles = ["BODY"]
    assert check_font_size(paras, ExactPreset(), roles=roles) == []


def test_legacy_role_null_behavior_remains_compatible():
    """Without roles, the legacy style-heuristic path still flags non-heading
    paragraphs using effective sizes."""

    class LegacyPreset(ExactPreset):
        profile_label = None
        LINE_SPACING_BODY = None
        SPACE_BEFORE_BODY = None
        SPACE_AFTER_BODY = None
        ALIGNMENT_BODY = None

    def build(doc):
        _normal_12(doc)
        doc.add_paragraph("legacy body")

    doc = _build_doc(build)
    paras = extract_paragraphs(doc)
    violations = check_font_size(paras, LegacyPreset(), roles=None)
    assert len(violations) == 1
    assert violations[0].severity == "MINOR"


def test_apa_allowed_pairs_remain_pairwise_and_grouped():
    """APA-style profiles carry no explicit sizes; pair validation flows
    through check_font_consistency (no size gate), exactly as in
    production. Family+size must validate together."""

    class ApaLikePreset:
        profile_label = '"APA 7 Student Paper"'
        FONT_FAMILY = None
        FONT_SIZE_BODY = None
        FONT_SIZE_H1 = None
        FONT_SIZE_H2 = None
        FONT_SIZE_H3 = None
        HEADING_FONT_FAMILY = None
        BODY_ALLOWED_FONT_COMBOS = (("Times New Roman", 12.0), ("Calibri", 11.0))
        HEADING_ALLOWED_FONT_COMBOS = (("Times New Roman", 12.0), ("Calibri", 11.0))

        def heading_expected_size(self, level):
            return None

        def is_font_pair_allowed(self, family, size, heading=False):
            if family is None or size is None:
                return True
            pairs = (self.HEADING_ALLOWED_FONT_COMBOS if heading
                     else self.BODY_ALLOWED_FONT_COMBOS)
            return any(f.lower() == family.lower() and abs(s - size) < 0.01
                       for f, s in pairs)

    def build(doc):
        normal = doc.styles["Normal"]
        normal.font.name = "Georgia"     # valid family…
        normal.font.size = Pt(11)        # …but paired with another family's size
        p = doc.add_paragraph()
        p.add_run("pair mismatch one ")
        p.add_run("pair mismatch two ")

    doc = _build_doc(build)
    paras = extract_paragraphs(doc)
    roles = classify_paragraphs(doc, paras)
    violations = check_font_consistency(paras, ApaLikePreset(), roles=roles)
    assert len(violations) == 1
    v = violations[0]
    assert v.rule_code == "FONT_CONSISTENCY"
    assert "Georgia 11.0pt" in v.actual_value
    assert v.location["run_indexes"] == [0, 1]
    # A valid pair never flags.
    def build_ok(d2):
        n = d2.styles["Normal"]
        n.font.name = "Calibri"
        n.font.size = Pt(11)
        d2.add_paragraph("valid calibri eleven")
    doc_ok = _build_doc(build_ok)
    paras_ok = extract_paragraphs(doc_ok)
    roles_ok = classify_paragraphs(doc_ok, paras_ok)
    assert check_font_consistency(paras_ok, ApaLikePreset(), roles=roles_ok) == []


def test_role_exemptions_remain_intact():
    """Cover/TOC/reference-ENTRY/empty/caption/figure-host/UNKNOWN roles are
    never font-flagged. REFERENCES_HEADING is heading-like by policy and
    legitimately receives the heading font requirement."""

    def build(doc):
        _normal_12(doc)
        doc.add_paragraph("cover line")                      # COVER-ish
        doc.add_paragraph("Contents entry", style="TOC Heading")
        doc.add_paragraph("References")
        doc.add_paragraph("Author, A. (2020). Some title.")  # reference entry
        doc.add_paragraph("")                                 # empty
        cap = doc.add_paragraph("Table 1: results")

    doc = _build_doc(build)
    paras = extract_paragraphs(doc)
    roles = classify_paragraphs(doc, paras)
    violations = check_font_size(paras, ExactPreset(), roles=roles)
    flagged = {v.location["paragraph_index"] for v in violations}
    exempt_roles = {"COVER", "TABLE_OF_CONTENTS_HEADING", "TABLE_OF_CONTENTS_ENTRY",
                    "REFERENCE_ENTRY", "EMPTY", "FIELD_ONLY",
                    "DISPLAYED_EQUATION", "FIGURE_HOST",
                    "CAPTION_TABLE", "CAPTION_FIGURE", "UNKNOWN"}
    for i, role in enumerate(roles):
        if role in exempt_roles:
            assert i not in flagged, f"role {role} at {i} must stay exempt"
    # And heading-like roles (incl. REFERENCES_HEADING) ARE flagged — pin it.
    ref_heading = [i for i, r in enumerate(roles) if r == "REFERENCES_HEADING"]
    assert ref_heading, "fixture must produce a REFERENCES_HEADING"
    for i in ref_heading:
        assert i in flagged


def test_inherited_family_creates_consistency_finding():
    def build(doc):
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"   # inherited family ≠ required
        normal.font.size = Pt(12)
        doc.add_paragraph("inherited wrong family")

    doc = _build_doc(build)
    paras = extract_paragraphs(doc)
    roles = classify_paragraphs(doc, paras)
    violations = check_font_consistency(paras, ExactPreset(), roles=roles)
    assert len(violations) == 1
    v = violations[0]
    assert v.rule_code == "FONT_CONSISTENCY"
    assert v.expected_value == "Times New Roman"
    assert v.actual_value == "Calibri"


def test_exact_family_compliant_no_findings():
    def build(doc):
        _normal_12(doc)  # TNR 12 inherited; family matches, size does not
        doc.add_paragraph("family compliant")

    doc = _build_doc(build)
    paras = extract_paragraphs(doc)
    roles = classify_paragraphs(doc, paras)
    # Family matches -> no consistency findings even though size mismatches.
    assert check_font_consistency(paras, ExactPreset(), roles=roles) == []
