"""Tests for extract_document_blocks — the Evidence-Linked Document Preview
paragraph-block extractor (Build 8A).

Invariants under test:
- block order/index == zero-based paragraph index;
- empty paragraphs preserved for index fidelity;
- style names and heading levels extracted truthfully;
- output compatible with extract_paragraphs by construction.
"""
from docx import Document

from app.services.document_parser import (
    extract_document_blocks,
    extract_paragraphs,
    get_heading_level,
)


def _make_doc(paragraphs):
    """paragraphs: list of (text, style_name|None)."""
    doc = Document()
    for text, style in paragraphs:
        p = doc.add_paragraph(text)
        if style:
            p.style = doc.styles[style]
    return doc


def test_order_and_zero_based_indexes():
    doc = _make_doc([("alpha", None), ("beta", None), ("gamma", None)])
    blocks = extract_document_blocks(doc)
    assert [b["order"] for b in blocks] == [0, 1, 2]
    assert [b["index"] for b in blocks] == [0, 1, 2]
    assert [b["type"] for b in blocks] == ["paragraph", "paragraph", "paragraph"]


def test_text_extraction():
    doc = _make_doc([("First sentence.", None), ("Second sentence.", None)])
    blocks = extract_document_blocks(doc)
    assert [b["text"] for b in blocks] == ["First sentence.", "Second sentence."]


def test_empty_paragraphs_preserved_for_index_fidelity():
    doc = _make_doc([("first", None), ("", None), ("third", None)])
    blocks = extract_document_blocks(doc)
    assert len(blocks) == 3
    assert blocks[0]["text"] == "first"
    assert blocks[1]["text"] == ""
    assert blocks[2]["text"] == "third"
    assert blocks[2]["index"] == 2


def test_style_names():
    doc = _make_doc([("plain", None), ("title", "Title"), ("body", "Normal")])
    blocks = extract_document_blocks(doc)
    # python-docx default template applies "Normal" when no style is set.
    assert blocks[0]["style_name"] == "Normal"
    assert blocks[1]["style_name"] == "Title"
    assert blocks[2]["style_name"] == "Normal"


def test_heading_levels():
    doc = _make_doc([
        ("H1 text", "Heading 1"),
        ("H2 text", "Heading 2"),
        ("body", "Normal"),
    ])
    blocks = extract_document_blocks(doc)
    assert blocks[0]["heading_level"] == 1
    assert blocks[1]["heading_level"] == 2
    assert blocks[2]["heading_level"] is None


def test_duplicate_paragraph_text_keeps_distinct_blocks():
    doc = _make_doc([("same", None), ("same", None), ("same", None)])
    blocks = extract_document_blocks(doc)
    assert len(blocks) == 3
    assert [b["index"] for b in blocks] == [0, 1, 2]
    assert all(b["text"] == "same" for b in blocks)


def test_compatible_with_extract_paragraphs():
    doc = _make_doc([("one", None), ("", "Heading 1"), ("two", "Normal")])
    blocks = extract_document_blocks(doc)
    paragraphs = extract_paragraphs(doc)
    assert len(blocks) == len(paragraphs)
    for block, para in zip(blocks, paragraphs):
        assert block["index"] == para["index"]
        assert block["order"] == para["index"]
        assert block["text"] == para["text"]
        assert block["style_name"] == para["style_name"]
        assert block["heading_level"] == get_heading_level(para["style_name"])


def test_large_synthetic_document_2000_paragraphs():
    doc = Document()
    for i in range(2000):
        doc.add_paragraph(f"Paragraph number {i}")
    blocks = extract_document_blocks(doc)
    assert len(blocks) == 2000
    assert blocks[0]["index"] == 0
    assert blocks[1999]["index"] == 1999
    assert blocks[1999]["text"] == "Paragraph number 1999"
    assert [b["order"] for b in blocks] == list(range(2000))


def test_blank_document_returns_no_blocks():
    doc = Document()
    assert extract_document_blocks(doc) == []
