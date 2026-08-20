"""Document Role Integration Phase 2B — Table and Image Caption eligibility.

Role-gated: Caption findings apply ONLY to SCHOLARLY_TABLE and academic
Figures. Administrative cover tables, assignment/assessment/rubric tables,
layout tables, UNKNOWN tables, cover logos, repeated header images, and
decorative images are exempt — no scored finding. Alt Text stays independent
of Caption eligibility.
"""
import base64
import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

from app.services.layout_engine import run_static_rules_engine
from app.services.role_classifier import classify_table_roles

_PNG = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAA"
    "BJRU5ErkJggg=="
)


def _base_doc():
    doc = Document()
    for name in ("Normal", "Heading 1", "Caption"):
        try:
            s = doc.styles[name]
            s.font.name = "Times New Roman"
            s.font.size = Pt(12)
        except KeyError:
            pass
    return doc


def _add_image(doc, alt=None):
    p = doc.add_paragraph()
    p.add_run().add_picture(io.BytesIO(base64.b64decode(_PNG)), width=Pt(12))
    if alt:
        ns = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
        for docpr in p._p.findall(".//" + ns + "docPr"):
            docpr.set("descr", alt)
    return p


def _add_word_caption(doc, label, text):
    cap = doc.add_paragraph()
    cap.style = doc.styles["Caption"]
    cap.add_run(f"{label} ")
    for piece in (" SEQ ", f"{label} ", "\\* ARABIC "):
        run = cap.add_run()
        t = OxmlElement("w:instrText"); t.set(qn("xml:space"), "preserve"); t.text = piece
        run._r.append(t)
    cap.add_run("1")
    cap.add_run(text)
    return cap


def _codes(viols):
    return {v.rule_code for v in viols}


def _caption_codes(viols):
    return {v.rule_code for v in viols if "CAPTION" in v.rule_code or "ALT" in v.rule_code}


def _bytes(doc):
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Table roles
# ---------------------------------------------------------------------------

def test_admin_cover_table_no_caption_finding():
    """Cover-information / assignment-detail / assessment tables are
    ADMINISTRATIVE_TABLE → no Caption finding."""
    doc = _base_doc()
    doc.add_paragraph("Assignment Cover", style="Title")
    t0 = doc.add_table(rows=3, cols=2)
    for r, (k, v) in enumerate([("Name", "A. Student"), ("ID", "S1"), ("Programme", "BSc")]):
        t0.rows[r].cells[0].text = k
        t0.rows[r].cells[1].text = v
    doc.add_paragraph("Assessment")
    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "Mark"
    t1.rows[0].cells[1].text = "Comments"
    t1.rows[1].cells[0].text = "85"
    t1.rows[1].cells[1].text = "Good"
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body text that is long enough to be prose here.")
    viols = run_static_rules_engine(_bytes(doc))
    assert "TABLE_CAPTION_MISSING" not in _codes(viols)
    assert "MANUAL_CAPTION" not in _codes(viols)


def test_rubric_table_no_caption_finding():
    """Marking-rubric table after References → RUBRIC_TABLE → no finding."""
    doc = _base_doc()
    doc.add_paragraph("Body text that is long enough to be prose here.")
    doc.add_paragraph("References", style="Heading 1")
    doc.add_paragraph("Author, A. (2020). Title.")
    t = doc.add_table(rows=3, cols=3)
    t.rows[0].cells[0].text = "Criteria"
    t.rows[0].cells[1].text = "Marks"
    t.rows[0].cells[2].text = "Grade"
    t.rows[1].cells[0].text = "Clarity"
    t.rows[1].cells[1].text = "5"
    t.rows[1].cells[2].text = "A"
    viols = run_static_rules_engine(_bytes(doc))
    assert "TABLE_CAPTION_MISSING" not in _codes(viols)


def test_layout_table_no_caption_finding():
    """A decorative/layout table (single cell, no structure) → UNKNOWN → no
    finding."""
    doc = _base_doc()
    doc.add_paragraph("Body text that is long enough to be prose here.")
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "decorative"
    viols = run_static_rules_engine(_bytes(doc))
    assert "TABLE_CAPTION_MISSING" not in _codes(viols)


def test_ambiguous_table_unknown_no_finding():
    """Data-like body table with no caption and no in-text reference →
    UNKNOWN → no scored finding."""
    doc = _base_doc()
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body text that is long enough to be prose here.")
    t = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            t.rows[r].cells[c].text = f"d{r}{c}"
    viols = run_static_rules_engine(_bytes(doc))
    assert "TABLE_CAPTION_MISSING" not in _codes(viols)
    assert "MANUAL_CAPTION" not in _codes(viols)


def test_scholarly_table_with_semantic_caption_valid():
    """Scholarly table (adjacent numbered caption) with a VALID semantic
    caption → no finding."""
    doc = _base_doc()
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body text that is long enough to be prose here.")
    _add_word_caption(doc, "Table", ": Results")
    t = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            t.rows[r].cells[c].text = f"r{r}c{c}"
    viols = run_static_rules_engine(_bytes(doc))
    assert "TABLE_CAPTION_MISSING" not in _codes(viols)
    assert "MANUAL_CAPTION" not in _codes(viols)


def test_scholarly_table_with_missing_caption_finding():
    """Scholarly table (in-text reference) WITHOUT a caption →
    TABLE_CAPTION_MISSING."""
    doc = _base_doc()
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("The results, as shown in Table 1, are conclusive.")
    t = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            t.rows[r].cells[c].text = f"r{r}c{c}"
    viols = run_static_rules_engine(_bytes(doc))
    assert "TABLE_CAPTION_MISSING" in _codes(viols)


def test_scholarly_table_with_manual_caption_flagged():
    """Scholarly table with a manual 'Table N' caption → MANUAL_CAPTION."""
    doc = _base_doc()
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body text that is long enough to be prose here.")
    doc.add_paragraph("Table 1: Results")
    t = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            t.rows[r].cells[c].text = f"r{r}c{c}"
    viols = run_static_rules_engine(_bytes(doc))
    assert "MANUAL_CAPTION" in _codes(viols)
    assert "TABLE_CAPTION_MISSING" not in _codes(viols)


# ---------------------------------------------------------------------------
# Figure roles
# ---------------------------------------------------------------------------

def test_cover_logo_no_figure_caption_finding():
    """Cover logo (image in COVER region) → no IMAGE_CAPTION_MISSING."""
    doc = _base_doc()
    doc.add_paragraph("Assignment Cover Page", style="Title")
    _add_image(doc, alt="University logo")
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body text that is long enough to be prose here.")
    viols = run_static_rules_engine(_bytes(doc))
    assert "IMAGE_CAPTION_MISSING" not in _codes(viols)
    assert "MANUAL_CAPTION" not in _codes(viols)


def test_header_logo_no_figure_caption_finding():
    """Repeated header logo — hosted in the header part, never a body
    paragraph → no Figure Caption finding."""
    doc = _base_doc()
    hdr = doc.sections[0].header
    hdr.paragraphs[0].add_run().add_picture(
        io.BytesIO(base64.b64decode(_PNG)), width=Pt(12)
    )
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body text that is long enough to be prose here.")
    viols = run_static_rules_engine(_bytes(doc))
    assert "IMAGE_CAPTION_MISSING" not in _codes(viols)


def test_academic_figure_with_caption_valid():
    """Academic Figure (image in BODY region) with a semantic caption → no
    finding."""
    doc = _base_doc()
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body text that is long enough to be prose here.")
    _add_image(doc, alt="Architecture diagram")
    _add_word_caption(doc, "Figure", ": Architecture")
    viols = run_static_rules_engine(_bytes(doc))
    assert "IMAGE_CAPTION_MISSING" not in _codes(viols)
    assert "MANUAL_CAPTION" not in _codes(viols)


def test_academic_figure_without_caption_finding():
    """Academic Figure (image in BODY region) without a caption →
    IMAGE_CAPTION_MISSING."""
    doc = _base_doc()
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body text that is long enough to be prose here.")
    _add_image(doc, alt="Data chart")
    viols = run_static_rules_engine(_bytes(doc))
    assert "IMAGE_CAPTION_MISSING" in _codes(viols)


def test_ambiguous_image_no_caption_finding():
    """An image in a cover region → no Caption finding. Body-first documents
    have no cover (cover_end=0), so their images are academic — the region
    rule decides, never image dimensions or format."""
    doc = _base_doc()
    doc.add_paragraph("Cover Page", style="Title")
    _add_image(doc, alt="Cover branding")
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body text that is long enough to be prose here.")
    viols = run_static_rules_engine(_bytes(doc))
    assert "IMAGE_CAPTION_MISSING" not in _codes(viols)


# ---------------------------------------------------------------------------
# Alt Text independence
# ---------------------------------------------------------------------------

def test_caption_exemption_does_not_exempt_alt_text():
    """Academic-region image: Caption and Alt Text are INDEPENDENT — an
    academic figure without caption AND without alt text gets BOTH findings.
    Cover-region logo alt-text policy is UNRESOLVED (no profile field) → no
    scored Alt Text finding for the logo (no invented universal rule)."""
    # Academic figure without caption and without alt text → both findings.
    doc = _base_doc()
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body text that is long enough to be prose here.")
    _add_image(doc, alt=None)
    viols = run_static_rules_engine(_bytes(doc))
    assert "IMAGE_CAPTION_MISSING" in _codes(viols)
    assert "IMAGE_ALT_TEXT_MISSING" in _codes(viols)

    # Cover logo: no Caption finding (exempt), and no scored Alt Text
    # finding (logo policy unresolved in the profile — no universal rule).
    doc2 = _base_doc()
    doc2.add_paragraph("Assignment Cover Page", style="Title")
    _add_image(doc2, alt=None)
    doc2.add_paragraph("1. Introduction", style="Heading 1")
    doc2.add_paragraph("Body text that is long enough to be prose here.")
    viols2 = run_static_rules_engine(_bytes(doc2))
    assert "IMAGE_CAPTION_MISSING" not in _codes(viols2)
    assert "IMAGE_ALT_TEXT_MISSING" not in _codes(viols2)


def test_academic_figure_caption_and_alt_text_findings_separate():
    """Academic Figure without caption AND without alt text → both findings,
    independently."""
    doc = _base_doc()
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body text that is long enough to be prose here.")
    _add_image(doc, alt=None)
    viols = run_static_rules_engine(_bytes(doc))
    assert "IMAGE_CAPTION_MISSING" in _codes(viols)
    assert "IMAGE_ALT_TEXT_MISSING" in _codes(viols)


# ---------------------------------------------------------------------------
# Table roles via classifier (persisted policy consumption)
# ---------------------------------------------------------------------------

def test_classifier_roles_match_caption_policy():
    """classify_table_roles agrees with the Caption rule's gating."""
    doc = _base_doc()
    doc.add_paragraph("Cover", style="Title")
    admin = doc.add_table(rows=2, cols=2)
    admin.rows[0].cells[0].text = "Name"
    admin.rows[0].cells[1].text = "Value"
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body text that is long enough to be prose here.")
    _add_word_caption(doc, "Table", ": Results")
    scholarly = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            scholarly.rows[r].cells[c].text = f"r{r}c{c}"
    roles = classify_table_roles(doc)
    assert roles[0] == "ADMINISTRATIVE_TABLE"
    assert roles[1] == "SCHOLARLY_TABLE"
    # The Caption rule exempts non-scholarly roles.
    assert roles[0] != "SCHOLARLY_TABLE"
