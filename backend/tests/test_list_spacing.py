"""Rule-aware list-spacing eligibility tests.

SPACE_BEFORE/SPACE_AFTER must be list-aware:
  - when the preset is silent on list spacing (LIST_SPACE_AFTER is None,
    the default), list items — direct <w:numPr> OR numbering inherited
    through the style chain (List Bullet / List Number), including nested
    lists — are exempt from BOTH checks; no requirement is invented;
  - when LIST_SPACE_AFTER is configured, list SPACE_AFTER validates
    against that value (SPACE_BEFORE stays unchecked for list items);
  - ordinary body paragraphs keep the body preset checks unchanged.
"""
import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from app.config import PresetConfig
from app.services.layout_engine import run_static_rules_engine
from app.services.layout_violation import LayoutViolation
from typing import Optional


def _add_numpr(p, ilvl=0, num_id=1):
    """Attach a direct <w:numPr> to a paragraph (bulleted/numbered list item)."""
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    lvl = OxmlElement("w:ilvl")
    lvl.set(qn("w:val"), str(ilvl))
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    numPr.append(lvl)
    numPr.append(nid)
    pPr.append(numPr)


def _doc(direct=False, style=None, nested=False, space_after=3, include_body=False):
    """Build a docx with list items (and optionally an ordinary body paragraph).

    `direct`: attach <w:numPr> to each paragraph (direct list formatting).
    `style`: apply the named style (List Bullet / List Number — numbering
             inherited through the style chain, no direct numPr).
    `nested`: list items carry ilvl=1 (nested level).
    """
    doc = Document()
    # Ensure the list styles have TNR body font so no font-noise findings.
    for name in ("Normal", "List Bullet", "List Number"):
        try:
            s = doc.styles[name]
            s.font.name = "Times New Roman"
            s.font.size = Pt(12)
        except KeyError:
            pass

    texts = ["First list item", "Second list item", "Final list item"]
    for t in texts:
        p = doc.add_paragraph(t)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.5
        if direct:
            _add_numpr(p, ilvl=1 if nested else 0)
        elif style:
            try:
                p.style = doc.styles[style]
            except KeyError:
                pass
        # nested style case: apply style AND bump ilvl via direct numPr
        if nested and style:
            _add_numpr(p, ilvl=1)

    if include_body:
        body = doc.add_paragraph("Ordinary body paragraph.")
        body.paragraph_format.space_after = Pt(space_after)
        body.paragraph_format.line_spacing = 1.5

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _codes(viols):
    return {v.rule_code for v in viols}


def _space_after_at(viols, index):
    return [
        v for v in viols
        if v.rule_code == "SPACE_AFTER" and v.location.get("paragraph_index") == index
    ]


# ---------------------------------------------------------------------------
# Default preset (LIST_SPACE_AFTER = None): list items exempt from SPACE checks
# ---------------------------------------------------------------------------

def test_direct_numpr_list_exempt_from_space_checks():
    """Direct numPr list items with space_after=3 must NOT produce SPACE_AFTER."""
    viols = run_static_rules_engine(_doc(direct=True))
    assert "SPACE_AFTER" not in _codes(viols)
    assert "SPACE_BEFORE" not in _codes(viols)


def test_style_inherited_bullet_list_exempt_from_space_checks():
    """List Bullet style (numbering on the style, no direct numPr) is exempt."""
    viols = run_static_rules_engine(_doc(style="List Bullet"))
    assert "SPACE_AFTER" not in _codes(viols)
    assert "SPACE_BEFORE" not in _codes(viols)


def test_style_inherited_numbered_list_exempt_from_space_checks():
    viols = run_static_rules_engine(_doc(style="List Number"))
    assert "SPACE_AFTER" not in _codes(viols)
    assert "SPACE_BEFORE" not in _codes(viols)


def test_nested_direct_list_exempt_from_space_checks():
    """Nested (ilvl=1) list items are exempt too."""
    viols = run_static_rules_engine(_doc(direct=True, nested=True))
    assert "SPACE_AFTER" not in _codes(viols)


def test_nested_style_list_exempt_from_space_checks():
    """Style-inherited list with a nested level override is exempt."""
    viols = run_static_rules_engine(_doc(style="List Bullet", nested=True))
    assert "SPACE_AFTER" not in _codes(viols)


def test_body_paragraph_after_list_keeps_space_after_finding():
    """Regression: ordinary body spacing findings are NOT suppressed."""
    viols = run_static_rules_engine(_doc(direct=True, include_body=True, space_after=3))
    # body paragraph is index 3 (after the 3 list items)
    body_findings = _space_after_at(viols, 3)
    assert len(body_findings) == 1
    assert body_findings[0].expected_value == "6pt"
    # list items themselves stay exempt
    assert _space_after_at(viols, 0) == []


def test_body_space_after_correct_still_clean():
    """Body with space_after=6 (preset) produces no SPACE_AFTER finding."""
    viols = run_static_rules_engine(_doc(direct=True, include_body=True, space_after=6))
    assert _space_after_at(viols, 3) == []


# ---------------------------------------------------------------------------
# Configured LIST_SPACE_AFTER: list SPACE_AFTER validates against it
# ---------------------------------------------------------------------------

class _ConfiguredPreset(PresetConfig):
    """Real preset with LIST_SPACE_AFTER configured to 6pt."""
    LIST_SPACE_AFTER: Optional[float] = 6


def _run_with_preset(file_bytes, preset):
    """Run the full engine with a preset override (monkeypatch settings)."""
    import app.services.layout_engine as le
    from app.config import settings

    original = settings.PRESET
    settings.PRESET = preset
    try:
        return run_static_rules_engine(file_bytes)
    finally:
        settings.PRESET = original


def test_configured_list_space_after_ok_no_finding():
    """LIST_SPACE_AFTER=6 and list space_after=6 → no SPACE_AFTER."""
    viols = _run_with_preset(_doc(direct=True, space_after=6), _ConfiguredPreset())
    assert "SPACE_AFTER" not in _codes(viols)


def test_configured_list_space_after_wrong_produces_finding():
    """LIST_SPACE_AFTER=6 and list space_after=3 → SPACE_AFTER expected 6pt."""
    viols = _run_with_preset(_doc(direct=True, space_after=3), _ConfiguredPreset())
    findings = _space_after_at(viols, 0)
    assert len(findings) == 1
    assert findings[0].expected_value == "6pt"
    assert findings[0].actual_value == "3.0pt"


def test_configured_list_space_after_style_inherited_too():
    """Configured list spacing applies to style-inherited lists as well."""
    viols = _run_with_preset(_doc(style="List Bullet", space_after=3), _ConfiguredPreset())
    assert len(_space_after_at(viols, 0)) == 1


def test_configured_list_space_before_still_not_checked():
    """SPACE_BEFORE is never checked for list items, even when configured."""
    doc = Document()
    p = doc.add_paragraph("List item")
    p.paragraph_format.space_before = Pt(10)  # non-zero before
    p.paragraph_format.space_after = Pt(6)    # matches configured after
    _add_numpr(p)
    buf = io.BytesIO()
    doc.save(buf)
    viols = _run_with_preset(buf.getvalue(), _ConfiguredPreset())
    assert "SPACE_BEFORE" not in _codes(viols)
    assert "SPACE_AFTER" not in _codes(viols)


# ---------------------------------------------------------------------------
# Caption spacing (semantic + manual) — CAPTION_SPACE_BEFORE / _AFTER
# ---------------------------------------------------------------------------

def _add_word_caption(doc, label, text, space_before=3, space_after=3):
    """Real Word caption: Caption style + SEQ Table/Figure field."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    cap = doc.add_paragraph()
    cap.style = doc.styles["Caption"]
    cap.paragraph_format.space_before = Pt(space_before)
    cap.paragraph_format.space_after = Pt(space_after)
    cap.add_run(f"{label} ")
    for piece in (" SEQ ", f"{label} ", "\\* ARABIC "):
        run = cap.add_run()
        t = OxmlElement("w:instrText")
        t.set(qn("xml:space"), "preserve")
        t.text = piece
        run._r.append(t)
    cap.add_run("1")
    cap.add_run(text)
    return cap


def _caption_doc():
    """Body + semantic Figure caption + manual Table caption + body, all
    with caption-like spacing (3pt before / 3pt after)."""
    doc = Document()
    for name in ("Normal", "Caption"):
        try:
            s = doc.styles[name]
            s.font.name = "Times New Roman"
            s.font.size = Pt(12)
        except KeyError:
            pass
    body = doc.add_paragraph("Ordinary body text.")
    body.paragraph_format.space_after = Pt(6)
    _add_word_caption(doc, "Figure", ": Semantically captioned chart")
    manual = doc.add_paragraph("Table 1: Results")
    manual.paragraph_format.space_before = Pt(3)
    manual.paragraph_format.space_after = Pt(3)
    body2 = doc.add_paragraph("Ordinary body text after.")
    body2.paragraph_format.space_after = Pt(6)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _space_before_at(viols, index):
    return [
        v for v in viols
        if v.rule_code == "SPACE_BEFORE" and v.location.get("paragraph_index") == index
    ]


def test_semantic_caption_exempt_when_preset_silent():
    """Default preset (CAPTION_SPACE_* = None): semantic Figure caption with
    non-body spacing produces NO SPACE_BEFORE/SPACE_AFTER findings."""
    viols = run_static_rules_engine(_caption_doc())
    # semantic caption is paragraph index 1
    assert _space_before_at(viols, 1) == []
    assert _space_after_at(viols, 1) == []


def test_manual_caption_exempt_when_preset_silent():
    """Default preset: manual 'Table 1: Results' caption (index 2) exempt."""
    viols = run_static_rules_engine(_caption_doc())
    assert _space_before_at(viols, 2) == []
    assert _space_after_at(viols, 2) == []


def test_body_spacing_unchanged_with_captions_present():
    """Ordinary body paragraphs keep body spacing findings regardless of the
    caption exemptions."""
    doc = Document()
    for name in ("Normal", "Caption"):
        try:
            s = doc.styles[name]
            s.font.name = "Times New Roman"
            s.font.size = Pt(12)
        except KeyError:
            pass
    body = doc.add_paragraph("Body with wrong spacing.")
    body.paragraph_format.space_before = Pt(3)
    body.paragraph_format.space_after = Pt(3)
    _add_word_caption(doc, "Table", ": Captioned table")
    buf = io.BytesIO()
    doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    # body is index 0 → SPACE_BEFORE 0pt and SPACE_AFTER 6pt expectations
    assert len(_space_before_at(viols, 0)) == 1
    assert len(_space_after_at(viols, 0)) == 1
    # caption (index 1) exempt
    assert _space_before_at(viols, 1) == []
    assert _space_after_at(viols, 1) == []


class _CaptionConfiguredPreset(PresetConfig):
    """Real preset with explicit Caption spacing: 6pt before / 12pt after."""
    CAPTION_SPACE_BEFORE: Optional[float] = 6
    CAPTION_SPACE_AFTER: Optional[float] = 12


def test_configured_caption_spacing_wrong_produces_findings():
    """Explicit Caption spacing validates semantic + manual captions."""
    viols = _run_with_preset(_caption_doc(), _CaptionConfiguredPreset())
    # semantic caption (index 1): actual 3pt vs required 6pt / 12pt
    before = _space_before_at(viols, 1)
    after = _space_after_at(viols, 1)
    assert len(before) == 1
    assert before[0].expected_value == "6pt"
    assert len(after) == 1
    assert after[0].expected_value == "12pt"
    # manual caption (index 2) validated the same way
    assert len(_space_before_at(viols, 2)) == 1
    assert len(_space_after_at(viols, 2)) == 1


def test_configured_caption_spacing_ok_no_finding():
    """Explicit Caption spacing matched → no findings."""
    doc = Document()
    for name in ("Normal", "Caption"):
        try:
            s = doc.styles[name]
            s.font.name = "Times New Roman"
            s.font.size = Pt(12)
        except KeyError:
            pass
    cap = _add_word_caption(doc, "Figure", ": Correct spacing", space_before=6, space_after=12)
    buf = io.BytesIO()
    doc.save(buf)
    viols = _run_with_preset(buf.getvalue(), _CaptionConfiguredPreset())
    assert _space_before_at(viols, 0) == []
    assert _space_after_at(viols, 0) == []


def test_configured_caption_space_before_none_skips_before_only():
    """CAPTION_SPACE_BEFORE None skips SPACE_BEFORE; CAPTION_SPACE_AFTER
    configured still validates SPACE_AFTER."""

    class _PartialPreset(PresetConfig):
        CAPTION_SPACE_BEFORE: Optional[float] = None
        CAPTION_SPACE_AFTER: Optional[float] = 6

    viols = _run_with_preset(_caption_doc(), _PartialPreset())
    # semantic caption index 1: before skipped, after validated (3 vs 6)
    assert _space_before_at(viols, 1) == []
    assert len(_space_after_at(viols, 1)) == 1


def test_caption_style_alone_is_caption_for_spacing():
    """Caption style WITHOUT a SEQ field (visible 'Table 1' text) is a
    semantic caption for spacing eligibility."""
    doc = Document()
    for name in ("Normal", "Caption"):
        try:
            s = doc.styles[name]
            s.font.name = "Times New Roman"
            s.font.size = Pt(12)
        except KeyError:
            pass
    p = doc.add_paragraph("Table 1: Results")
    p.style = doc.styles["Caption"]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    buf = io.BytesIO()
    doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    assert _space_before_at(viols, 0) == []
    assert _space_after_at(viols, 0) == []
