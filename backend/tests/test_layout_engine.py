"""Tests for the static layout engine and scoring math."""
import io

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.layout_engine import run_static_rules_engine
from app.services.scoring import calculate_weighted_score
from app.services.layout_violation import LayoutViolation
from app.config import settings


# ---------------------------------------------------------------------------
# Scoring math
# ---------------------------------------------------------------------------

def test_calculate_weighted_score_clean_returns_100():
    assert calculate_weighted_score([]) == 100


def test_calculate_weighted_score_one_major_returns_85():
    v = LayoutViolation("X", "MAJOR", {}, "msg")
    assert calculate_weighted_score([v]) == 85


def test_calculate_weighted_score_one_minor_returns_97():
    v = LayoutViolation("X", "MINOR", {}, "msg")
    assert calculate_weighted_score([v]) == 97


def test_calculate_weighted_score_mixed_math():
    viols = [
        LayoutViolation("X", "MAJOR", {}, "m"),
        LayoutViolation("X", "MAJOR", {}, "m"),
        LayoutViolation("X", "MINOR", {}, "m"),
    ]
    # 100 - 15 - 15 - 3 = 67
    assert calculate_weighted_score(viols) == 67


def test_calculate_weighted_score_floors_at_zero():
    viols = [LayoutViolation("X", "MAJOR", {}, "m") for _ in range(8)]
    # 100 - 8*15 = -20, floor 0
    assert calculate_weighted_score(viols) == 0


def test_calculate_weighted_score_unknown_severity_ignored():
    v = LayoutViolation("X", "WEIRD", {}, "m")
    assert calculate_weighted_score([v]) == 100


# ---------------------------------------------------------------------------
# Layout engine — citation sensor wiring
# ---------------------------------------------------------------------------

def test_citation_mismatch_included_in_engine_output(docx_factory):
    body = [
        "Intro paragraph.",
        "Orphan (Garcia, 2018) text.",
    ]
    # no references header at all -> orphan flagged
    file_bytes = docx_factory(paragraphs=body, references=None)
    viols = run_static_rules_engine(file_bytes)
    codes = [v.rule_code for v in viols]
    assert "CITATION_MISMATCH" in codes
    cm = next(v for v in viols if v.rule_code == "CITATION_MISMATCH")
    assert cm.severity == "MAJOR"
    assert "Garcia" in cm.message


def test_present_citation_not_flagged(docx_factory):
    body = ["Smith (2020) wrote…"]
    refs = ["Smith, J. (2020). Title. Press."]
    file_bytes = docx_factory(paragraphs=body, references=refs)
    viols = run_static_rules_engine(file_bytes)
    codes = [v.rule_code for v in viols]
    assert "CITATION_MISMATCH" not in codes


# ---------------------------------------------------------------------------
# Layout engine — preset conformance
# ---------------------------------------------------------------------------

def test_wrong_margin_triggers_major(docx_factory):
    # SUC left = 1.5in; we set 1.0in
    body = ["Body text."]
    refs = []  # avoid citation noise
    file_bytes = docx_factory(
        paragraphs=body, references=refs, margins={"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0}
    )
    viols = run_static_rules_engine(file_bytes)
    codes = [v.rule_code for v in viols]
    assert "MARGIN_LEFT" in codes
    mv = next(v for v in viols if v.rule_code == "MARGIN_LEFT")
    assert mv.severity == "MAJOR"


def test_wrong_body_font_triggers_minor(docx_factory):
    body = ["Body text using Arial."]
    refs = []
    file_bytes = docx_factory(
        paragraphs=body, references=refs,
        margins={"left": 1.5, "right": 1.0, "top": 1.0, "bottom": 1.0},
        font_name="Arial",
    )
    viols = run_static_rules_engine(file_bytes)
    codes = [v.rule_code for v in viols]
    assert "FONT_CONSISTENCY" in codes
    fc = next(v for v in viols if v.rule_code == "FONT_CONSISTENCY")
    assert fc.severity == "MINOR"


def test_uncaptioned_table_triggers_minor(docx_factory):
    body = ["Body text."]
    refs = []
    file_bytes = docx_factory(
        paragraphs=body, references=refs,
        margins={"left": 1.5, "right": 1.0, "top": 1.0, "bottom": 1.0},
        tables=[[["A", "B"], ["1", "2"]]],
        with_caption=False,
    )
    viols = run_static_rules_engine(file_bytes)
    codes = [v.rule_code for v in viols]
    assert "TABLE_CAPTION_MISSING" in codes
    tc = next(v for v in viols if v.rule_code == "TABLE_CAPTION_MISSING")
    assert tc.severity == "MINOR"


# ---------------------------------------------------------------------------
# Semantic caption detection (FR-2.6)
# ---------------------------------------------------------------------------

def _add_seq_para(doc, label_text, seq_label, style=None, text_suffix=": Test"):
    """Caption paragraph with a SEQ field split across three instrText runs."""
    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style]
    p.add_run(f"{label_text} ")
    for piece in (" SEQ ", f"{seq_label} ", "\\* ARABIC "):
        run = p.add_run()
        t = OxmlElement("w:instrText")
        t.set(qn("xml:space"), "preserve")
        t.text = piece
        run._r.append(t)
    p.add_run("1")
    p.add_run(text_suffix)
    return p


def _add_list_para(doc, text):
    """Numbered list item paragraph (numPr in pPr)."""
    p = doc.add_paragraph(text)
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    return p


def _doc_with_table(caption_builder=None, below=False):
    """Body + optional caption + one 2x2 table."""
    doc = Document()
    doc.add_paragraph("Body text.")
    if caption_builder and not below:
        caption_builder(doc)
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "A"
    t.rows[1].cells[1].text = "B"
    if caption_builder and below:
        caption_builder(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_PNG_1PX = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8"
    "z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)


def _PNG_RGB(index):
    """Build a distinct 1x1 PNG (index -> different colour) in-memory."""
    import base64
    import struct
    import zlib

    def chunk(tag, data):
        c = tag + data
        return (
            struct.pack(">I", len(data)) + c
            + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        )

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    rgb = ((255, 0, 0), (0, 0, 255), (0, 128, 0), (128, 0, 128))[index % 4]
    raw = b"\x00" + bytes(rgb)
    return (
        b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw)) + chunk(b"IEND", b"")
    )


def _doc_with_image(caption_builder=None, below=True, alt_text="Test alt"):
    """Body + optional caption + one inline picture paragraph."""
    import base64
    doc = Document()
    doc.add_paragraph("Body text.")
    if caption_builder and not below:
        caption_builder(doc)
    para = doc.add_paragraph()
    para.add_run().add_picture(io.BytesIO(base64.b64decode(_PNG_1PX)), width=Inches(1))
    if alt_text:
        ns = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
        for docpr in para._p.findall(".//" + ns + "docPr"):
            docpr.set("descr", alt_text)
    if caption_builder and below:
        caption_builder(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _convert_inline_to_anchor(inline):
    """Convert one wp:inline drawing element into a wp:anchor (in place)."""
    wp = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
    a_ns = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
    drawing = inline.getparent()
    anchor = OxmlElement("wp:anchor")
    for attr, val in (
        ("distT", "0"), ("distB", "0"), ("distL", "114300"), ("distR", "114300"),
        ("simplePos", "0"), ("relativeHeight", "251658240"), ("behindDoc", "0"),
        ("locked", "0"), ("layoutInCell", "1"), ("allowOverlap", "1"),
    ):
        anchor.set(attr, val)
    for tag in ("simplePos", "positionH", "positionV"):
        el = OxmlElement("wp:" + tag)
        if tag == "simplePos":
            el.set("x", "0")
            el.set("y", "0")
        else:
            el.set("relativeFrom", "column" if tag == "positionH" else "paragraph")
            off = OxmlElement("wp:posOffset")
            off.text = "0"
            el.append(off)
        anchor.append(el)
    extent = inline.find(wp + "extent")
    if extent is not None:
        anchor.append(extent)
    eff = inline.find(wp + "effectExtent")
    if eff is not None:
        anchor.append(eff)
    anchor.append(OxmlElement("wp:wrapNone"))
    docpr = inline.find(wp + "docPr")
    if docpr is not None:
        anchor.append(docpr)
    cnv = inline.find(wp + "cNvGraphicFramePr")
    if cnv is not None:
        anchor.append(cnv)
    graphic = inline.find(a_ns + "graphic")
    if graphic is not None:
        anchor.append(graphic)
    drawing.replace(inline, anchor)


def _png_bytes(index):
    """Distinct 1x1 PNG per index — python-docx dedupes identical images."""
    import base64
    if index == 0:
        return base64.b64decode(_PNG_1PX)
    return _PNG_RGB(index)


def _doc_with_images(alt_texts, one_paragraph=False, anchored=False):
    """Build docx with len(alt_texts) images.

    `alt_texts`: per-image docPr@descr; None means "no alt text".
    `one_paragraph`: all images in a single paragraph (order preserved).
    `anchored`: use wp:anchor instead of wp:inline for every drawing.
    """
    doc = Document()
    doc.add_paragraph("Body text.")
    if one_paragraph:
        para = doc.add_paragraph()
        for i in range(len(alt_texts)):
            para.add_run().add_picture(
                io.BytesIO(_png_bytes(i)), width=Inches(1)
            )
    else:
        for i in range(len(alt_texts)):
            para = doc.add_paragraph()
            para.add_run().add_picture(
                io.BytesIO(_png_bytes(i)), width=Inches(1)
            )

    wp = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
    if anchored:
        for inline in doc.element.body.findall(".//" + wp + "inline"):
            _convert_inline_to_anchor(inline)

    docprs = doc.element.body.findall(".//" + wp + "docPr")
    for docpr, alt in zip(docprs, alt_texts):
        if alt is not None:
            docpr.set("descr", alt)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _alt_missing_indexes(viols):
    return sorted(
        v.location.get("image_index")
        for v in viols if v.rule_code == "IMAGE_ALT_TEXT_MISSING"
    )


def _caption_codes(viols):
    return {v.rule_code for v in viols if "CAPTION" in v.rule_code}


def test_semantic_table_caption_accepted(docx_factory):
    """Real Word table caption (Caption style + SEQ Table) is VALID."""
    file_bytes = docx_factory(
        paragraphs=["Body text."], references=[],
        tables=[[["A", "B"], ["1", "2"]]],
    )
    viols = run_static_rules_engine(file_bytes)
    assert "TABLE_CAPTION_MISSING" not in _caption_codes(viols)
    assert "MANUAL_CAPTION" not in _caption_codes(viols)


def test_semantic_figure_caption_accepted(docx_factory):
    """Real Word figure caption (Caption style + SEQ Figure, below) is VALID."""
    file_bytes = docx_factory(
        paragraphs=["Body text."], references=[],
        images=1,
    )
    viols = run_static_rules_engine(file_bytes)
    assert "IMAGE_CAPTION_MISSING" not in _caption_codes(viols)
    assert "MANUAL_CAPTION" not in _caption_codes(viols)


def test_manual_table_caption_flagged_minor():
    """Typed 'Table N' text without semantics -> MANUAL_CAPTION (MINOR)."""
    file_bytes = _doc_with_table(
        caption_builder=lambda doc: doc.add_paragraph("Table 1: Results"),
    )
    viols = run_static_rules_engine(file_bytes)
    codes = _caption_codes(viols)
    assert "TABLE_CAPTION_MISSING" not in codes
    assert "MANUAL_CAPTION" in codes
    mc = next(v for v in viols if v.rule_code == "MANUAL_CAPTION")
    assert mc.severity == "MINOR"
    assert "Insert Caption" in mc.message
    # The caption paragraph identity rides along so the frontend can
    # associate by document order instead of requiring the typed number to
    # match the table ordinal.
    assert mc.location.get("table_index") == 0
    assert isinstance(mc.location.get("paragraph_index"), int)
    assert mc.location["paragraph_index"] >= 0


def test_manual_table_caption_below_table_still_carries_caption_index():
    """Caption BELOW the table is associated the same way (both directions)."""
    file_bytes = _doc_with_table(
        caption_builder=lambda doc: doc.add_paragraph("Table 1: Results"),
        below=True,
    )
    viols = run_static_rules_engine(file_bytes)
    mc = next(v for v in viols if v.rule_code == "MANUAL_CAPTION")
    assert mc.location.get("table_index") == 0
    assert isinstance(mc.location.get("paragraph_index"), int)


def test_manual_figure_caption_flagged_minor():
    """Typed 'Figure N' text below an image without semantics -> MANUAL_CAPTION."""
    file_bytes = _doc_with_image(
        caption_builder=lambda doc: doc.add_paragraph("Figure 1: Architecture"),
    )
    viols = run_static_rules_engine(file_bytes)
    codes = _caption_codes(viols)
    assert "IMAGE_CAPTION_MISSING" not in codes
    assert "MANUAL_CAPTION" in codes


def test_wrong_label_type_not_accepted():
    """A Figure caption beside a table does not satisfy the table rule."""
    file_bytes = _doc_with_table(
        caption_builder=lambda doc: doc.add_paragraph("Figure 1: Wrong label"),
    )
    viols = run_static_rules_engine(file_bytes)
    codes = _caption_codes(viols)
    assert "TABLE_CAPTION_MISSING" in codes
    assert "MANUAL_CAPTION" not in codes


def test_semantic_caption_below_table_accepted():
    """Caption below the table (opposite of convention) is still VALID."""
    file_bytes = _doc_with_table(
        caption_builder=lambda doc: _add_seq_para(doc, "Table", "Table"),
        below=True,
    )
    viols = run_static_rules_engine(file_bytes)
    assert "TABLE_CAPTION_MISSING" not in _caption_codes(viols)
    assert "MANUAL_CAPTION" not in _caption_codes(viols)


def test_heading_resembling_caption_not_accepted():
    """Heading-styled 'Table N' text is never a valid caption."""
    file_bytes = _doc_with_table(
        caption_builder=lambda doc: _add_seq_para(
            doc, "Table", "Table", style="Heading 1", text_suffix=": Results"),
    )
    viols = run_static_rules_engine(file_bytes)
    codes = _caption_codes(viols)
    assert "MANUAL_CAPTION" in codes
    assert "TABLE_CAPTION_MISSING" not in codes


def test_list_item_resembling_caption_not_accepted():
    """Numbered list item 'Figure N' text is never a valid caption."""
    file_bytes = _doc_with_image(
        caption_builder=lambda doc: _add_list_para(doc, "Figure 1: Results"),
    )
    viols = run_static_rules_engine(file_bytes)
    codes = _caption_codes(viols)
    assert "MANUAL_CAPTION" in codes
    assert "IMAGE_CAPTION_MISSING" not in codes


def test_split_seq_field_runs_detected():
    """SEQ instruction split across runs still validates the caption."""
    file_bytes = _doc_with_table(
        caption_builder=lambda doc: _add_seq_para(doc, "Table", "Table"),
    )
    viols = run_static_rules_engine(file_bytes)
    assert "TABLE_CAPTION_MISSING" not in _caption_codes(viols)
    assert "MANUAL_CAPTION" not in _caption_codes(viols)


def test_caption_style_alone_is_valid():
    """Caption style without a SEQ field is a valid caption."""
    def builder(doc):
        p = doc.add_paragraph("Table 1: Test")
        p.style = doc.styles["Caption"]
    file_bytes = _doc_with_table(caption_builder=builder)
    viols = run_static_rules_engine(file_bytes)
    assert "TABLE_CAPTION_MISSING" not in _caption_codes(viols)
    assert "MANUAL_CAPTION" not in _caption_codes(viols)


def test_uncaptioned_image_triggers_minor(docx_factory):
    """Image without any adjacent caption -> IMAGE_CAPTION_MISSING."""
    file_bytes = docx_factory(
        paragraphs=["Body text."], references=[],
        images=1, with_image_caption=False,
    )
    viols = run_static_rules_engine(file_bytes)
    assert "IMAGE_CAPTION_MISSING" in _caption_codes(viols)


# ---------------------------------------------------------------------------
# Image alt-text — per-image docPr@descr resolution (defect fix)
# ---------------------------------------------------------------------------

def test_alt_text_first_image_only():
    """Only the image WITHOUT alt text is flagged — not its neighbour."""
    file_bytes = _doc_with_images(["First image alt", None])
    viols = run_static_rules_engine(file_bytes)
    assert _alt_missing_indexes(viols) == [1]


def test_alt_text_second_image_only():
    file_bytes = _doc_with_images([None, "Second image alt"])
    viols = run_static_rules_engine(file_bytes)
    assert _alt_missing_indexes(viols) == [0]


def test_alt_text_both_images_distinct():
    """Distinct alt text on both images -> no IMAGE_ALT_TEXT_MISSING."""
    file_bytes = _doc_with_images(["First alt", "Second alt"])
    viols = run_static_rules_engine(file_bytes)
    assert _alt_missing_indexes(viols) == []


def test_alt_text_neither_image():
    file_bytes = _doc_with_images([None, None])
    viols = run_static_rules_engine(file_bytes)
    assert _alt_missing_indexes(viols) == [0, 1]


def test_alt_text_multiple_images_one_paragraph():
    """Multiple images in a single paragraph still resolve individually."""
    file_bytes = _doc_with_images([None, "Second alt"], one_paragraph=True)
    viols = run_static_rules_engine(file_bytes)
    assert _alt_missing_indexes(viols) == [0]


def test_alt_text_anchored_images():
    """Anchored drawings resolve alt text per image too."""
    file_bytes = _doc_with_images(["First alt", None], anchored=True)
    viols = run_static_rules_engine(file_bytes)
    assert _alt_missing_indexes(viols) == [1]


def test_alt_text_anchored_images_both_have_alt():
    file_bytes = _doc_with_images(["A", "B"], anchored=True)
    viols = run_static_rules_engine(file_bytes)
    assert _alt_missing_indexes(viols) == []


def test_image_index_stable_across_rules():
    """image_index ordering matches for alt-text and caption findings."""
    file_bytes = _doc_with_images(["Alt", None])
    viols = run_static_rules_engine(file_bytes)
    alt = _alt_missing_indexes(viols)
    cap = sorted(
        v.location["image_index"]
        for v in viols if v.rule_code == "IMAGE_CAPTION_MISSING"
    )
    assert alt == [1]
    assert cap == [0, 1]
