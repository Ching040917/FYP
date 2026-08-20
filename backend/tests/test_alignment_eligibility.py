"""Rule-aware ALIGNMENT eligibility tests.

ALIGNMENT must apply only to eligible visible-text paragraphs. Empty,
field-only, image-only, caption, Title/Subtitle, list (direct and style-
inherited), and References-section paragraphs are skipped; ordinary body
text (justify valid; center/left findings) and headings (preset alignment)
keep validating. Mixed text-and-image paragraphs validate when visible
text exists.
"""
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from app.services.layout_engine import run_static_rules_engine
from app.services.layout_violation import LayoutViolation

# 1x1 px transparent PNG
_PNG = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAA"
    "BJRU5ErkJggg=="
)


def _doc(blocks):
    """Build a docx from (text, style, align, kind) tuples.

    kind: 'text' | 'image' | 'field' | 'direct-list'
    """
    import base64

    doc = Document()
    for name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Caption", "Title", "Subtitle", "List Bullet"):
        try:
            s = doc.styles[name]
            s.font.name = "Times New Roman"
            s.font.size = Pt(12)
        except KeyError:
            pass

    for text, style, align, kind in blocks:
        p = doc.add_paragraph()
        if style:
            try:
                p.style = doc.styles[style]
            except KeyError:
                pass
        if align is not None:
            p.alignment = align
        if kind == "image":
            p.add_run().add_picture(io.BytesIO(base64.b64decode(_PNG)), width=Pt(2))
        elif kind == "field":
            run = p.add_run()
            f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText"); instr.text = " PAGE "
            f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
            run._r.append(f1); run._r.append(instr); run._r.append(f2)
        elif kind == "direct-list":
            run = p.add_run(text)
            pPr = p._p.get_or_add_pPr()
            numPr = OxmlElement("w:numPr")
            numId = OxmlElement("w:numId"); numId.set(qn("w:val"), "1")
            numPr.append(numId)
            pPr.append(numPr)
        else:
            run = p.add_run(text)
            if run.font is not None:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _alignment_at(viols, index):
    return [v for v in viols if v.rule_code == "ALIGNMENT" and v.location.get("paragraph_index") == index]


def _alignment_codes(viols):
    return {v.rule_code for v in viols}


AM = WD_ALIGN_PARAGRAPH


def test_justified_body_is_valid():
    viols = run_static_rules_engine(_doc([("Body text.", None, AM.JUSTIFY, "text")]))
    assert _alignment_at(viols, 0) == []


def test_centered_body_remains_a_finding():
    viols = run_static_rules_engine(_doc([("Body text.", None, AM.CENTER, "text")]))
    assert len(_alignment_at(viols, 0)) == 1


def test_left_body_remains_a_finding():
    viols = run_static_rules_engine(_doc([("Body text.", None, AM.LEFT, "text")]))
    assert len(_alignment_at(viols, 0)) == 1


def test_empty_paragraph_skipped_even_with_explicit_center():
    viols = run_static_rules_engine(_doc([("", None, AM.CENTER, "text")]))
    assert _alignment_at(viols, 0) == []


def test_field_only_paragraph_skipped():
    viols = run_static_rules_engine(_doc([("", None, AM.CENTER, "field")]))
    assert _alignment_at(viols, 0) == []


def test_image_only_host_paragraph_skipped():
    viols = run_static_rules_engine(_doc([("", None, AM.CENTER, "image")]))
    assert _alignment_at(viols, 0) == []


def test_mixed_text_and_image_paragraph_skipped_as_figure_host():
    """Phase 2A: a drawing-host paragraph is FIGURE_HOST — mixed visible text
    never silently converts it to BODY, so it is NOT alignment-eligible."""
    # two paragraphs: image-only (skip) then mixed text+image (also skip —
    # the classifier marks any drawing host FIGURE_HOST)
    import base64
    from docx.shared import Pt as _Pt

    doc = Document()
    p = doc.add_paragraph()
    p.alignment = AM.CENTER
    r = p.add_run("See figure below.")
    r.font.name = "Times New Roman"
    p.add_run().add_picture(io.BytesIO(base64.b64decode(_PNG)), width=Pt(2))
    buf = io.BytesIO()
    doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    assert len([v for v in viols if v.rule_code == "ALIGNMENT"]) == 0


def test_semantic_caption_style_skipped():
    viols = run_static_rules_engine(_doc([("Table 1: Results", "Caption", AM.CENTER, "text")]))
    assert _alignment_at(viols, 0) == []


def test_manual_caption_text_skipped():
    viols = run_static_rules_engine(_doc([("Table 1: Results", None, AM.CENTER, "text")]))
    assert _alignment_at(viols, 0) == []


def test_title_and_subtitle_skipped():
    viols = run_static_rules_engine(_doc([
        ("Thesis Title", "Title", AM.CENTER, "text"),
        ("A Study", "Subtitle", AM.CENTER, "text"),
    ]))
    assert _alignment_at(viols, 0) == []
    assert _alignment_at(viols, 1) == []


def test_direct_numpr_list_skipped():
    viols = run_static_rules_engine(_doc([("Item", None, AM.LEFT, "direct-list")]))
    assert _alignment_at(viols, 0) == []


def test_style_inherited_list_skipped():
    viols = run_static_rules_engine(_doc([("Item", "List Bullet", AM.LEFT, "text")]))
    assert _alignment_at(viols, 0) == []


def test_heading_centered_remains_finding_under_preset():
    viols = run_static_rules_engine(_doc([
        ("Chapter One", "Heading 1", AM.CENTER, "text"),
        ("Section", "Heading 2", AM.CENTER, "text"),
    ]))
    assert len(_alignment_at(viols, 0)) == 1  # H1 center vs preset left
    assert len(_alignment_at(viols, 1)) == 1  # H2 center vs preset left


def test_heading_left_is_valid():
    viols = run_static_rules_engine(_doc([
        ("Chapter One", "Heading 1", AM.LEFT, "text"),
        ("Sub", "Heading 3", None, "text"),
    ]))
    assert _alignment_at(viols, 0) == []
    assert _alignment_at(viols, 1) == []


def test_references_entries_skipped_but_body_before_still_validated():
    viols = run_static_rules_engine(_doc([
        ("Body before references.", None, AM.LEFT, "text"),      # finding
        ("References", "Heading 1", AM.LEFT, "text"),             # header — skipped
        ("Garcia, A. (2018). A book. Publisher.", None, AM.LEFT, "text"),   # entry — skipped
        ("Lee, B. (2020). Another book. Publisher.", None, AM.LEFT, "text"),  # entry — skipped
    ]))
    assert len(_alignment_at(viols, 0)) == 1
    assert _alignment_at(viols, 1) == []
    assert _alignment_at(viols, 2) == []
    assert _alignment_at(viols, 3) == []


def test_references_entry_justified_is_valid_anyway():
    viols = run_static_rules_engine(_doc([
        ("References", None, AM.LEFT, "text"),
        ("Garcia, A. (2018). A book. Publisher.", None, AM.JUSTIFY, "text"),
    ]))
    assert _alignment_at(viols, 1) == []


def test_other_typography_findings_not_suppressed_for_skipped_paragraphs():
    """Phase 2A: a drawing-host paragraph (FIGURE_HOST) skips ALIGNMENT AND
    LINE_SPACING — the role is authoritative. Ordinary BODY paragraphs keep
    all typography findings."""
    import base64
    from docx.shared import Pt as _Pt

    doc = Document()
    p = doc.add_paragraph()
    p.alignment = AM.CENTER
    p.paragraph_format.line_spacing = 2.0  # wrong vs required 1.5
    p.add_run().add_picture(io.BytesIO(base64.b64decode(_PNG)), width=Pt(2))
    buf = io.BytesIO()
    doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    # FIGURE_HOST → both alignment and line-spacing skipped (role-gated).
    assert "ALIGNMENT" not in _alignment_codes(viols)
    assert "LINE_SPACING" not in _alignment_codes(viols)
    # A genuine BODY paragraph still gets its typography findings.
    doc2 = Document()
    b = doc2.add_paragraph("Ordinary body text that is long enough to be prose.")
    b.alignment = AM.CENTER
    b.paragraph_format.line_spacing = 2.0
    buf2 = io.BytesIO()
    doc2.save(buf2)
    viols2 = run_static_rules_engine(buf2.getvalue())
    assert "ALIGNMENT" in _alignment_codes(viols2)
    assert "LINE_SPACING" in _alignment_codes(viols2)
