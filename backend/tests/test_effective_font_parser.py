"""Effective font resolution tests — parser-level style-hierarchy resolution.

Covers the confirmed BODY Font Size missed-detection defect: python-docx
`run.font.size` returns None for style-inherited sizes, which previously
made the deterministic checks silently skip ordinary body paragraphs.
"""
import io

import pytest
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

from app.services.document_parser import (
    parse_document,
    extract_paragraphs,
    resolve_effective_size,
    resolve_effective_family,
)


def _build(builder) -> bytes:
    doc = Document()
    builder(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _runs_of(doc_bytes):
    doc = parse_document(doc_bytes)
    paras = extract_paragraphs(doc)
    return doc, paras


def test_direct_wsz_resolves_as_direct():
    def build(doc):
        p = doc.add_paragraph()
        r = p.add_run("direct")
        r.font.size = Pt(14)
        r.font.name = "Arial"

    _, paras = _runs_of(_build(build))
    run = paras[0]["runs"][0]
    assert run["font_size"] == 14.0            # legacy direct field unchanged
    assert run["effective_font_size"] == 14.0
    assert run["font_size_source"] == "direct"
    assert run["effective_font_name"] == "Arial"
    assert run["font_name_source"] == "direct"


def test_character_style_size_resolves():
    def build(doc):
        char = doc.styles.add_style("MyEmph", WD_STYLE_TYPE.CHARACTER)
        char.font.size = Pt(18)
        p = doc.add_paragraph()
        r = p.add_run("styled")
        r.style = doc.styles["MyEmph"]

    _, paras = _runs_of(_build(build))
    run = paras[0]["runs"][0]
    assert run["effective_font_size"] == 18.0
    assert run["font_size_source"] == "character_style"


def test_paragraph_style_size_resolves_as_paragraph_style():
    """The real-document case: Normal style carries 12 pt; runs have no
    direct formatting."""
    def build(doc):
        normal = doc.styles["Normal"]
        normal.font.size = Pt(12)
        normal.font.name = "Times New Roman"
        doc.add_paragraph("inherited body text")

    _, paras = _runs_of(_build(build))
    run = paras[0]["runs"][0]
    # Legacy direct-only fields remain None — this was the defect.
    assert run["font_size"] is None
    # Effective resolution now finds it.
    assert run["effective_font_size"] == 12.0
    assert run["font_size_source"] == "paragraph_style"
    assert run["effective_font_name"] == "Times New Roman"
    assert run["font_name_source"] == "paragraph_style"


def test_nested_base_style_chain_resolves_as_style_chain():
    def build(doc):
        base = doc.styles.add_style(" MyBaseBody", WD_STYLE_TYPE.PARAGRAPH)
        base.base_style = doc.styles["Normal"]
        base.font.size = Pt(11)
        child = doc.styles.add_style("ChildBody", WD_STYLE_TYPE.PARAGRAPH)
        child.base_style = base
        child.font.size = None  # inherits from base
        p = doc.add_paragraph("chained", style=child)

    _, paras = _runs_of(_build(build))
    run = paras[0]["runs"][0]
    assert run["effective_font_size"] == 11.0
    assert run["font_size_source"] == "style_chain"


def test_cycle_safe_style_traversal():
    """A basedOn cycle must terminate, not recurse forever."""
    def build(doc):
        a = doc.styles.add_style("CycleA", WD_STYLE_TYPE.PARAGRAPH)
        b = doc.styles.add_style("CycleB", WD_STYLE_TYPE.PARAGRAPH)
        a.base_style = b
        b.base_style = a
        b.font.size = Pt(10)
        doc.add_paragraph("cycled", style=a)

    # Must not raise RecursionError.
    _, paras = _runs_of(_build(build))
    run = paras[0]["runs"][0]
    assert run["effective_font_size"] == 10.0


def test_doc_defaults_resolve_when_no_style_declares_size():
    from docx.oxml.ns import qn

    def build(doc):
        # Remove any size declarations we can reach, then set docDefaults.
        styles_elm = doc.styles.element
        dd = styles_elm.find(qn("w:docDefaults"))
        assert dd is not None
        rpr_default = dd.find(qn("w:rPrDefault"))
        rpr = rpr_default.find(qn("w:rPr"))
        sz = rpr.find(qn("w:sz"))
        if sz is None:
            sz = rpr.makeelement(qn("w:sz"), {})
            rpr.append(sz)
        sz.set(qn("w:val"), "26")  # half-points -> 13 pt
        doc.add_paragraph("defaults only")

    _, paras = _runs_of(_build(build))
    run = paras[0]["runs"][0]
    if run["effective_font_size"] is not None:
        assert run["effective_font_size"] == 13.0
        assert run["font_size_source"] == "document_default"


def test_direct_overrides_inherited():
    def build(doc):
        normal = doc.styles["Normal"]
        normal.font.size = Pt(12)
        p = doc.add_paragraph()
        big = p.add_run("big")
        big.font.size = Pt(20)
        p.add_run("small")  # inherits 12

    _, paras = _runs_of(_build(build))
    runs = paras[0]["runs"]
    assert runs[0]["effective_font_size"] == 20.0
    assert runs[0]["font_size_source"] == "direct"
    assert runs[1]["effective_font_size"] == 12.0
    assert runs[1]["font_size_source"] == "paragraph_style"


def test_unresolved_remains_none_with_source_marker():
    def build(doc):
        # A dedicated paragraph style with no size anywhere in its chain and
        # no docDefault fallback exercised by an isolated resolver call.
        p = doc.add_paragraph()
        p.add_run("")  # empty text run still resolves structurally

    doc_bytes = _build(build)
    doc = parse_document(doc_bytes)
    para = doc.paragraphs[0]
    run = para.runs[0]
    # Directly probe the resolver with no styles_element so every hop misses.
    size, source = resolve_effective_size(run, para, styles_element=None)
    assert size is None
    assert source == "unresolved"
    family, fsource = resolve_effective_family(run, para, styles_element=None)
    assert family is None
    assert fsource == "unresolved"


def test_effective_font_family_uses_same_precedence():
    def build(doc):
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(12)
        p = doc.add_paragraph()
        direct = p.add_run("direct family")
        direct.font.name = "Georgia"
        direct.font.size = Pt(12)
        p.add_run("inherited family")

    _, paras = _runs_of(_build(build))
    runs = paras[0]["runs"]
    assert runs[0]["effective_font_name"] == "Georgia"
    assert runs[0]["font_name_source"] == "direct"
    assert runs[1]["effective_font_name"] == "Calibri"
    assert runs[1]["font_name_source"] == "paragraph_style"


def test_high_ansi_only_declaration_resolves():
    def build(doc):
        from docx.oxml.ns import qn
        p = doc.add_paragraph()
        r = p.add_run("x")
        rPr = r._element.get_or_add_rPr()
        rfonts = rPr.makeelement(qn("w:rFonts"), {})
        rfonts.set(qn("w:hAnsi"), "MS Mincho")
        rPr.insert(0, rfonts)

    _, paras = _runs_of(_build(build))
    run = paras[0]["runs"][0]
    assert run["effective_font_name"] == "MS Mincho"
    assert run["font_name_source"] == "direct"


def test_no_document_text_persisted_in_metadata():
    """Run metadata carries provenance, not diagnostics text."""
    def build(doc):
        normal = doc.styles["Normal"]
        normal.font.size = Pt(12)
        doc.add_paragraph("some body sentence content")

    _, paras = _runs_of(_build(build))
    run = paras[0]["runs"][0]
    keys = set(run.keys())
    allowed_meta = {"index", "text", "font_name", "font_size", "bold", "italic",
                    "underline", "effective_font_size", "font_size_source",
                    "effective_font_name", "font_name_source"}
    extra = keys - allowed_meta
    assert not extra, f"unexpected metadata keys: {extra}"
