"""Shared pytest fixtures for the FastAPI backend.

Uses an in-memory SQLite engine, a no-op background AI task, and a
`make_docx_bytes()` factory for building valid .docx payloads in tests.
"""
import base64
import io
import os

# Isolate tests from the real development database BEFORE any app module is
# imported and creates a SQLAlchemy engine. pydantic-settings gives the env
# var priority over backend/.env, so app.database.engine becomes an in-memory
# engine instead of targeting backend/audit.db.
os.environ["DATABASE_URL"] = "sqlite:///:memory:"

import pytest
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.database import Base, get_db
from app.config import settings, PresetConfig
from app.main import app
from app.api import routes as api_routes


# ---------------------------------------------------------------------------
# DB fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def test_engine():
    """Fresh in-memory SQLite engine per test."""
    engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    try:
        yield engine
    finally:
        Base.metadata.drop_all(bind=engine)
        engine.dispose()


# ---------------------------------------------------------------------------
# Settings / app fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def small_file_cap(monkeypatch):
    """Shrink MAX_FILE_SIZE so oversize tests don't need multi-MB payloads.
    Use via `with small_file_cap: client = client_with_cap()` or just apply
    the monkeypatch and rebuild the client with `make_client_with_cap`.
    """
    monkeypatch.setattr(settings, "MAX_FILE_SIZE", 1024)  # 1 KB
    return settings


def _build_client(test_engine):
    """Construct a TestClient with the given engine, no settings mutation."""
    Session = sessionmaker(autocommit=False, autoflush=False, bind=test_engine)

    def _override_get_db():
        db = Session()
        try:
            yield db
        finally:
            db.close()

    app.dependency_overrides[get_db] = _override_get_db
    return TestClient(app)


@pytest.fixture
def client(test_engine, mock_ai_task, mock_init_db):
    """TestClient with get_db overridden to the in-memory engine.
    Uses real MAX_FILE_SIZE — do NOT use for oversize tests.
    """
    c = _build_client(test_engine)
    try:
        with c:
            yield c
    finally:
        app.dependency_overrides.clear()


@pytest.fixture
def client_file_db(tmp_path, mock_ai_task, mock_init_db):
    """TestClient backed by a temporary FILE-backed SQLite database.

    Unlike the in-memory `client` (which shares ONE StaticPool connection
    across threads), a file-backed SQLite uses SQLAlchemy's default
    QueuePool: each concurrent request thread gets its own pooled
    connection and session, so `db.refresh()` after commit cannot race.

    Purpose: concurrency tests only — the production app keeps its own
    single-connection behavior untouched.
    """
    db_path = tmp_path / "concurrent-tests.db"
    engine = create_engine(
        f"sqlite:///{db_path}",
        connect_args={"check_same_thread": False},
    )
    Base.metadata.create_all(bind=engine)
    Session = sessionmaker(autocommit=False, autoflush=False, bind=engine)

    def _override_get_db():
        db = Session()
        try:
            yield db
        finally:
            db.close()

    app.dependency_overrides[get_db] = _override_get_db
    c = TestClient(app)
    try:
        with c:
            yield c
    finally:
        app.dependency_overrides.clear()
        Session.close_all()
        engine.dispose()
        # tmp_path removal is handled by pytest; nothing is written to
        # backend/audit.db (mock_init_db prevents startup writes).


@pytest.fixture
def client_with_small_cap(test_engine, small_file_cap, mock_ai_task, mock_init_db):
    """TestClient variant with a tiny MAX_FILE_SIZE for oversize tests."""
    c = _build_client(test_engine)
    try:
        with c:
            yield c
    finally:
        app.dependency_overrides.clear()


@pytest.fixture
def mock_ai_task(monkeypatch):
    """No-op the AI citation task so tests never hit the network.

    Returns a completed, empty local result — the default deterministic
    outcome for tests that do not care about AI specifics.
    """
    from app.services.ai_citation import AiCitationResult

    async def _no_op(*args, **kwargs):
        return AiCitationResult(
            status="COMPLETED_NO_SUGGESTIONS",
            provider="LOCAL_OLLAMA",
            suggestions=[],
        )
    monkeypatch.setattr(api_routes, "async_ai_citation_task", _no_op)
    return _no_op


@pytest.fixture
def mock_init_db(monkeypatch):
    """Prevent the startup event from writing to ./audit.db.

    Also no-ops startup stale-audit reconciliation so the app's own
    (separate) in-memory engine is never queried without its tables.
    """
    monkeypatch.setattr("app.main.init_db", lambda: None)
    monkeypatch.setattr(
        "app.main.reconcile_stale_audits",
        lambda *args, **kwargs: 0,
    )


# (client fixtures defined above after small_file_cap)


# ---------------------------------------------------------------------------
# docx factory
# ---------------------------------------------------------------------------

# 1x1 transparent PNG — smallest embeddable picture for figure tests.
_PNG_1PX = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8"
    "z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)


def _add_word_caption(doc, label, text):
    """Insert a real Word caption: Caption style + SEQ field.

    The field instruction is split across two <w:instrText> runs (as Word
    does) so caption detection must reassemble split field runs. The cached
    field result ('1') is written as visible text like Word does.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    cap = doc.add_paragraph()
    try:
        cap.style = doc.styles["Caption"]
    except KeyError:
        pass
    label_run = cap.add_run(f"{label} ")
    begin = cap.add_run()
    f = OxmlElement("w:fldChar")
    f.set(qn("w:fldCharType"), "begin")
    begin._r.append(f)
    instr1 = cap.add_run()
    t = OxmlElement("w:instrText")
    t.set(qn("xml:space"), "preserve")
    t.text = " SEQ "
    instr1._r.append(t)
    instr2 = cap.add_run()
    t = OxmlElement("w:instrText")
    t.set(qn("xml:space"), "preserve")
    t.text = f"{label} \\* ARABIC "
    instr2._r.append(t)
    sep = cap.add_run()
    f = OxmlElement("w:fldChar")
    f.set(qn("w:fldCharType"), "separate")
    sep._r.append(f)
    result_run = cap.add_run("1")
    end = cap.add_run()
    f = OxmlElement("w:fldChar")
    f.set(qn("w:fldCharType"), "end")
    end._r.append(f)
    cap.add_run(text)
    for run in (label_run, result_run):
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return cap


def _set_alt_text(para, descr):
    """Set docPr@descr on the first drawing in a paragraph."""
    ns = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
    for docpr in para._p.findall(".//" + ns + "docPr"):
        docpr.set("descr", descr)
        return


def make_docx_bytes(
    paragraphs=None,
    references=None,
    margins=None,
    font_name="Times New Roman",
    body_size=12,
    line_spacing=1.5,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    tables=None,
    with_caption=True,
    images=None,
    with_image_caption=True,
    with_alt_text=True,
):
    """Build a docx in-memory, return bytes.

    `paragraphs`: list[str]  or  list[tuple[str, style_name]]
    `references`: list[str] — bibliography entries (preceded by "References" header)
    `margins`: dict with keys left/right/top/bottom in inches
    `tables`: list[list[list[str]]] — table contents; auto-preceded by a real
              Word caption (Caption style + SEQ Table field) when with_caption
    `with_caption`: if False, skip the auto-caption (for caption-missing tests)
    `images`: int — number of 1x1 PNG inline pictures to embed
    `with_image_caption`: real Word Figure caption below each image
    `with_alt_text`: docPr@descr alt-text on each image
    """
    doc = Document()
    if margins:
        sec = doc.sections[0]
        if "left" in margins:
            sec.left_margin = Inches(margins["left"])
        if "right" in margins:
            sec.right_margin = Inches(margins["right"])
        if "top" in margins:
            sec.top_margin = Inches(margins["top"])
        if "bottom" in margins:
            sec.bottom_margin = Inches(margins["bottom"])

    if paragraphs:
        for p in paragraphs:
            text = p[0] if isinstance(p, tuple) else p
            style = p[1] if isinstance(p, tuple) and len(p) > 1 else None
            para = doc.add_paragraph(text)
            if style:
                try:
                    para.style = doc.styles[style]
                except KeyError:
                    pass
            for run in para.runs:
                run.font.name = font_name
                run.font.size = Pt(body_size)
            pf = para.paragraph_format
            pf.line_spacing = line_spacing
            pf.alignment = alignment

    if tables:
        for tbl_idx, tbl in enumerate(tables):
            if with_caption:
                _add_word_caption(doc, "Table", f": Test table {tbl_idx + 1}")
            t = doc.add_table(rows=len(tbl), cols=max(len(r) for r in tbl))
            for r_i, row in enumerate(tbl):
                for c_i, val in enumerate(row):
                    t.rows[r_i].cells[c_i].text = val

    if images:
        for img_idx in range(images):
            pic_para = doc.add_paragraph()
            pic_para.add_run().add_picture(io.BytesIO(_PNG_1PX), width=Inches(1))
            if with_alt_text:
                _set_alt_text(pic_para, "Test image alt text")
            if with_image_caption:
                _add_word_caption(doc, "Figure", f": Test figure {img_idx + 1}")

    if references:
        # The References header classifies as REFERENCES_HEADING (heading-
        # like), so it carries an H1 font requirement once effective-font
        # resolution sees through style inheritance. Give both the header
        # and the entries explicit conforming formatting so "clean
        # document" fixtures stay clean under correct detection.
        refs_header = doc.add_paragraph("References")
        for run in refs_header.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(16)  # matches the default preset's H1 size
        for ref in references:
            ref_para = doc.add_paragraph(ref)
            for run in ref_para.runs:
                run.font.name = font_name
                run.font.size = Pt(body_size)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def docx_factory():
    return make_docx_bytes
