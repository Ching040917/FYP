"""Document Role Integration — Phase 2A: Font and Paragraph Typography
eligibility driven by the authoritative paragraph role.

Contract under test:
  - BODY (and APPENDIX_BODY prose) keep body font/size/alignment/
    line-spacing/space-before/space-after findings.
  - HEADING_1/2/3 keep configured heading font/size/alignment/spacing.
  - COVER / TITLE / SUBTITLE produce no BODY typography findings.
  - TOC heading/entry produce no BODY typography findings.
  - LIST_ITEM keeps visible-text font checks; list spacing policy is
    preserved (LIST_SPACE_AFTER=None ⇒ no list-spacing finding).
  - CAPTION_TABLE/CAPTION_FIGURE use only Caption-specific spacing.
  - REFERENCE_ENTRY uses REFERENCES_LINE_SPACING (2.0 default); 1.5/1.0
    produce the role-specific finding; body requirements never apply.
  - DISPLAYED_EQUATION / FIGURE_HOST / EMPTY / FIELD_ONLY are skipped.
  - UNKNOWN produces no deduction and is never converted to BODY.
  - role == null (legacy) retains the old behavior.
  - Paragraph indexes and evidence navigation remain stable.
"""
import base64
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from app.config import PresetConfig
from app.services.layout_engine import run_static_rules_engine

_PNG = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAA"
    "BJRU5ErkJggg=="
)

AM = WD_ALIGN_PARAGRAPH


def _base_doc():
    doc = Document()
    for name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3",
                 "Caption", "TOC Heading", "List Bullet", "List Number"):
        try:
            s = doc.styles[name]
            s.font.name = "Times New Roman"
            s.font.size = Pt(12)
        except KeyError:
            pass
    return doc


def _add_numpr(p, ilvl=0):
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    lvl = OxmlElement("w:ilvl"); lvl.set(qn("w:val"), str(ilvl))
    nid = OxmlElement("w:numId"); nid.set(qn("w:val"), "1")
    numPr.append(lvl); numPr.append(nid)
    pPr.append(numPr)


def _add_toc_field(doc, text):
    p = doc.add_paragraph()
    r = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    r._r.append(f1); r._r.append(instr); r._r.append(f2)
    p.add_run(text)
    r2 = p.add_run()
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    r2._r.append(f3)
    return p


def _add_equation(doc, text):
    p = doc.add_paragraph()
    m = OxmlElement("m:oMath")
    r = OxmlElement("m:r")
    t = OxmlElement("m:t"); t.text = text
    r.append(t); m.append(r)
    p._p.append(m)
    return p


def _add_image_para(doc, text=None):
    p = doc.add_paragraph()
    if text:
        p.add_run(text)
    p.add_run().add_picture(io.BytesIO(base64.b64decode(_PNG)), width=Pt(12))
    return p


def _codes(viols):
    return {v.rule_code for v in viols}


def _at(viols, code, index):
    return [v for v in viols if v.rule_code == code and v.location.get("paragraph_index") == index]


def _body_text(doc, text="Ordinary body text that is long enough to be prose."):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.alignment = AM.JUSTIFY
    return p


# ---------------------------------------------------------------------------
# BODY / HEADING retention
# ---------------------------------------------------------------------------

def test_body_retains_valid_typography_findings():
    doc = _base_doc()
    p = doc.add_paragraph("Body text with wrong typography here.")
    p.alignment = AM.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(3)
    for r in p.runs:
        r.font.name = "Arial"
    buf = io.BytesIO(); doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    assert "FONT_CONSISTENCY" in _codes(viols)
    assert "ALIGNMENT" in _codes(viols)
    assert "LINE_SPACING" in _codes(viols)
    assert "SPACE_BEFORE" in _codes(viols)


def test_headings_retain_configured_findings():
    doc = _base_doc()
    for i, (text, style, size) in enumerate([
        ("Chapter One", "Heading 1", 16),
        ("Section", "Heading 2", 14),
        ("Subsection", "Heading 3", 12),
    ]):
        p = doc.add_paragraph(text, style=style)
        p.alignment = AM.CENTER  # wrong vs preset left
        for r in p.runs:
            r.font.size = Pt(size + 4)  # wrong size
    buf = io.BytesIO(); doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    assert len(_at(viols, "ALIGNMENT", 0)) == 1
    assert len(_at(viols, "ALIGNMENT", 1)) == 1
    assert len(_at(viols, "ALIGNMENT", 2)) == 1
    assert len(_at(viols, "FONT_SIZE", 0)) == 1


# ---------------------------------------------------------------------------
# Cover / Title / Subtitle / TOC skip
# ---------------------------------------------------------------------------

def test_cover_labels_produce_no_body_findings():
    """Cover-region paragraphs (with Arial, center, 1.0 spacing) produce no
    FONT/typography findings — the role is COVER."""
    doc = _base_doc()
    p = doc.add_paragraph("Assignment Title", style="Title")
    for r in p.runs: r.font.name = "Arial"
    p2 = doc.add_paragraph("Student Name Here")
    p2.alignment = AM.CENTER
    p2.paragraph_format.line_spacing = 1.0
    p3 = doc.add_paragraph("Detail:")
    p3.paragraph_format.line_spacing = 1.0
    doc.add_paragraph("1. Introduction", style="Heading 1")
    _body_text(doc)
    buf = io.BytesIO(); doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    # Title (0) and cover labels (1,2) produce no BODY typography findings.
    for idx in (0, 1, 2):
        assert _at(viols, "FONT_CONSISTENCY", idx) == []
        assert _at(viols, "LINE_SPACING", idx) == []
        assert _at(viols, "ALIGNMENT", idx) == []
    # Heading + body still validate.
    assert len(_at(viols, "ALIGNMENT", 4)) == 0  # body justified → clean
    assert _at(viols, "LINE_SPACING", 4) == []


def test_title_and_subtitle_produce_no_body_findings():
    doc = _base_doc()
    t = doc.add_paragraph("Thesis Title", style="Title")
    t.alignment = AM.CENTER
    t.paragraph_format.line_spacing = 1.0
    for r in t.runs: r.font.name = "Arial"
    s = doc.add_paragraph("A Study", style="Subtitle")
    s.alignment = AM.CENTER
    s.paragraph_format.line_spacing = 1.0
    for r in s.runs: r.font.name = "Arial"
    doc.add_paragraph("1. Introduction", style="Heading 1")
    _body_text(doc)
    buf = io.BytesIO(); doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    for idx in (0, 1):
        assert _at(viols, "FONT_CONSISTENCY", idx) == []
        assert _at(viols, "LINE_SPACING", idx) == []
        assert _at(viols, "ALIGNMENT", idx) == []


def test_toc_entries_produce_no_body_findings():
    doc = _base_doc()
    doc.add_paragraph("Table of Contents", style="Heading 1")
    p = _add_toc_field(doc, "1. Introduction")
    p.alignment = AM.CENTER
    p.paragraph_format.line_spacing = 1.0
    for r in p.runs: r.font.name = "Arial"
    doc.add_paragraph("1. Introduction", style="Heading 1")
    _body_text(doc)
    buf = io.BytesIO(); doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    # TOC entry (index 1) — no body typography findings.
    assert _at(viols, "FONT_CONSISTENCY", 1) == []
    assert _at(viols, "LINE_SPACING", 1) == []
    assert _at(viols, "ALIGNMENT", 1) == []


# ---------------------------------------------------------------------------
# Lists
# ---------------------------------------------------------------------------

def test_list_keeps_font_checks_but_no_spacing_findings():
    """LIST_ITEM: visible-text font checks apply; with LIST_SPACE_AFTER=None
    (default) no list SPACE_BEFORE/SPACE_AFTER finding exists. Body
    line-spacing continues to apply to lists (legacy policy preserved)."""
    doc = _base_doc()
    p = doc.add_paragraph("First list item")
    _add_numpr(p)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs: r.font.name = "Arial"
    buf = io.BytesIO(); doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    assert len(_at(viols, "FONT_CONSISTENCY", 0)) == 1  # visible-text font check
    assert _at(viols, "SPACE_AFTER", 0) == []            # silent policy
    assert _at(viols, "SPACE_BEFORE", 0) == []
    assert _at(viols, "ALIGNMENT", 0) == []              # list alignment exempt
    assert _at(viols, "LINE_SPACING", 0) == []           # 1.5 is correct anyway


def test_configured_list_space_after_still_validates():
    class _ListPreset(PresetConfig):
        LIST_SPACE_AFTER = 6

    doc = _base_doc()
    p = doc.add_paragraph("List item")
    _add_numpr(p)
    p.paragraph_format.space_after = Pt(3)
    buf = io.BytesIO(); doc.save(buf)

    import app.services.layout_engine as le
    from app.config import settings
    original = settings.PRESET
    settings.PRESET = _ListPreset()
    try:
        viols = run_static_rules_engine(buf.getvalue())
    finally:
        settings.PRESET = original
    assert len(_at(viols, "SPACE_AFTER", 0)) == 1


# ---------------------------------------------------------------------------
# Captions
# ---------------------------------------------------------------------------

def _add_word_caption(doc, label, text, style=True, space_before=3, space_after=3):
    cap = doc.add_paragraph()
    if style:
        cap.style = doc.styles["Caption"]
    cap.paragraph_format.space_before = Pt(space_before)
    cap.paragraph_format.space_after = Pt(space_after)
    cap.paragraph_format.line_spacing = 1.0  # wrong for body — must be ignored
    cap.alignment = AM.CENTER                # wrong for body — must be ignored
    cap.add_run(f"{label} ")
    for piece in (" SEQ ", f"{label} ", "\\* ARABIC "):
        run = cap.add_run()
        t = OxmlElement("w:instrText"); t.set(qn("xml:space"), "preserve"); t.text = piece
        run._r.append(t)
    cap.add_run("1")
    cap.add_run(text)
    return cap


def test_semantic_caption_uses_only_caption_spacing():
    """CAPTION role: no body font/alignment/line-spacing; caption spacing
    only when configured (default silent → no findings)."""
    doc = _base_doc()
    _body_text(doc)
    _add_word_caption(doc, "Figure", ": Chart")
    _body_text(doc, "More ordinary body text after the caption here.")
    buf = io.BytesIO(); doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    # caption at index 1
    assert _at(viols, "ALIGNMENT", 1) == []
    assert _at(viols, "LINE_SPACING", 1) == []
    assert _at(viols, "SPACE_BEFORE", 1) == []
    assert _at(viols, "SPACE_AFTER", 1) == []
    assert _at(viols, "FONT_CONSISTENCY", 1) == []


def test_manual_caption_uses_only_caption_spacing():
    doc = _base_doc()
    _body_text(doc)
    p = doc.add_paragraph("Table 1: Results")
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = AM.CENTER
    p.paragraph_format.line_spacing = 1.0
    buf = io.BytesIO(); doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    assert _at(viols, "ALIGNMENT", 1) == []
    assert _at(viols, "LINE_SPACING", 1) == []
    assert _at(viols, "SPACE_BEFORE", 1) == []
    assert _at(viols, "SPACE_AFTER", 1) == []


def test_configured_caption_spacing_validates():
    class _CapPreset(PresetConfig):
        CAPTION_SPACE_BEFORE = 6
        CAPTION_SPACE_AFTER = 12

    doc = _base_doc()
    _body_text(doc)
    _add_word_caption(doc, "Figure", ": Chart")  # 3/3 vs 6/12
    buf = io.BytesIO(); doc.save(buf)

    import app.services.layout_engine as le
    from app.config import settings
    original = settings.PRESET
    settings.PRESET = _CapPreset()
    try:
        viols = run_static_rules_engine(buf.getvalue())
    finally:
        settings.PRESET = original
    assert len(_at(viols, "SPACE_BEFORE", 1)) == 1
    assert len(_at(viols, "SPACE_AFTER", 1)) == 1


# ---------------------------------------------------------------------------
# References
# ---------------------------------------------------------------------------

def _refs_doc(line_spacing):
    doc = _base_doc()
    _body_text(doc)
    doc.add_paragraph("References", style="Heading 1")
    for ref in ["Smith, J. (2020). Title. Press.", "Garcia, A. (2018). Book. Publisher."]:
        p = doc.add_paragraph(ref)
        p.paragraph_format.line_spacing = line_spacing
    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


def test_references_2_0_accepted():
    """REFERENCES_LINE_SPACING=2.0 (approved profile) → no finding."""
    viols = run_static_rules_engine(_refs_doc(2.0))
    assert "LINE_SPACING" not in _codes(viols)


def test_references_1_5_produces_role_specific_finding():
    viols = run_static_rules_engine(_refs_doc(1.5))
    ls = [v for v in viols if v.rule_code == "LINE_SPACING"]
    assert len(ls) == 2  # both entries flagged
    assert all(v.expected_value == "2.0" for v in ls)
    assert all(v.message.startswith("References line spacing") for v in ls)


def test_references_1_0_produces_role_specific_finding():
    viols = run_static_rules_engine(_refs_doc(1.0))
    ls = [v for v in viols if v.rule_code == "LINE_SPACING"]
    assert len(ls) == 2
    assert all(v.expected_value == "2.0" for v in ls)


def test_references_no_body_alignment_or_spacing_findings():
    """Reference entries never get BODY alignment/paragraph-spacing."""
    doc = _base_doc()
    _body_text(doc)
    doc.add_paragraph("References", style="Heading 1")
    p = doc.add_paragraph("Smith, J. (2020). Title. Press.")
    p.alignment = AM.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 2.0
    buf = io.BytesIO(); doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    assert _at(viols, "ALIGNMENT", 2) == []
    assert _at(viols, "SPACE_BEFORE", 2) == []
    assert _at(viols, "SPACE_AFTER", 2) == []
    assert _at(viols, "LINE_SPACING", 2) == []


# ---------------------------------------------------------------------------
# Equations / empty / field-only / UNKNOWN / legacy
# ---------------------------------------------------------------------------

def test_equations_skipped():
    doc = _base_doc()
    _body_text(doc)
    p = _add_equation(doc, "score = w * m")
    p.alignment = AM.CENTER
    p.paragraph_format.line_spacing = 1.0
    buf = io.BytesIO(); doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    assert _at(viols, "ALIGNMENT", 1) == []
    assert _at(viols, "LINE_SPACING", 1) == []
    assert _at(viols, "FONT_CONSISTENCY", 1) == []


def test_empty_field_only_image_only_skipped():
    doc = _base_doc()
    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.text = " PAGE "
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    r._r.append(f1); r._r.append(instr); r._r.append(f2)
    _add_image_para(doc)
    buf = io.BytesIO(); doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    for idx in (0, 1, 2):
        assert _at(viols, "ALIGNMENT", idx) == []
        assert _at(viols, "LINE_SPACING", idx) == []
        assert _at(viols, "FONT_CONSISTENCY", idx) == []


def test_unknown_produces_no_deduction():
    """Heading 4+ classifies UNKNOWN — no typography deduction, and it is
    never converted to BODY."""
    doc = _base_doc()
    p = doc.add_paragraph("Level Four", style="Heading 4") if "Heading 4" in [s.name for s in doc.styles] else doc.add_paragraph("Level Four")
    if "Heading 4" in [s.name for s in doc.styles]:
        p.style = doc.styles["Heading 4"]
    p.alignment = AM.CENTER
    p.paragraph_format.line_spacing = 1.0
    buf = io.BytesIO(); doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue())
    assert _at(viols, "ALIGNMENT", 0) == []
    assert _at(viols, "LINE_SPACING", 0) == []


def test_legacy_role_null_retains_behavior():
    """Direct rule invocation without roles (legacy) keeps old heuristics."""
    from app.services.layout_engine import (
        check_font_consistency,
        check_font_size,
        check_paragraph_typography,
    )
    from app.services.document_parser import extract_paragraphs
    from app.config import settings

    doc = _base_doc()
    p = doc.add_paragraph("Body text here.")
    p.alignment = AM.CENTER
    p.paragraph_format.line_spacing = 1.0
    for r in p.runs: r.font.name = "Arial"
    buf = io.BytesIO(); doc.save(buf)

    from app.services.document_parser import parse_document
    d = parse_document(buf.getvalue())
    paras = extract_paragraphs(d)
    preset = settings.PRESET
    # No roles argument → legacy path: font/typography findings fire.
    assert len(check_font_consistency(paras, preset)) == 1
    assert len(check_paragraph_typography(paras, preset, doc=d)) >= 2  # align + LS


# ---------------------------------------------------------------------------
# Sanitized real-document regression (Phase 2A)
# ---------------------------------------------------------------------------

def _sanitized_real_doc():
    """The sanitized real regression document: cover title area (Arial,
    centered, 1.0 spacing), admin tables, assessment labels, academic body,
    and References at 2.0 line spacing. The logo lives in the header."""
    doc = _base_doc()
    hdr = doc.sections[0].header
    hdr.paragraphs[0].add_run().add_picture(
        io.BytesIO(base64.b64decode(_PNG)), width=Pt(12)
    )
    # Cover title area — Normal style, Arial, center, 1.0 spacing.
    for title_line in ("THE IMPACT OF ARTIFICIAL INTELLIGENCE", "ON HIGHER EDUCATION"):
        p = doc.add_paragraph(title_line)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = AM.CENTER
        for r in p.runs:
            r.font.name = "Arial"
    # Administrative info tables.
    t0 = doc.add_table(rows=3, cols=2)
    for r, (k, v) in enumerate([("Name", "Sanitized Student"), ("Matric No", "A12345"),
                                ("Programme", "Bachelor of Computing")]):
        t0.rows[r].cells[0].text = k
        t0.rows[r].cells[1].text = v
    p = doc.add_paragraph("Assignment Details")
    p.paragraph_format.line_spacing = 1.0
    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "Module"
    t1.rows[0].cells[1].text = "Research Methods"
    t1.rows[1].cells[0].text = "Submission Date"
    t1.rows[1].cells[1].text = "15 January 2025"
    # Assessment / mark area.
    p = doc.add_paragraph("Assessment")
    p.paragraph_format.line_spacing = 1.0
    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "Mark"
    t2.rows[0].cells[1].text = "Comments"
    t2.rows[1].cells[0].text = "78"
    t2.rows[1].cells[1].text = "Well structured"
    # Academic content.
    doc.add_paragraph("1. Introduction", style="Heading 1")
    p = doc.add_paragraph("Artificial intelligence is transforming higher education in many observable ways.")
    p.paragraph_format.line_spacing = 1.5
    doc.add_paragraph("2. Literature Review", style="Heading 1")
    p = doc.add_paragraph("Prior work has examined the role of machine learning in universities.")
    p.paragraph_format.line_spacing = 1.5
    # References — 2.0 line spacing (approved profile).
    doc.add_paragraph("References", style="Heading 1")
    for ref in ["Smith, J. (2020). Artificial intelligence in education. Press.",
                "Garcia, A. (2018). Machine learning foundations. Publisher."]:
        p = doc.add_paragraph(ref)
        p.paragraph_format.line_spacing = 2.0
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _role_at(viols, index):
    from app.services.document_parser import parse_document, extract_paragraphs
    from app.services.role_classifier import classify_paragraphs
    doc = parse_document(_sanitized_real_doc())
    roles = classify_paragraphs(doc, extract_paragraphs(doc))
    return roles[index]


def test_sanitized_real_doc_cover_findings_removed():
    """Cover Font/Line-Spacing/Alignment findings are removed; BODY and
    Heading findings are retained; References 2.0 findings removed."""
    viols = run_static_rules_engine(_sanitized_real_doc())
    # The cover region (0..3) has NO font/typography findings.
    from app.services.document_parser import parse_document, extract_paragraphs
    from app.services.role_classifier import classify_paragraphs
    doc = parse_document(_sanitized_real_doc())
    paras = extract_paragraphs(doc)
    roles = classify_paragraphs(doc, paras)
    for idx, role in enumerate(roles):
        if role in ("COVER", "TITLE", "SUBTITLE"):
            assert _at(viols, "FONT_CONSISTENCY", idx) == [], f"cover {idx}"
            assert _at(viols, "LINE_SPACING", idx) == [], f"cover {idx}"
            assert _at(viols, "ALIGNMENT", idx) == [], f"cover {idx}"
    # References entries: 2.0 → no LINE_SPACING finding.
    ref_idxs = [i for i, r in enumerate(roles) if r == "REFERENCE_ENTRY"]
    assert ref_idxs, "expected reference entries"
    for idx in ref_idxs:
        assert _at(viols, "LINE_SPACING", idx) == [], f"ref {idx}"
    # Valid BODY/Heading findings retained — the academic body is clean here,
    # so assert the engine still reports them when violated (checked in the
    # dedicated tests); here assert no BODY/HEADING paragraph lost findings
    # that should exist: body paragraphs are all compliant in this fixture.
    body_idxs = [i for i, r in enumerate(roles) if r == "BODY"]
    assert body_idxs
    for idx in body_idxs:
        assert _at(viols, "LINE_SPACING", idx) == []
        assert _at(viols, "ALIGNMENT", idx) == []
    # The role-identity contract: every role is from the known enum.
    from app.services.role_classifier import PARAGRAPH_ROLES
    assert all(r in PARAGRAPH_ROLES for r in roles)


def test_sanitized_real_doc_score_improves():
    """Score trajectory across builds: 68 (pre-2A) → 78 (Phase 2A cover +
    references typography removed) → 84 (Phase 2B: 3 administrative/rubric
    Table Caption false positives removed, each MINOR=2pt) → 72 (effective-
    font resolution: three style-inherited H1 runs now correctly resolve to
    12 pt vs the required 16 pt — genuine MAJOR FONT_SIZE findings that were
    previously invisible when run.font.size returned None)."""
    from app.services.scoring import calculate_weighted_score_detailed
    viols = run_static_rules_engine(_sanitized_real_doc())
    score = calculate_weighted_score_detailed(viols).total
    assert score == 72, f"expected 72, got {score}"
    # The three newly detected violations are heading-size findings on
    # distinct paragraphs — traceable, not aggregated.
    font_findings = [v for v in viols if v.rule_code == "FONT_SIZE"]
    assert len(font_findings) == 3
    assert {v.location["paragraph_index"] for v in font_findings} == {4, 6, 8}
