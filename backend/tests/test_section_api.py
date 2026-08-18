"""API compatibility tests for the persisted section metadata.

Verifies:
- POST /api/audit persists `sections` (list of SectionMetadata);
- GET /api/audit/{id} returns the IDENTICAL metadata (POST + GET consistent,
  so refresh / History / AuditPage refetch never lose it);
- the structure is nullable and backward-compatible (never required);
- historical rows without metadata return `sections: null`;
- existing response fields are untouched (no breaking change);
- the migration upgrade / downgrade / re-upgrade round-trips safely.
"""
import io
import uuid
from pathlib import Path
from sqlalchemy import create_engine, inspect
from docx import Document
from docx.shared import Inches


def test_submit_response_includes_section_metadata(client, docx_factory):
    doc = Document()
    doc.sections[0].left_margin = Inches(1.5)
    doc.add_paragraph("alpha body paragraph")
    doc.add_section()
    doc.add_paragraph("beta body paragraph")
    buf = io.BytesIO()
    doc.save(buf)

    resp = client.post(
        "/api/audit",
        files={"file": ("sections.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200
    data = resp.json()
    assert "sections" in data
    assert isinstance(data["sections"], list)
    assert len(data["sections"]) == 2
    s0 = data["sections"][0]
    assert s0["section_index"] == 0
    assert s0["start_paragraph_index"] == 0
    assert s0["break_type"] in ("nextPage", "continuous", "oddPage", "evenPage")
    # no text/paths leak
    blob = str(data["sections"])
    assert "alpha" not in blob and "beta" not in blob
    assert ".docx" not in blob and "C:" not in blob


def test_submit_response_backward_compatible_optional(client, docx_factory):
    resp = client.post(
        "/api/audit",
        files={"file": ("plain.docx", docx_factory(["just a paragraph"]), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200
    data = resp.json()
    # sections present but harmless for a single-section doc
    assert "sections" in data
    assert len(data["sections"]) == 1
    assert data["sections"][0]["end_paragraph_index"] is None


def test_get_returns_identical_metadata_to_post(client, docx_factory):
    """POST → GET must return the SAME sections (the real AuditPage refetch
    replaces the submit result with a GET result — metadata must survive)."""
    doc = Document()
    doc.add_paragraph("one")
    doc.add_section()
    doc.add_paragraph("two")
    doc.add_section()
    doc.add_paragraph("three")
    buf = io.BytesIO()
    doc.save(buf)

    post_resp = client.post(
        "/api/audit",
        files={"file": ("multi.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert post_resp.status_code == 200
    audit_id = post_resp.json()["audit_id"]
    post_sections = post_resp.json()["sections"]

    get_resp = client.get(f"/api/audit/{audit_id}")
    assert get_resp.status_code == 200
    get_sections = get_resp.json().get("sections")

    assert get_sections is not None, "GET must return the persisted sections"
    assert len(get_sections) == len(post_sections)
    assert get_sections == post_sections


def test_history_reopen_preserves_metadata(client, docx_factory):
    """A later GET (as if reopening from History) still carries sections."""
    resp = client.post(
        "/api/audit",
        files={"file": ("plain.docx", docx_factory(["hello world body"]), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    audit_id = resp.json()["audit_id"]
    first = client.get(f"/api/audit/{audit_id}").json()
    second = client.get(f"/api/audit/{audit_id}").json()
    assert first.get("sections") is not None
    assert first["sections"] == second["sections"]
    assert second["sections"][0]["section_index"] == 0


def test_get_audit_null_sections_for_historical_without_column(client, test_engine):
    """Rows created before section_metadata existed return null — truthful
    unavailable, never a fabricated range. Simulated by inserting a row
    whose section_metadata is None (as a legacy row would be)."""
    from app.models.audit import AuditRecord
    from sqlalchemy.orm import sessionmaker

    Session = sessionmaker(bind=test_engine)
    with Session() as db:
        legacy = AuditRecord(
            id=str(uuid.uuid4()),
            filename="legacy.docx",
            file_size=10,
            weighted_score=0,
            deploy_mode="LOCAL",
            status="completed",
            section_metadata=None,
        )
        db.add(legacy)
        db.commit()
        legacy_id = legacy.id

    resp = client.get(f"/api/audit/{legacy_id}")
    assert resp.status_code == 200
    data = resp.json()
    assert data.get("sections") is None


def test_audit_deletion_removes_section_metadata_safely(client, docx_factory):
    """Deleting an audit that carries section metadata must not error."""
    resp = client.post(
        "/api/audit",
        files={"file": ("plain.docx", docx_factory(["body one"]), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200
    audit_id = resp.json()["audit_id"]
    assert resp.json().get("sections") is not None

    del_resp = client.delete(f"/api/audit/{audit_id}")
    assert del_resp.status_code == 200
    # subsequent GET is 404 — row and metadata removed together
    assert client.get(f"/api/audit/{audit_id}").status_code == 404


def test_migration_roundtrip_section_metadata(tmp_path, monkeypatch):
    from alembic import command
    from alembic.config import Config

    from app.config import settings as app_settings

    db_path = tmp_path / "section.db"
    url = f"sqlite:///{db_path}"
    backend_dir = Path(__file__).resolve().parents[1]

    cfg = Config()
    cfg.set_main_option("script_location", str(backend_dir / "alembic"))

    def _columns():
        return {c["name"] for c in inspect(create_engine(url)).get_columns("audit_records")}

    monkeypatch.setattr(app_settings, "DATABASE_URL", url)

    # fresh install → head has the column
    command.upgrade(cfg, "head")
    assert "section_metadata" in _columns()

    # downgrade to the previous head removes it
    command.downgrade(cfg, "d4e8d19e2b3f")
    assert "section_metadata" not in _columns()

    # re-upgrade restores it (upgrade/downgrade safe)
    command.upgrade(cfg, "head")
    assert "section_metadata" in _columns()


def test_submit_response_existing_fields_untouched(client, docx_factory):
    resp = client.post(
        "/api/audit",
        files={"file": ("plain.docx", docx_factory(["body", "more body"]), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200
    data = resp.json()
    for key in ("status", "audit_id", "weighted_compliance_score", "physical_layout_errors",
                "ai_citation_tooltips", "score_breakdown", "document_stats"):
        assert key in data
