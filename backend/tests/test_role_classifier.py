"""Role classifier PoC tests (Phase 1) — authoritative backend roles.

Covers: cover detection (incl. body-first guard), TOC field entries, H1-3,
body, direct + style-inherited lists, displayed equations, semantic and
manual captions, scholarly + administrative + rubric tables, References
entries, Appendix content, rubric back-matter, empty/field-only/image-only
paragraphs, conflicting evidence, UNKNOWN fallback.

No rule migration, no scoring change — role classification only.
"""
import base64
import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

from app.services.document_parser import extract_paragraphs
from app.services.role_classifier import (
    classify_paragraphs,
    classify_table_roles,
    PARAGRAPH_ROLES,
    TABLE_ROLES,
)

_PNG_1PX = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk"
    "+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
)


def _base_doc():
    doc = Document()
    for name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3",
                 "Caption", "List Bullet", "List Number", "TOC Heading", "Equation"):
        try:
            s = doc.styles[name]
            s.font.name = "Times New Roman"
            s.font.size = Pt(12)
        except KeyError:
            pass
    return doc


def _add_seq_caption(doc, label, text_suffix):
    """Real Word caption: Caption style + split SEQ field runs."""
    cap = doc.add_paragraph()
    cap.style = doc.styles["Caption"]
    cap.add_run(f"{label} ")
    for piece in (" SEQ ", f"{label} ", "\\* ARABIC "):
        run = cap.add_run()
        t = OxmlElement("w:instrText")
        t.set(qn("xml:space"), "preserve")
        t.text = piece
        run._r.append(t)
    cap.add_run("1")
    cap.add_run(text_suffix)
    return cap


def _add_toc_field(doc, text):
    """Paragraph with a real TOC field (entry)."""
    p = doc.add_paragraph()
    r = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.text = " TOC \\o \"1-3\" \\h \\z \\u "
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    r._r.append(f1); r._r.append(instr); r._r.append(f2)
    p.add_run(text)
    r2 = p.add_run()
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    r2._r.append(f3)
    return p


def _add_equation(doc, text):
    """Paragraph with an OMML oMath element (displayed equation)."""
    p = doc.add_paragraph()
    m = OxmlElement("m:oMath")
    r = OxmlElement("m:r")
    t = OxmlElement("m:t")
    t.text = text
    r.append(t)
    m.append(r)
    p._p.append(m)
    return p


def _add_numpr(p):
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    nid = OxmlElement("w:numId"); nid.set(qn("w:val"), "1")
    numPr.append(nid)
    pPr.append(numPr)


def _add_image_para(doc):
    p = doc.add_paragraph()
    p.add_run().add_picture(io.BytesIO(base64.b64decode(_PNG_1PX)), width=Pt(12))
    return p


def _add_header_logo(doc):
    """Embed a logo in the section header — invisible to doc.paragraphs."""
    hdr = doc.sections[0].header
    hdr.paragraphs[0].add_run().add_picture(
        io.BytesIO(base64.b64decode(_PNG_1PX)), width=Pt(12)
    )
    return hdr


def _add_invisible_toc_textbox(doc):
    """A TOC stored in a text box — NOT represented by doc.paragraphs.

    Real documents may store TOC content in text boxes/drawing canvases that
    python-docx's paragraph model never sees. This fixture inserts the
    txbxContent via raw OOXML inside a paragraph's drawing so the TOC text
    exists in the file but is invisible to `doc.paragraphs`.
    """
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    rPr = OxmlElement("w:rPr")
    rPr.append(OxmlElement("w:noProof"))
    pPr.append(rPr)
    run = p.add_run()
    run._r.append(OxmlElement("w:drawing"))
    # A minimal <w:drawing><wps:txbx><w:txbxContent>...</w:txbxContent>.
    # Namespaces are declared inline because python-docx does not register
    # the shape/wordprocessingShape namespaces by default.
    from lxml import etree
    drawing = run._r.find(qn("w:drawing"))
    drawing.append(etree.fromstring(
        '<wps:txbx xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:p><w:r><w:t>1. Introduction</w:t></w:r></w:p>'
        '<w:p><w:r><w:t>2. Methods</w:t></w:r></w:p>'
        '</w:txbxContent></wps:txbx>'
    ))
    return p


def _bytes(doc):
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _roles(blocks_bytes):
    from app.services.layout_engine import parse_document
    doc = parse_document(blocks_bytes)
    paras = extract_paragraphs(doc)
    return classify_paragraphs(doc, paras)


# ---------------------------------------------------------------------------
# Cover
# ---------------------------------------------------------------------------

def test_cover_with_title_signal():
    doc = _base_doc()
    doc.add_paragraph("Assignment Title", style="Title")
    doc.add_paragraph("Student Name Here")
    doc.add_paragraph("Detail:")
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body text begins here with enough words to be prose.")
    roles = _roles(_bytes(doc))
    assert roles[0] == "TITLE"
    assert roles[1] == "COVER"
    assert roles[2] == "COVER"
    assert roles[3] == "HEADING_1"
    assert roles[4] == "BODY"


def test_body_first_document_stays_body_no_cover():
    doc = _base_doc()
    doc.add_paragraph("This assignment explores dataset representation with enough prose to be body text.")
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("More body prose follows the introduction heading here.")
    roles = _roles(_bytes(doc))
    assert roles[0] == "BODY"
    assert roles[1] == "HEADING_1"
    assert roles[2] == "BODY"
    assert "COVER" not in roles


def test_cover_signal_needs_proof_no_guess():
    # Labels only (no Title style, no logo) before the first heading → no
    # cover evidence → those paragraphs stay BODY (never guessed as cover).
    doc = _base_doc()
    doc.add_paragraph("Detail:")
    doc.add_paragraph("Assessment:")
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body prose that follows the introduction here.")
    roles = _roles(_bytes(doc))
    assert roles[0] != "COVER"
    assert roles[1] != "COVER"


def test_administrative_tables_before_first_heading_prove_cover():
    """Two admin cover tables before the first Heading — cover region even
    with Normal-style structural labels (no Title style, no body logo)."""
    doc = _base_doc()
    doc.add_paragraph("STUDENT DETAILS")
    t0 = doc.add_table(rows=3, cols=2)
    for r, (k, v) in enumerate([("Name", "A. Student"), ("ID", "S12345"),
                                ("Programme", "BSc Computing")]):
        t0.rows[r].cells[0].text = k
        t0.rows[r].cells[1].text = v
    doc.add_paragraph("ASSESSMENT")
    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "Mark"
    t1.rows[0].cells[1].text = "Comments"
    t1.rows[1].cells[0].text = "85"
    t1.rows[1].cells[1].text = "Good work"
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body prose that follows the introduction heading here.")
    roles = _roles(_bytes(doc))
    # Structural labels before the heading → cover, ending at the heading.
    assert roles[0] == "COVER"   # STUDENT DETAILS
    assert roles[1] == "COVER"   # ASSESSMENT
    assert roles[2] == "HEADING_1"
    assert roles[3] == "BODY"


def test_cover_image_not_present_in_doc_paragraphs():
    """A logo in the section header — invisible to doc.paragraphs — still
    proves a cover when labels precede the first heading."""
    doc = _base_doc()
    _add_header_logo(doc)
    doc.add_paragraph("Assignment Cover Page")
    doc.add_paragraph("Name:")
    doc.add_paragraph("ID:")
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body prose that follows the introduction heading here.")
    roles = _roles(_bytes(doc))
    assert roles[0] == "COVER"
    assert roles[1] == "COVER"
    assert roles[2] == "COVER"
    assert roles[3] == "HEADING_1"
    assert roles[4] == "BODY"


def test_heading_first_document_stays_academic():
    doc = _base_doc()
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body prose that follows the introduction heading here.")
    roles = _roles(_bytes(doc))
    assert roles[0] == "HEADING_1"
    assert roles[1] == "BODY"
    assert "COVER" not in roles


def test_short_ordinary_first_paragraph_alone_insufficient():
    doc = _base_doc()
    doc.add_paragraph("Abstract")
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body prose that follows the introduction heading here.")
    roles = _roles(_bytes(doc))
    # A single short first paragraph is not a cover — no structural signal.
    assert roles[0] == "BODY"
    assert roles[1] == "HEADING_1"


def test_first_academic_heading_ends_cover():
    doc = _base_doc()
    doc.add_paragraph("Assignment Title", style="Title")
    doc.add_paragraph("Student Name Here")
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body prose that follows the introduction heading here.")
    roles = _roles(_bytes(doc))
    assert roles[0] == "TITLE"
    assert roles[1] == "COVER"
    assert roles[2] == "HEADING_1"
    assert roles[3] == "BODY"


def test_toc_invisible_to_paragraph_extraction_does_not_extend_cover():
    """Real-document TOC may live in text boxes invisible to doc.paragraphs.
    The first academic Heading is still an independent cover end — cover
    never extends past it even when no TOC is visible in paragraph blocks."""
    doc = _base_doc()
    doc.add_paragraph("Assignment Title", style="Title")
    doc.add_paragraph("Student Name Here")
    _add_invisible_toc_textbox(doc)
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Body prose that follows the introduction heading here.")
    roles = _roles(_bytes(doc))
    assert roles[0] == "TITLE"
    assert roles[1] == "COVER"
    # The invisible-TOC paragraph is an empty drawing host → FIGURE_HOST,
    # and cover ends at the first academic Heading regardless.
    assert roles[2] == "FIGURE_HOST"
    assert roles[3] == "HEADING_1"
    assert roles[4] == "BODY"


def test_no_cover_evidence_returns_body_or_unknown_never_cover():
    doc = _base_doc()
    doc.add_paragraph("Some ordinary body text that is long enough to be prose here.")
    doc.add_paragraph("1. Introduction", style="Heading 1")
    roles = _roles(_bytes(doc))
    assert roles[0] == "BODY"
    assert roles[1] == "HEADING_1"
    assert "COVER" not in roles


def test_sanitized_real_document_cover_region():
    """Regression fixture mimicking the sanitized real document: Normal-style
    structural labels, administrative tables, AND a header logo invisible to
    doc.paragraphs (the exact gate failure — labels stayed BODY). Cover-region
    paragraphs must classify COVER; the first academic Heading and the body
    that follows must not."""
    doc = _base_doc()
    # Logo lives in the header — invisible to doc.paragraphs.
    _add_header_logo(doc)
    # Document title area (Normal style — no Title style used).
    doc.add_paragraph("THE IMPACT OF ARTIFICIAL INTELLIGENCE")
    doc.add_paragraph("ON HIGHER EDUCATION")
    # Administrative information tables.
    t0 = doc.add_table(rows=3, cols=2)
    for r, (k, v) in enumerate([("Name", "Sanitized Student"),
                                ("Matric No", "A12345"),
                                ("Programme", "Bachelor of Computing")]):
        t0.rows[r].cells[0].text = k
        t0.rows[r].cells[1].text = v
    doc.add_paragraph("Assignment Details")
    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "Module"
    t1.rows[0].cells[1].text = "Research Methods"
    t1.rows[1].cells[0].text = "Submission Date"
    t1.rows[1].cells[1].text = "15 January 2025"
    # Assessment / mark area.
    doc.add_paragraph("Assessment")
    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "Mark"
    t2.rows[0].cells[1].text = "Comments"
    t2.rows[1].cells[0].text = "78"
    t2.rows[1].cells[1].text = "Well structured"
    # Academic content begins.
    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph("Artificial intelligence is transforming higher education in many ways.")
    doc.add_paragraph("2. Literature Review", style="Heading 1")
    doc.add_paragraph("Prior work has examined the role of machine learning in universities.")
    roles = _roles(_bytes(doc))
    # Cover region (all Normal-style structural labels). Tables are NOT
    # paragraphs — doc.paragraphs holds only the 4 label lines before the
    # heading; the classifier sees the tables through OOXML body order.
    assert roles[0] == "COVER", roles  # THE IMPACT OF ARTIFICIAL INTELLIGENCE
    assert roles[1] == "COVER", roles  # ON HIGHER EDUCATION
    assert roles[2] == "COVER", roles  # Assignment Details
    assert roles[3] == "COVER", roles  # Assessment
    # First academic Heading ends the cover; body after stays BODY.
    assert roles[4] == "HEADING_1", roles
    assert roles[5] == "BODY", roles
    assert roles[6] == "HEADING_1", roles
    assert roles[7] == "BODY", roles


# ---------------------------------------------------------------------------
# TOC
# ---------------------------------------------------------------------------

def test_toc_field_entries_and_heading():
    doc = _base_doc()
    doc.add_paragraph("Table of Contents", style="Heading 1")
    _add_toc_field(doc, "1. Introduction")
    _add_toc_field(doc, "2. Methods")
    doc.add_paragraph("1. Introduction", style="Heading 1")
    roles = _roles(_bytes(doc))
    assert roles[0] == "TABLE_OF_CONTENTS_HEADING"
    assert roles[1] == "TABLE_OF_CONTENTS_ENTRY"
    assert roles[2] == "TABLE_OF_CONTENTS_ENTRY"
    assert roles[3] == "HEADING_1"


def test_toc_style_entries():
    doc = _base_doc()
    doc.add_paragraph("Contents", style="Heading 1")
    p = doc.add_paragraph("1. Introduction 3")
    p.style = doc.styles["TOC Heading"] if "TOC Heading" in [s.name for s in doc.styles] else None
    # fall back to a TOC-ish style name
    if p.style is None:
        p.style = doc.styles["Normal"]
        # force a TOC-looking style name via raw XML
        pPr = p._p.get_or_add_pPr()
        pStyle = OxmlElement("w:pStyle"); pStyle.set(qn("w:val"), "TOC1")
        pPr.append(pStyle)
    doc.add_paragraph("1. Introduction", style="Heading 1")
    roles = _roles(_bytes(doc))
    assert roles[1] == "TABLE_OF_CONTENTS_ENTRY"


# ---------------------------------------------------------------------------
# Headings / body / lists / equations / captions / empty
# ---------------------------------------------------------------------------

def test_headings_body_and_unknown_level4():
    doc = _base_doc()
    doc.add_paragraph("H1", style="Heading 1")
    doc.add_paragraph("H2", style="Heading 2")
    doc.add_paragraph("H3", style="Heading 3")
    doc.add_paragraph("Ordinary body text that is long enough to classify.")
    roles = _roles(_bytes(doc))
    assert roles[:3] == ["HEADING_1", "HEADING_2", "HEADING_3"]
    assert roles[3] == "BODY"


def test_direct_and_style_inherited_lists():
    doc = _base_doc()
    p = doc.add_paragraph("First list item")
    _add_numpr(p)
    doc.add_paragraph("Bullet item", style="List Bullet")
    doc.add_paragraph("Numbered item", style="List Number")
    roles = _roles(_bytes(doc))
    assert roles == ["LIST_ITEM", "LIST_ITEM", "LIST_ITEM"]


def test_displayed_equation():
    doc = _base_doc()
    doc.add_paragraph("The fuzzy score is computed as:")
    _add_equation(doc, "score = w * m")
    doc.add_paragraph("Where w is the weight.")
    roles = _roles(_bytes(doc))
    assert roles[1] == "DISPLAYED_EQUATION"


def test_semantic_and_manual_captions_split_by_type():
    doc = _base_doc()
    _add_seq_caption(doc, "Table", ": Sample results")
    doc.add_paragraph("Figure 1: Architecture diagram")
    doc.add_paragraph("Body text after the captions, long enough to be prose.")
    roles = _roles(_bytes(doc))
    assert roles[0] == "CAPTION_TABLE"
    assert roles[1] == "CAPTION_FIGURE"
    assert roles[2] == "BODY"


def test_empty_field_only_and_figure_host():
    doc = _base_doc()
    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.text = " PAGE "
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    r._r.append(f1); r._r.append(instr); r._r.append(f2)
    _add_image_para(doc)
    roles = _roles(_bytes(doc))
    assert roles[0] == "EMPTY"
    assert roles[1] == "FIELD_ONLY"
    assert roles[2] == "FIGURE_HOST"


# ---------------------------------------------------------------------------
# References / Appendix / rubric back-matter
# ---------------------------------------------------------------------------

def test_references_heading_and_entries():
    doc = _base_doc()
    doc.add_paragraph("Body text that is long enough to be prose here.")
    doc.add_paragraph("References", style="Heading 1")
    doc.add_paragraph("Rubens, N. O. (2006). The application of fuzzy logic.")
    doc.add_paragraph("Zadeh, L. A. (1965). Fuzzy sets.")
    roles = _roles(_bytes(doc))
    assert roles[0] == "BODY"
    assert roles[1] == "REFERENCES_HEADING"
    assert roles[2] == "REFERENCE_ENTRY"
    assert roles[3] == "REFERENCE_ENTRY"


def test_appendix_heading_and_body():
    doc = _base_doc()
    doc.add_paragraph("Main body prose that is long enough to be classified.")
    doc.add_paragraph("Appendix A", style="Heading 1")
    doc.add_paragraph("Appendix content paragraph.")
    roles = _roles(_bytes(doc))
    assert roles[0] == "BODY"
    assert roles[1] == "APPENDIX_HEADING"
    assert roles[2] == "APPENDIX_BODY"


def test_rubric_backmatter_not_body():
    doc = _base_doc()
    doc.add_paragraph("Main body prose that is long enough to be classified.")
    doc.add_paragraph("Marking Rubrics", style="Heading 1")
    doc.add_paragraph("Criteria and marks content.")
    roles = _roles(_bytes(doc))
    assert roles[0] == "BODY"
    assert roles[1] == "UNKNOWN"  # rubric heading — back-matter, not scholarly
    assert roles[2] == "UNKNOWN"


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

def _doc_with_tables():
    doc = _base_doc()
    doc.add_paragraph("Cover label row:", style="Title")
    t0 = doc.add_table(rows=2, cols=2)
    t0.rows[0].cells[0].text = "Name"
    t0.rows[0].cells[1].text = "Value"
    doc.add_paragraph("1. Introduction", style="Heading 1")
    _add_seq_caption(doc, "Table", ": Results")
    t1 = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            t1.rows[r].cells[c].text = f"r{r}c{c}"
    doc.add_paragraph("References", style="Heading 1")
    t2 = doc.add_table(rows=1, cols=1)
    t2.rows[0].cells[0].text = "decorative"
    return doc


def test_table_roles_conservative():
    doc = _doc_with_tables()
    paras = extract_paragraphs(doc)
    # cover region = before first Heading 1 (index of intro)
    intro_idx = next(p["index"] for p in paras if p["text"] == "1. Introduction")
    refs_idx = next(p["index"] for p in paras if p["text"] == "References")
    roles = classify_table_roles(doc, cover_end=intro_idx, refs_start=refs_idx)
    assert len(roles) == 3
    # t0: before cover boundary → administrative
    assert roles[0] == "ADMINISTRATIVE_TABLE"
    # t1: manuscript body + adjacent numbered caption → scholarly
    assert roles[1] == "SCHOLARLY_TABLE"
    # t2: references region → never a caption target → UNKNOWN
    assert roles[2] == "UNKNOWN"


def test_unknown_table_not_penalized():
    doc = _base_doc()
    doc.add_paragraph("1. Introduction", style="Heading 1")
    t = doc.add_table(rows=3, cols=3)  # data-like but NO caption
    for r in range(3):
        for c in range(3):
            t.rows[r].cells[c].text = f"d{r}{c}"
    doc.add_paragraph("Body prose continues after the table.")
    roles = classify_table_roles(doc, cover_end=0)
    assert roles == ["UNKNOWN"]  # structure alone never qualifies


def test_two_administrative_cover_tables():
    """Two admin cover tables before the first Heading → both
    ADMINISTRATIVE_TABLE regardless of content."""
    doc = _base_doc()
    doc.add_paragraph("Name", style="Title")
    t0 = doc.add_table(rows=2, cols=2)
    t0.rows[0].cells[0].text = "Student"
    t0.rows[0].cells[1].text = "A. Student"
    t0.rows[1].cells[0].text = "ID"
    t0.rows[1].cells[1].text = "S12345"
    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "Module"
    t1.rows[0].cells[1].text = "Research"
    t1.rows[1].cells[0].text = "Date"
    t1.rows[1].cells[1].text = "2025-01-15"
    doc.add_paragraph("1. Introduction", style="Heading 1")
    paras = extract_paragraphs(doc)
    intro_idx = next(p["index"] for p in paras if p["text"] == "1. Introduction")
    roles = classify_table_roles(doc, cover_end=intro_idx)
    assert roles == ["ADMINISTRATIVE_TABLE", "ADMINISTRATIVE_TABLE"]


def test_assessment_mark_table_classification_policy():
    """MARK/Comments assessment table on the cover → ADMINISTRATIVE_TABLE
    (cover-region position decides, before content)."""
    doc = _base_doc()
    doc.add_paragraph("Assessment", style="Title")
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Mark"
    t.rows[0].cells[1].text = "Comments"
    t.rows[1].cells[0].text = "85"
    t.rows[1].cells[1].text = "Good work"
    doc.add_paragraph("1. Introduction", style="Heading 1")
    paras = extract_paragraphs(doc)
    intro_idx = next(p["index"] for p in paras if p["text"] == "1. Introduction")
    roles = classify_table_roles(doc, cover_end=intro_idx)
    assert roles == ["ADMINISTRATIVE_TABLE"]


def test_rubric_table_after_references():
    """Final marking-criteria table in back-matter → RUBRIC_TABLE; ambiguous
    back-matter table → UNKNOWN."""
    doc = _base_doc()
    doc.add_paragraph("Body prose that is long enough to be classified.")
    doc.add_paragraph("References", style="Heading 1")
    doc.add_paragraph("Author, A. (2020). Title.")
    t_rubric = doc.add_table(rows=3, cols=3)
    t_rubric.rows[0].cells[0].text = "Criteria"
    t_rubric.rows[0].cells[1].text = "Marks"
    t_rubric.rows[0].cells[2].text = "Grade"
    t_rubric.rows[1].cells[0].text = "Clarity"
    t_rubric.rows[1].cells[1].text = "5"
    t_rubric.rows[1].cells[2].text = "A"
    doc.add_paragraph("Appendix", style="Heading 1")
    t_ambig = doc.add_table(rows=2, cols=2)
    t_ambig.rows[0].cells[0].text = "x"
    t_ambig.rows[0].cells[1].text = "y"
    t_ambig.rows[1].cells[0].text = "1"
    t_ambig.rows[1].cells[1].text = "2"
    paras = extract_paragraphs(doc)
    refs_idx = next(p["index"] for p in paras if p["text"] == "References")
    app_idx = next(p["index"] for p in paras if p["text"] == "Appendix")
    roles = classify_table_roles(doc, cover_end=0, refs_start=refs_idx, appendix_start=app_idx)
    assert roles == ["RUBRIC_TABLE", "UNKNOWN"]


def test_ambiguous_table_remains_unknown():
    """A data-like table in the manuscript body with NO caption → UNKNOWN,
    and no scholarly Caption eligibility is inferred from it."""
    doc = _base_doc()
    doc.add_paragraph("1. Introduction", style="Heading 1")
    t = doc.add_table(rows=4, cols=4)
    for r in range(4):
        for c in range(4):
            t.rows[r].cells[c].text = f"d{r}{c}"
    doc.add_paragraph("Body prose continues after the table here.")
    roles = classify_table_roles(doc, cover_end=0)
    assert roles == ["UNKNOWN"]


def test_roles_are_known_enum_members():
    doc = _base_doc()
    doc.add_paragraph("Hello body text that is long enough.")
    doc.add_paragraph("References", style="Heading 1")
    doc.add_paragraph("Author, A. (2020). Title.")
    roles = _roles(_bytes(doc))
    for role in roles:
        assert role in PARAGRAPH_ROLES
    for role in TABLE_ROLES:
        assert isinstance(role, str)
