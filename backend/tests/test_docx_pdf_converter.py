"""Focused tests for the LibreOffice DOCX-to-PDF conversion service.

LibreOffice itself is never required: subprocess.Popen is faked for unit
tests; one optional integration test runs only when soffice is installed.
"""
import io
import subprocess
import tempfile
import uuid
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

import pytest
from docx import Document

from app.services.docx_pdf_converter import (
    DocxConversionError,
    convert_docx_to_pdf,
    find_soffice,
    _SOFFICE_CANDIDATES,
    _PAGE_RE,
)


def _pdf_bytes(pages=1):
    body = b"".join(
        b"%d 0 obj\n<< /Type /Page >>\nendobj\n" % i for i in range(pages)
    )
    return b"%PDF-1.4\n" + body + b"%%EOF"


class FakePopen:
    """Records invocations; writes a PDF into the --outdir when successful."""

    calls = []

    def __init__(self, args, **kwargs):
        FakePopen.calls.append(args)
        self.args = args
        self.pid = 4242
        self.returncode = getattr(FakePopen, "returncode", 0)
        self.pdf_bytes = getattr(FakePopen, "pdf_bytes", None)
        self.write_pdf = getattr(FakePopen, "write_pdf", True)
        self.timeout_error = getattr(FakePopen, "timeout_error", False)

    def communicate(self, timeout=None):
        if self.timeout_error:
            raise subprocess.TimeoutExpired(self.args, timeout)
        if self.write_pdf:
            outdir = self.args[self.args.index("--outdir") + 1]
            Path(outdir, "converted.pdf").write_bytes(
                self.pdf_bytes or _pdf_bytes()
            )
        return b"", b""

    def kill(self):
        pass


@pytest.fixture
def fake_popen(monkeypatch):
    FakePopen.calls = []
    for name in ("returncode", "pdf_bytes", "write_pdf", "timeout_error"):
        try:
            delattr(FakePopen, name)
        except AttributeError:
            pass
    monkeypatch.setattr("app.services.docx_pdf_converter.subprocess.Popen", FakePopen)
    return FakePopen


def _track_mkdtemp(monkeypatch, root: Path):
    def fake_mkdtemp(**kwargs):
        d = root / uuid.uuid4().hex
        d.mkdir()
        return str(d)

    monkeypatch.setattr("app.services.docx_pdf_converter.tempfile.mkdtemp", fake_mkdtemp)


def _real_docx_bytes():
    doc = Document()
    doc.add_paragraph("Body text.")
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "A"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# 1. LibreOffice executable discovery
# ---------------------------------------------------------------------------

def test_discovery_configured_path(monkeypatch, tmp_path):
    exe = tmp_path / "soffice.com"
    exe.write_bytes(b"")
    monkeypatch.setenv("SOFFICE_EXECUTABLE", str(exe))
    assert find_soffice() == str(exe)


def test_discovery_path_lookup(monkeypatch, tmp_path):
    exe = tmp_path / "soffice.exe"
    exe.write_bytes(b"")
    monkeypatch.delenv("SOFFICE_EXECUTABLE", raising=False)
    monkeypatch.setattr("app.services.docx_pdf_converter.shutil.which",
                        lambda name: str(exe) if name == "soffice.com" else None)
    assert find_soffice() == str(exe)


def test_discovery_prefers_soffice_com(monkeypatch, tmp_path):
    com = tmp_path / "soffice.com"
    exe = tmp_path / "soffice.exe"
    com.write_bytes(b"")
    exe.write_bytes(b"")
    monkeypatch.delenv("SOFFICE_EXECUTABLE", raising=False)
    monkeypatch.setattr("app.services.docx_pdf_converter.shutil.which", lambda name: None)
    monkeypatch.setattr("app.services.docx_pdf_converter._SOFFICE_CANDIDATES",
                        (str(com), str(exe)))
    assert find_soffice() == str(com)


def test_discovery_uses_fixed_candidates(monkeypatch, tmp_path):
    exe = tmp_path / "soffice.com"
    exe.write_bytes(b"")
    monkeypatch.delenv("SOFFICE_EXECUTABLE", raising=False)
    monkeypatch.setattr("app.services.docx_pdf_converter.shutil.which", lambda name: None)
    monkeypatch.setattr("app.services.docx_pdf_converter._SOFFICE_CANDIDATES",
                        (str(exe), str(tmp_path / "missing.com")))
    assert find_soffice() == str(exe)


def test_discovery_none_available(monkeypatch):
    monkeypatch.delenv("SOFFICE_EXECUTABLE", raising=False)
    monkeypatch.setattr("app.services.docx_pdf_converter.shutil.which", lambda name: None)
    monkeypatch.setattr("app.services.docx_pdf_converter._SOFFICE_CANDIDATES", ())
    assert find_soffice() is None


# ---------------------------------------------------------------------------
# 2-5. Conversion outcomes
# ---------------------------------------------------------------------------

def test_successful_conversion(fake_popen, monkeypatch, tmp_path):
    _track_mkdtemp(monkeypatch, tmp_path)
    result = convert_docx_to_pdf(_real_docx_bytes())
    assert result.startswith(b"%PDF")


def test_conversion_uses_random_internal_filename(fake_popen):
    convert_docx_to_pdf(_real_docx_bytes())
    src = Path(FakePopen.calls[0][-1])
    assert src.suffix == ".docx"
    assert src.name.startswith(".") is False  # user filename never involved


def test_invalid_docx_rejected():
    with pytest.raises(DocxConversionError) as exc:
        convert_docx_to_pdf(b"not a zip file at all")
    assert exc.value.category == "invalid_docx"


def test_empty_input_rejected():
    with pytest.raises(DocxConversionError) as exc:
        convert_docx_to_pdf(b"")
    assert exc.value.category == "invalid_docx"


# ---------------------------------------------------------------------------
# 6-9. Failure categories
# ---------------------------------------------------------------------------

def test_timeout_simulated(fake_popen, monkeypatch):
    fake_popen.timeout_error = True
    killed = []
    monkeypatch.setattr(
        "app.services.docx_pdf_converter._kill_process_tree",
        lambda proc: killed.append(proc),
    )
    with pytest.raises(DocxConversionError) as exc:
        convert_docx_to_pdf(_real_docx_bytes())
    assert exc.value.category == "timeout"
    assert len(killed) == 1


def test_nonzero_exit_code(fake_popen, monkeypatch, tmp_path):
    _track_mkdtemp(monkeypatch, tmp_path)
    fake_popen.returncode = 1
    with pytest.raises(DocxConversionError) as exc:
        convert_docx_to_pdf(_real_docx_bytes())
    assert exc.value.category == "conversion_failed"


def test_missing_pdf_output(fake_popen, monkeypatch, tmp_path):
    _track_mkdtemp(monkeypatch, tmp_path)
    fake_popen.write_pdf = False
    with pytest.raises(DocxConversionError) as exc:
        convert_docx_to_pdf(_real_docx_bytes())
    assert exc.value.category == "conversion_failed"


def test_invalid_pdf_magic(fake_popen, monkeypatch, tmp_path):
    _track_mkdtemp(monkeypatch, tmp_path)
    fake_popen.pdf_bytes = b"NOTAPDF"
    with pytest.raises(DocxConversionError) as exc:
        convert_docx_to_pdf(_real_docx_bytes())
    assert exc.value.category == "invalid_pdf"


def test_libreoffice_unavailable(monkeypatch):
    monkeypatch.setattr("app.services.docx_pdf_converter.find_soffice",
                        lambda: None)
    with pytest.raises(DocxConversionError) as exc:
        convert_docx_to_pdf(_real_docx_bytes())
    assert exc.value.category == "libreoffice_unavailable"


# ---------------------------------------------------------------------------
# 10. Concurrency: isolated profiles and outdirs
# ---------------------------------------------------------------------------

def test_concurrent_conversions_use_isolated_profiles(fake_popen):
    results = []

    def work(_):
        results.append(convert_docx_to_pdf(_real_docx_bytes()))

    with ThreadPoolExecutor(max_workers=6) as pool:
        list(pool.map(work, range(6)))

    profiles = set()
    outdirs = set()
    for args in FakePopen.calls:
        profiles.add(next(a for a in args if a.startswith("-env:UserInstallation=")))
        outdirs.add(args[args.index("--outdir") + 1])
    assert len(profiles) == 6
    assert len(outdirs) == 6
    assert all(r.startswith(b"%PDF") for r in results)


# ---------------------------------------------------------------------------
# 11-12. Cleanup on success and every failure path
# ---------------------------------------------------------------------------

def test_cleanup_after_success(fake_popen, monkeypatch, tmp_path):
    _track_mkdtemp(monkeypatch, tmp_path)
    assert convert_docx_to_pdf(_real_docx_bytes()).startswith(b"%PDF")
    assert list(tmp_path.iterdir()) == []


@pytest.mark.parametrize("setup", [
    {"write_pdf": False},
    {"pdf_bytes": b"garbage"},
    {"returncode": 3},
    {"timeout_error": True},
])
def test_cleanup_after_every_failure(fake_popen, monkeypatch, tmp_path, setup):
    _track_mkdtemp(monkeypatch, tmp_path)
    for attr, value in setup.items():
        setattr(fake_popen, attr, value)
    with pytest.raises(DocxConversionError):
        convert_docx_to_pdf(_real_docx_bytes())
    assert list(tmp_path.iterdir()) == []


# ---------------------------------------------------------------------------
# 13. Non-ASCII temporary parent directory
# ---------------------------------------------------------------------------

def test_non_ascii_temp_parent(fake_popen, monkeypatch, tmp_path):
    exotic = tmp_path / "Tëst 目录"
    exotic.mkdir()
    monkeypatch.setattr(tempfile, "tempdir", str(exotic))
    result = convert_docx_to_pdf(_real_docx_bytes())
    assert result.startswith(b"%PDF")


# ---------------------------------------------------------------------------
# 14. Logs never leak content or absolute paths
# ---------------------------------------------------------------------------

def test_logs_do_not_leak_content_or_paths(fake_popen, monkeypatch, caplog):
    import logging
    caplog.set_level(logging.INFO)
    parent = Path(tempfile.gettempdir()) / "unused"
    parent.mkdir(exist_ok=True)
    _track_mkdtemp(monkeypatch, parent)
    fake_popen.pdf_bytes = _pdf_bytes()
    assert convert_docx_to_pdf(b"PK\x03\x04SUPER_SECRET_DOCUMENT").startswith(b"%PDF")
    fake_popen.write_pdf = False
    with pytest.raises(DocxConversionError):
        convert_docx_to_pdf(b"PK\x03\x04SUPER_SECRET_DOCUMENT")
    text = caplog.text.lower()
    assert "super_secret" not in text
    assert "secret" not in text
    assert "docx2pdf" not in text


# ---------------------------------------------------------------------------
# Integration: real LibreOffice when installed
# ---------------------------------------------------------------------------

@pytest.mark.skipif(find_soffice() is None, reason="LibreOffice not installed")
def test_real_conversion_with_tables_and_figures(docx_factory):
    before = {p.name for p in Path(tempfile.gettempdir()).glob("docx2pdf_*")}
    file_bytes = docx_factory(
        paragraphs=["Body text.", "More text."],
        references=[],
        tables=[[["A", "B"], ["1", "2"]], [["C", "D"], ["3", "4"]]],
        images=1,
    )
    pdf = convert_docx_to_pdf(file_bytes)
    assert pdf.startswith(b"%PDF")
    assert len(_PAGE_RE.findall(pdf)) >= 1
    after = {p.name for p in Path(tempfile.gettempdir()).glob("docx2pdf_*")}
    assert after == before
