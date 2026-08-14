"""Focused tests for the APA 7 guidance improvement (guidance build).

Covers the fabricated-detail-safe assembly contract:
- structured suggestion hierarchy
- journal / book / webpage template structure
- webpage template only for a known webpage source
- unknown source type is never guessed
- placeholders instead of invented bibliographic details
- placeholder warning always present
- reason sanitisation
- same-paragraph findings keep distinct guidance
- end-to-end POST assembly with the real task and a mocked provider
"""
import json

import pytest

from app.services import ai_citation
from app.services.ai_citation import (
    AI_STATUS_WITH_SUGGESTIONS,
    _validate_source_type,
    _sanitise_reason,
    _reject_subjective_reason,
    _build_personalised_correction,
    build_apa_suggestion,
    async_ai_citation_task,
)

WARNING = "Formatting example only. Replace all placeholders with verified source information before submission."
JOURNAL_TEMPLATE = "Author, A. A. (Year). Title of the article. Journal Name, volume(issue), page\u2013page. https://doi.org/xxxxx"
BOOK_TEMPLATE = "Author, A. A. (Year). Title of the book. Publisher."
WEBPAGE_TEMPLATE = "Author, A. A. (Year, Month Day). Title of the page. Site Name. URL"

REASON = "APA 7 requires every in-text citation to have a matching References entry. Add the missing entry using the source details you already have."


@pytest.fixture
def ai_session(test_engine):
    from sqlalchemy.orm import sessionmaker
    Session = sessionmaker(bind=test_engine)
    s = Session()
    try:
        yield s
    finally:
        s.close()


def _finding(idx, author="Smith", year="2020"):
    return {
        "finding_key": f"fk-{idx}-{author}-{year}",
        "paragraph_index": idx,
        "rule_code": "CITATION_MISMATCH",
        "severity": "MAJOR",
        "snippet": f"({author}, {year})",
        "message": f"Citation '{author} ({year})' was found in text, but no matching entry was found in the References bibliography.",
        "expected_value": f"Reference entry for {author}",
        "actual_value": f"({author}, {year})",
    }


# ---------------------------------------------------------------------------
# Source-type validation — never guessed
# ---------------------------------------------------------------------------

def test_source_type_whitelist():
    assert _validate_source_type("journal_article") == "journal_article"
    assert _validate_source_type("book") == "book"
    assert _validate_source_type("webpage") == "webpage"
    assert _validate_source_type("JOURNAL_ARTICLE") == "journal_article"
    assert _validate_source_type("  Book  ") == "book"


def test_source_type_unknown_not_guessed():
    assert _validate_source_type("conference paper") == "unknown"
    assert _validate_source_type("") == "unknown"
    assert _validate_source_type(None) == "unknown"
    assert _validate_source_type(123) == "unknown"


def test_reason_sanitisation():
    assert _sanitise_reason("  Suggestion: Fix it now.  ") == "Fix it now."
    assert _sanitise_reason("Line one\n\nLine two") == "Line one Line two"
    assert _sanitise_reason("   ") is None
    assert _sanitise_reason(None) is None
    assert _sanitise_reason(42) is None
    long = "x" * 800
    assert len(_sanitise_reason(long)) == 700


# ---------------------------------------------------------------------------
# Neutral wording — subjective/accusatory guidance is rejected
# ---------------------------------------------------------------------------

def test_neutral_reason_kept():
    reason = "No matching References entry was found. Add a matching reference entry and verify the source details."
    assert _reject_subjective_reason(reason) == reason


@pytest.mark.parametrize("phrase", [
    "This is an academic integrity violation.",
    "This citation error is misconduct.",
    "This mistake risks losing credibility.",
    "A citation mismatch damages your credibility.",
    "You must create the missing citation now.",
    "Please create a citation for this source.",
    "Create citations for every source you use.",
])
def test_subjective_reason_rejected(phrase):
    assert _reject_subjective_reason(f"Reason text. {phrase} More text.") is None


def test_subjective_reason_not_case_sensitive():
    assert _reject_subjective_reason("ACADEMIC INTEGRITY matters here.") is None


def test_reference_entry_phrasing_present_in_prompt():
    """Guidance must say 'reference entry', never 'create citation'."""
    prompt = ai_citation.CITATION_GUIDANCE_PROMPT
    assert "reference entry" in prompt
    assert "create the missing citation" in prompt  # listed as forbidden
    assert "create a citation" in prompt


async def test_subjective_guidance_dropped_from_results(ai_session, monkeypatch):
    """Subjective reason is rejected end-to-end — no suggestion persisted."""
    findings = [_finding(0)]

    async def _fake(prompt):
        return json.dumps([
            {"finding_key": findings[0]["finding_key"],
             "reason": "This is an academic integrity violation.",
             "source_type": "book", "confidence": 0.9},
        ])

    monkeypatch.setattr(ai_citation, "call_ollama_local", _fake)

    out = await async_ai_citation_task(
        audit_id="tpl-5", db=ai_session, cloud=False, citation_findings=findings)
    ai_session.commit()
    assert out.suggestions == []
    assert out.status == "COMPLETED_NO_SUGGESTIONS"


def test_shared_template_blocks_identical_across_findings():
    """Same-paragraph findings share one template/checklist block, distinct reasons.

    The stored suggestions differ only in the personalised correction;
    the shared verification checklist + templates are byte-identical so the
    UI can render them once instead of duplicating per finding.
    """
    sug_a = build_apa_suggestion("Correct the Garcia (2018) entry.", "journal_article")
    sug_b = build_apa_suggestion("Correct the Lee (2021) entry.", "journal_article")

    def _split(sug):
        sections = sug.split("\n\n")
        correction = sections[0].replace("Recommended correction\n", "")
        shared = "\n\n".join(sections[1:])
        return correction, shared

    corr_a, shared_a = _split(sug_a)
    corr_b, shared_b = _split(sug_b)
    assert corr_a != corr_b, "personalised corrections must stay distinct"
    assert shared_a == shared_b, "shared checklist + templates must not be duplicated per finding"
    assert "Garcia" not in shared_a and "Lee" not in shared_b


# ---------------------------------------------------------------------------
# Deterministic personalised correction wording
# ---------------------------------------------------------------------------

def test_correction_states_no_matching_entry_with_author_year():
    corr = _build_personalised_correction("Garcia", "2018")
    assert corr.startswith(
        "No matching References entry was found for Garcia (2018) in this document."
    )
    assert "in this document" in corr


def test_correction_includes_verify_and_add_actions():
    corr = _build_personalised_correction("Garcia", "2018")
    assert "Verify the original source details" in corr
    assert "add the corresponding APA 7 reference entry" in corr


def test_correction_includes_correct_or_remove_alternative():
    corr = _build_personalised_correction("Garcia", "2018")
    assert "refers to the wrong source" in corr
    assert "correct or remove it instead" in corr


def test_correction_never_names_other_authors():
    garcia = _build_personalised_correction("Garcia", "2018")
    lee = _build_personalised_correction("Lee", "2021")
    assert "Garcia" in garcia and "Lee" not in garcia
    assert "Lee" in lee and "Garcia" not in lee


@pytest.mark.parametrize("bad", [
    "No matching reference exists.",
    "You must add a missing reference entry.",
    "Create the missing citation.",
    "This is an academic integrity violation.",
])
def test_correction_excludes_absolute_and_subjective_wording(bad):
    corr = _build_personalised_correction("Garcia", "2018")
    assert bad.lower() not in corr.lower()
    for phrase in ("academic integrity", "credibility", "misconduct"):
        assert phrase not in corr.lower()


async def test_garcia_and_lee_same_paragraph_get_distinct_deterministic_guidance(ai_session, monkeypatch):
    """Same-paragraph Garcia + Lee findings map to separate deterministic corrections."""
    findings = [_finding(21, "Garcia", "2018"), _finding(21, "Lee", "2021")]

    async def _fake(prompt):
        return json.dumps([
            {"finding_key": findings[0]["finding_key"], "reason": "Neutral reason A.", "source_type": "unknown", "confidence": 0.9},
            {"finding_key": findings[1]["finding_key"], "reason": "Neutral reason B.", "source_type": "unknown", "confidence": 0.9},
        ])

    monkeypatch.setattr(ai_citation, "call_ollama_local", _fake)

    out = await async_ai_citation_task(
        audit_id="tpl-6", db=ai_session, cloud=False, citation_findings=findings)
    ai_session.commit()

    assert out.status == AI_STATUS_WITH_SUGGESTIONS
    assert len(out.suggestions) == 2
    sug_garcia = out.suggestions[0]["suggestion"]
    sug_lee = out.suggestions[1]["suggestion"]
    assert "Garcia (2018)" in sug_garcia and "Lee" not in sug_garcia
    assert "Lee (2021)" in sug_lee and "Garcia" not in sug_lee
    assert "in this document" in sug_garcia and "in this document" in sug_lee
    assert out.suggestions[0]["paragraph_index"] == 21
    assert out.suggestions[1]["paragraph_index"] == 21


def test_score_recalculated_without_false_positive():
    """Fixture score: removing the false Smith finding leaves 2 mismatches.

    citation_apa MAJOR weight = 5, so 2 mismatches deduct 10 → score 90.
    (With the false Smith finding it would have been 3 × 5 = 15 → 85.)
    """
    from app.services.scoring import calculate_weighted_score_detailed
    from docx import Document
    from app.services.citation_sensor import run_citation_sensor
    from app.services.document_parser import extract_paragraphs

    doc = Document()
    doc.add_paragraph("Smith (2020) reported the first finding.")
    doc.add_paragraph("However, Garcia (2018) and Lee (2021) disagreed.")
    doc.add_paragraph("References")
    doc.add_paragraph("1. Smith, J. (2020). On things. Press.")
    doc.add_paragraph("2. Jones, A. (2019). Other things. Press.")
    doc.add_paragraph("Appendix A: Manifest")
    doc.add_paragraph("Manifest lists Smith (2020) and Garcia (2018) as verified sources.")

    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert len(viols) == 2
    result = calculate_weighted_score_detailed(viols)
    assert result.total == 90
    apa = next(b for b in result.breakdown if b.category == "citation_apa")
    assert apa.major == 2
    assert apa.deduction == 10


# ---------------------------------------------------------------------------
# Template structure
# ---------------------------------------------------------------------------

def test_journal_template_structurally_correct():
    sug = build_apa_suggestion(REASON, "journal_article")
    assert "Journal article:" in sug
    assert JOURNAL_TEMPLATE in sug
    assert BOOK_TEMPLATE not in sug
    assert WEBPAGE_TEMPLATE not in sug
    assert WARNING in sug


def test_book_template_structurally_correct():
    sug = build_apa_suggestion(REASON, "book")
    assert "Book:" in sug
    assert BOOK_TEMPLATE in sug
    assert JOURNAL_TEMPLATE not in sug
    assert WARNING in sug


def test_webpage_template_only_when_source_known():
    web = build_apa_suggestion(REASON, "webpage")
    assert WEBPAGE_TEMPLATE in web
    assert WARNING in web

    unknown = build_apa_suggestion(REASON, "unknown")
    assert WEBPAGE_TEMPLATE not in unknown, "webpage template must not appear for unknown source"


def test_unknown_source_type_shows_labelled_alternatives():
    sug = build_apa_suggestion(REASON, "unknown")
    assert "Identify the original source type" in sug
    assert "Journal article:" in sug
    assert JOURNAL_TEMPLATE in sug
    assert "Book:" in sug
    assert BOOK_TEMPLATE in sug
    assert WEBPAGE_TEMPLATE not in sug
    assert WARNING in sug


def test_placeholder_warning_always_present_with_templates():
    for source_type in ("journal_article", "book", "webpage", "unknown"):
        assert WARNING in build_apa_suggestion(REASON, source_type)


def test_suggestion_hierarchy_order():
    sug = build_apa_suggestion(REASON, "unknown")
    order = [sug.index(part) for part in (
        "Recommended correction", "What to verify", "APA 7 formatting example", WARNING)]
    assert order == sorted(order), "sections must appear in hierarchy order"
    for item in ("- Author name and initials", "- Publication year", "- Source title",
                 "- Journal or publisher", "- Volume, issue, and pages where applicable",
                 "- DOI or URL where applicable"):
        assert item in sug


def test_placeholders_not_invented_details():
    sug = build_apa_suggestion(REASON, "journal_article")
    for placeholder in ("A. A.", "(Year)", "xxxxx"):
        assert placeholder in sug
    # No fabricated bibliographic details may appear in the template surface
    for fabricated in ("Garcia", "2018", "vol. 12", "pp. 1-10", "https://doi.org/10.1"):
        assert fabricated not in sug


# ---------------------------------------------------------------------------
# Task-level: missing metadata, distinct same-paragraph guidance
# ---------------------------------------------------------------------------

async def test_item_without_reason_is_rejected(ai_session, monkeypatch):
    findings = [_finding(0)]

    async def _fake(prompt):
        return json.dumps([
            {"finding_key": findings[0]["finding_key"], "source_type": "book", "confidence": 0.9},
        ])

    monkeypatch.setattr(ai_citation, "call_ollama_local", _fake)

    out = await async_ai_citation_task(
        audit_id="tpl-1", db=ai_session, cloud=False, citation_findings=findings)
    assert out.suggestions == []
    assert out.status == "COMPLETED_NO_SUGGESTIONS"


async def test_missing_metadata_yields_placeholders(ai_session, monkeypatch):
    """Model supplies a reason but no source type — placeholders, no guesses."""
    findings = [_finding(0)]

    async def _fake(prompt):
        return json.dumps([
            {"finding_key": findings[0]["finding_key"], "reason": REASON, "confidence": 0.9},
        ])

    monkeypatch.setattr(ai_citation, "call_ollama_local", _fake)

    out = await async_ai_citation_task(
        audit_id="tpl-2", db=ai_session, cloud=False, citation_findings=findings)
    ai_session.commit()

    assert out.status == AI_STATUS_WITH_SUGGESTIONS
    sug = out.suggestions[0]["suggestion"]
    assert "Identify the original source type" in sug
    assert BOOK_TEMPLATE in sug
    assert WARNING in sug
    assert out.suggestions[0]["confidence"] == 0.9


async def test_same_paragraph_findings_keep_distinct_guidance(ai_session, monkeypatch):
    """Two findings in one paragraph keep separate suggestions and messages."""
    findings = [_finding(5, "Smith", "2020"), _finding(5, "Jones", "2019")]

    async def _fake(prompt):
        return json.dumps([
            {"finding_key": findings[0]["finding_key"], "reason": "Add the Smith (2020) entry.", "source_type": "unknown", "confidence": 0.9},
            {"finding_key": findings[1]["finding_key"], "reason": "Add the Jones (2019) entry.", "source_type": "unknown", "confidence": 0.8},
        ])

    monkeypatch.setattr(ai_citation, "call_ollama_local", _fake)

    out = await async_ai_citation_task(
        audit_id="tpl-3", db=ai_session, cloud=False, citation_findings=findings)
    ai_session.commit()

    assert out.status == AI_STATUS_WITH_SUGGESTIONS
    assert len(out.suggestions) == 2
    assert out.suggestions[0]["paragraph_index"] == 5
    assert out.suggestions[1]["paragraph_index"] == 5
    assert "Smith" in out.suggestions[0]["suggestion"]
    assert "Jones" in out.suggestions[1]["suggestion"]
    assert out.suggestions[0]["message"] != out.suggestions[1]["message"]


async def test_confidence_not_forced_to_one(ai_session, monkeypatch):
    """Model confidence is preserved; absence of confidence stays None."""
    findings = [_finding(0)]

    async def _fake(prompt):
        return json.dumps([
            {"finding_key": findings[0]["finding_key"], "reason": REASON, "source_type": "book"},
        ])

    monkeypatch.setattr(ai_citation, "call_ollama_local", _fake)

    out = await async_ai_citation_task(
        audit_id="tpl-4", db=ai_session, cloud=False, citation_findings=findings)
    ai_session.commit()

    assert out.status == AI_STATUS_WITH_SUGGESTIONS
    assert out.suggestions[0]["confidence"] is None, "missing confidence must not become 1.0"


# ---------------------------------------------------------------------------
# End-to-end POST assembly with the real task and a mocked provider
# ---------------------------------------------------------------------------

def test_post_assembles_structured_guidance(client, docx_factory, monkeypatch):
    import app.api.routes as api_routes

    async def _fake(prompt):
        # The prompt embeds the example array before the real findings — take
        # the LAST array (the real findings) so the echoed finding_key matches.
        start = prompt.rindex("[")
        end = prompt.rindex("]") + 1
        findings = json.loads(prompt[start:end])
        return json.dumps([
            {"finding_key": findings[0]["finding_key"], "reason": REASON,
             "source_type": "book", "confidence": 0.9},
        ])

    monkeypatch.setattr(ai_citation, "call_ollama_local", _fake)
    # conftest's client fixture no-ops the route task — route the real task
    # through so the assembled suggestion is exercised end-to-end.
    real_task = ai_citation.async_ai_citation_task

    async def _real(*args, **kwargs):
        return await real_task(*args, **kwargs)

    monkeypatch.setattr(api_routes, "async_ai_citation_task", _real)

    file_bytes = docx_factory(paragraphs=["Orphan (Garcia, 2018) text."], references=None)
    post = client.post(
        "/api/audit",
        files={"file": ("t.docx", file_bytes, "application/octet-stream")},
    )
    assert post.status_code == 200
    data = post.json()
    assert data["ai_review_status"] == AI_STATUS_WITH_SUGGESTIONS
    assert len(data["ai_citation_tooltips"]) == 1
    tooltip = data["ai_citation_tooltips"][0]
    assert tooltip["confidence"] == 0.9
    sug = tooltip["suggestion"]
    assert "Recommended correction" in sug
    assert "What to verify" in sug
    assert "APA 7 formatting example" in sug
    assert BOOK_TEMPLATE in sug
    assert WARNING in sug
    # Deterministic correction names the source; templates never fabricate details.
    assert "No matching References entry was found for Garcia (2018) in this document." in sug
    shared = "\n\n".join(sug.split("\n\n")[1:])
    assert "Garcia" not in shared, "model reason must not fabricate source details into templates"
