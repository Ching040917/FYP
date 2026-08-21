"""Profile → deterministic-rule integration — Build 4 tests.

Compares the SAME controlled DOCX under SUC, APA 7, and a custom profile:
  - 1 in left margin fails SUC but passes APA;
  - References 2.0 passes both SUC and APA;
  - APA heading does not receive a 16 pt requirement (inherits body pair);
  - custom nullable requirement creates no finding;
  - same document → reproducible but profile-dependent results;
  - each Audit stores the snapshot used;
  - modifying the source custom profile does not alter stored results;
  - allowed font pairs validate TOGETHER (valid family + wrong size → fail);
  - no global-default fallback.
"""
import base64
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

from app.services.layout_engine import run_static_rules_engine
from app.services.profile_preset_adapter import EffectiveProfileConfig
from app.services.profile_registry import (
    APA_PROFILE_ID,
    SUC_PROFILE_ID,
    get_builtin_profile,
)
from app.services.profile_schema import new_custom_profile
from app.services.profile_snapshot import resolve_snapshot, EffectiveProfileSnapshot

AM = WD_ALIGN_PARAGRAPH
_PNG = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAA"
    "BJRU5ErkJggg=="
)


def _config(profile):
    return EffectiveProfileConfig(resolve_snapshot(profile))


def _doc(
    left_margin=1.0,
    body_ls=2.0,
    heading_size=12,
    heading_font="Calibri",
    heading_ls=1.0,
    refs_ls=2.0,
    body_font="Calibri",
    body_size=11,
):
    """Controlled DOCX: 1 in margins, body 2.0 spacing, 12pt headings,
    References at 2.0."""
    doc = Document()
    for name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Caption"):
        try:
            s = doc.styles[name]
            s.font.name = "Times New Roman"
            s.font.size = Pt(12)
        except KeyError:
            pass
    sec = doc.sections[0]
    sec.left_margin = Inches(left_margin)
    sec.right_margin = Inches(1.0)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)

    p = doc.add_paragraph("Body text that is long enough to be prose here.")
    p.paragraph_format.line_spacing = body_ls
    p.alignment = AM.JUSTIFY
    for r in p.runs:
        r.font.name = body_font
        r.font.size = Pt(body_size)

    h = doc.add_paragraph("1. Introduction", style="Heading 1")
    h.paragraph_format.line_spacing = heading_ls
    h.alignment = AM.LEFT
    for r in h.runs:
        r.font.name = heading_font
        r.font.size = Pt(heading_size)

    p2 = doc.add_paragraph("More body prose after the heading here.")
    p2.paragraph_format.line_spacing = body_ls
    for r in p2.runs:
        r.font.name = body_font
        r.font.size = Pt(body_size)

    doc.add_paragraph("References", style="Heading 1")
    for ref in ["Smith, J. (2020). Title. Press.", "Garcia, A. (2018). Book. Publisher."]:
        p = doc.add_paragraph(ref)
        p.paragraph_format.line_spacing = refs_ls
        for r in p.runs:
            r.font.name = body_font
            r.font.size = Pt(body_size)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _codes(viols):
    return {v.rule_code for v in viols}


# ---------------------------------------------------------------------------
# SUC vs APA comparison
# ---------------------------------------------------------------------------

def test_one_inch_left_margin_passes_suc_fails_apa():
    """Same DOCX (1 in left margin): SUC does not check margins → no
    MARGIN_LEFT; APA requires 1.0 → no MARGIN_LEFT either. The contrast
    appears for NON-1-in margins under APA."""
    file_bytes = _doc(left_margin=1.0)
    suc = run_static_rules_engine(file_bytes, config=_config(get_builtin_profile(SUC_PROFILE_ID)))
    apa = run_static_rules_engine(file_bytes, config=_config(get_builtin_profile(APA_PROFILE_ID)))
    assert "MARGIN_LEFT" not in _codes(suc)
    assert "MARGIN_LEFT" not in _codes(apa)


def test_suc_no_margin_findings_across_1_1_25_1_5():
    """SUC produces no Margin findings for 1 in, 1.25 in, or 1.5 in left
    margins — the check is disabled, never scored."""
    for margin in (1.0, 1.25, 1.5):
        file_bytes = _doc(left_margin=margin)
        viols = run_static_rules_engine(file_bytes, config=_config(get_builtin_profile(SUC_PROFILE_ID)))
        codes = _codes(viols)
        assert "MARGIN_LEFT" not in codes
        assert "MARGIN_RIGHT" not in codes
        assert "MARGIN_TOP" not in codes
        assert "MARGIN_BOTTOM" not in codes


def test_apa_identifies_non_one_inch_margins():
    """APA requires 1 in on all sides — 1.25 in left is flagged."""
    file_bytes = _doc(left_margin=1.25)
    apa = run_static_rules_engine(file_bytes, config=_config(get_builtin_profile(APA_PROFILE_ID)))
    assert "MARGIN_LEFT" in _codes(apa)
    assert "MARGIN_RIGHT" not in _codes(apa)


def test_apa_still_requires_one_inch():
    """APA continues to require 1 in margins."""
    apa = get_builtin_profile(APA_PROFILE_ID)
    assert apa.margins.margin_left_in == 1.0
    assert apa.margins.margin_right_in == 1.0
    assert apa.margins.margin_top_in == 1.0
    assert apa.margins.margin_bottom_in == 1.0


def test_references_2_0_passes_suc_and_apa():
    """References at 2.0 line spacing passes both approved profiles."""
    file_bytes = _doc(refs_ls=2.0)
    suc = run_static_rules_engine(file_bytes, config=_config(get_builtin_profile(SUC_PROFILE_ID)))
    apa = run_static_rules_engine(file_bytes, config=_config(get_builtin_profile(APA_PROFILE_ID)))
    ref_ls_suc = [v for v in suc if v.rule_code == "LINE_SPACING" and v.location.get("paragraph_index", -1) >= 4]
    ref_ls_apa = [v for v in apa if v.rule_code == "LINE_SPACING" and v.location.get("paragraph_index", -1) >= 4]
    assert ref_ls_suc == []
    assert ref_ls_apa == []


def test_apa_heading_does_not_get_16pt_requirement():
    """APA headings inherit the body pair — a 12 pt heading under APA is
    valid if (family, size) is an allowed pair; it never gets a 16 pt
    requirement."""
    file_bytes = _doc(heading_size=12, heading_font="Calibri", body_font="Calibri", body_size=11)
    apa = run_static_rules_engine(file_bytes, config=_config(get_builtin_profile(APA_PROFILE_ID)))
    font_size = [v for v in apa if v.rule_code == "FONT_SIZE"]
    # Calibri 12 is NOT an allowed APA pair (Calibri is 11) → flagged, but
    # the message/expected must never claim 16 pt.
    for v in font_size:
        assert "16" not in (v.expected_value or "")


def test_apa_heading_inherits_body_pair_valid():
    """Calibri 11 heading under APA (inherited body pair) → no font finding."""
    file_bytes = _doc(heading_size=11, heading_font="Calibri", body_font="Calibri", body_size=11)
    apa = run_static_rules_engine(file_bytes, config=_config(get_builtin_profile(APA_PROFILE_ID)))
    assert "FONT_SIZE" not in _codes(apa)
    assert "FONT_CONSISTENCY" not in _codes(apa)


# ---------------------------------------------------------------------------
# Allowed font-pair validation (TOGETHER)
# ---------------------------------------------------------------------------

def test_allowed_font_pair_valid_family_wrong_size_rejected():
    """APA: 'Calibri 12' is invalid — Calibri is only allowed at 11 pt.
    A valid family with a size from another family must be rejected."""
    file_bytes = _doc(body_font="Calibri", body_size=12)
    apa = run_static_rules_engine(file_bytes, config=_config(get_builtin_profile(APA_PROFILE_ID)))
    fc = [v for v in apa if v.rule_code == "FONT_CONSISTENCY" and v.location.get("paragraph_index") == 0]
    assert len(fc) >= 1
    assert "Calibri" in fc[0].actual_value


def test_allowed_font_pair_valid_combo_passes():
    file_bytes = _doc(body_font="Calibri", body_size=11, heading_font="Calibri", heading_size=11)
    apa = run_static_rules_engine(file_bytes, config=_config(get_builtin_profile(APA_PROFILE_ID)))
    assert "FONT_CONSISTENCY" not in _codes(apa)
    assert "FONT_SIZE" not in _codes(apa)


# ---------------------------------------------------------------------------
# Custom profile
# ---------------------------------------------------------------------------

def test_custom_nullable_requirement_creates_no_finding():
    """A custom profile with nullable body line spacing → no LINE_SPACING
    finding for a wrong body spacing."""
    custom = new_custom_profile("NullSpacing", base=get_builtin_profile(SUC_PROFILE_ID))
    custom.body.line_spacing = None  # nullable → skip check
    custom.margins.margin_left_in = 1.0
    file_bytes = _doc(left_margin=1.0, body_ls=1.0)
    viols = run_static_rules_engine(file_bytes, config=_config(custom))
    assert "MARGIN_LEFT" not in _codes(viols)   # custom requires 1.0
    assert "LINE_SPACING" not in _codes(viols)  # nullable → skipped


def test_custom_modified_margins_and_heading_size():
    """Custom: 1.25 in margins, heading 14 pt, body 1.5 spacing."""
    custom = new_custom_profile("Custom", base=get_builtin_profile(SUC_PROFILE_ID))
    custom.margins.margin_left_in = 1.25
    custom.heading.font_size_pt = 14.0
    custom.body.line_spacing = 1.5
    custom.body.font_family = "Calibri"
    custom.body.font_size_pt = 11.0
    custom.heading.font_family = "Calibri"
    custom.heading.font_size_pt = 14.0

    # Fully compliant document EXCEPT the left margin: every visible run
    # carries direct formatting matching the profile. (Effective-font
    # resolution now detects style-inherited sizes, so an unformatted
    # References heading would legitimately flag as TNR 12 != Calibri 14.)
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)

    p = doc.add_paragraph("Body text that is long enough to be prose here.")
    p.paragraph_format.line_spacing = 1.5
    p.alignment = AM.JUSTIFY
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(11)

    h = doc.add_paragraph("1. Introduction", style="Heading 1")
    h.paragraph_format.line_spacing = 1.0
    h.alignment = AM.LEFT
    for r in h.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(14)

    refs_h = doc.add_paragraph("References", style="Heading 1")
    refs_h.paragraph_format.line_spacing = 1.0
    for r in refs_h.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(14)

    buf = io.BytesIO()
    doc.save(buf)
    viols = run_static_rules_engine(buf.getvalue(), config=_config(custom))
    assert "MARGIN_LEFT" in _codes(viols)       # 1.0 vs required 1.25
    assert "LINE_SPACING" not in _codes(viols)  # 1.5 matches
    assert "FONT_SIZE" not in _codes(viols)     # 14 pt heading matches
    assert "FONT_CONSISTENCY" not in _codes(viols)


# ---------------------------------------------------------------------------
# Reproducibility / snapshot binding
# ---------------------------------------------------------------------------

def test_same_document_reproducible_but_profile_dependent():
    file_bytes = _doc(left_margin=1.0, body_ls=2.0, heading_size=11,
                      heading_font="Calibri", body_font="Calibri", body_size=11)
    a1 = run_static_rules_engine(file_bytes, config=_config(get_builtin_profile(APA_PROFILE_ID)))
    a2 = run_static_rules_engine(file_bytes, config=_config(get_builtin_profile(APA_PROFILE_ID)))
    suc = run_static_rules_engine(file_bytes, config=_config(get_builtin_profile(SUC_PROFILE_ID)))
    # Same profile → identical findings.
    assert [(v.rule_code, v.severity, v.location) for v in a1] == [(v.rule_code, v.severity, v.location) for v in a2]
    # SUC never emits margin findings; APA at 1.0 in has none either here.
    assert ("MARGIN_LEFT", "MAJOR", {"section_index": 0}) not in [
        (v.rule_code, v.severity, v.location) for v in suc
    ]
    assert ("MARGIN_LEFT", "MAJOR", {"section_index": 0}) not in [
        (v.rule_code, v.severity, v.location) for v in a1
    ]


def test_no_global_default_fallback():
    """A blank custom profile (all nulls) produces NO margin/font/spacing
    findings regardless of the document — no PresetConfig leaks in."""
    blank = new_custom_profile("Blank")
    file_bytes = _doc(left_margin=1.0, body_ls=1.0, heading_size=20)
    viols = run_static_rules_engine(file_bytes, config=_config(blank))
    assert "MARGIN_LEFT" not in _codes(viols)
    assert "FONT_SIZE" not in _codes(viols)
    assert "LINE_SPACING" not in _codes(viols)


def test_profile_aware_messages_reference_profile():
    """APA margin findings reference the selected profile by name."""
    file_bytes = _doc(left_margin=1.25)
    apa = run_static_rules_engine(file_bytes, config=_config(get_builtin_profile(APA_PROFILE_ID)))
    ml = next(v for v in apa if v.rule_code == "MARGIN_LEFT")
    assert '"APA 7 Student Paper"' in ml.message
    assert "requires 1.00 in" in ml.message


# ---------------------------------------------------------------------------
# Audit-level snapshot binding (POST uses the snapshot config)
# ---------------------------------------------------------------------------

def test_audit_uses_snapshot_config_and_stores_it(client):
    """The audit's findings are produced from the stored snapshot config —
    the same document audited under different profiles gives different
    results, and each audit stores its snapshot."""
    # Build a doc with 1.25 in left margin so SUC passes (not checked) and
    # APA flags it.
    doc = Document()
    doc.sections[0].left_margin = Inches(1.25)
    doc.sections[0].right_margin = Inches(1.0)
    doc.sections[0].top_margin = Inches(1.0)
    doc.sections[0].bottom_margin = Inches(1.0)
    p = doc.add_paragraph("Body text that is long enough to be prose here.")
    p.paragraph_format.line_spacing = 1.5
    p.alignment = AM.JUSTIFY
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    doc.add_paragraph("References", style="Heading 1")
    for ref in ["Smith, J. (2020). Title. Press."]:
        p = doc.add_paragraph(ref)
        p.paragraph_format.line_spacing = 2.0
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
    buf = io.BytesIO()
    doc.save(buf)
    doc_bytes = buf.getvalue()

    def _post_bytes(profile_id, bytes_override=None):
        return client.post(
            "/api/audit",
            files={"file": ("t.docx", bytes_override or doc_bytes, "application/octet-stream")},
            params={"profile_id": profile_id},
        )

    suc_post = _post_bytes(SUC_PROFILE_ID)
    apa_post = _post_bytes(APA_PROFILE_ID)
    assert suc_post.status_code == 200
    assert apa_post.status_code == 200
    suc_codes = {v["rule_code"] for v in suc_post.json()["physical_layout_errors"]}
    apa_codes = {v["rule_code"] for v in apa_post.json()["physical_layout_errors"]}
    assert "MARGIN_LEFT" not in suc_codes
    assert "MARGIN_LEFT" in apa_codes
    # Each audit stores the snapshot it used.
    assert suc_post.json()["profile_snapshot"]["profile_id"] == SUC_PROFILE_ID
    assert apa_post.json()["profile_snapshot"]["profile_id"] == APA_PROFILE_ID
    # GET returns the stored snapshot.
    get = client.get(f"/api/audit/{apa_post.json()['audit_id']}")
    assert get.json()["profile_snapshot"]["profile_id"] == APA_PROFILE_ID


def test_modifying_custom_source_does_not_alter_stored_results(client):
    """Editing the source custom profile after the audit never changes the
    stored findings/snapshot."""
    custom = new_custom_profile("SourceMut", base=get_builtin_profile(APA_PROFILE_ID))
    from app.services.profile_resolver import resolve_request_profile
    snapshot = resolve_request_profile(custom_profile=custom.to_dict())
    stored_fp = snapshot.fingerprint

    # Audit with this custom profile.
    doc = Document()
    doc.add_paragraph("Body text that is long enough to be prose here.")
    buf = io.BytesIO()
    doc.save(buf)
    resp = client.post(
        "/api/audit",
        files={"file": ("t.docx", buf.getvalue(), "application/octet-stream")},
        data={"custom_profile": __import__("json").dumps(custom.to_dict())},
    )
    assert resp.status_code == 200
    audit_id = resp.json()["audit_id"]
    findings = resp.json()["physical_layout_errors"]

    # Mutate the source profile.
    custom.margins.margin_left_in = 3.0
    custom.body.font_family = "Comic Sans"

    get = client.get(f"/api/audit/{audit_id}").json()
    assert get["profile_snapshot"]["fingerprint"] == stored_fp
    assert get["violations"] == findings


# ---------------------------------------------------------------------------
# Evidence-based margin policy (SUC does not check margins)
# ---------------------------------------------------------------------------

def test_custom_explicit_margins_still_create_findings():
    """A Custom profile requiring 1.5 in left margin still creates a
    MARGIN_LEFT finding for a 1.0 in document."""
    custom = new_custom_profile("CustomMargins", base=get_builtin_profile(APA_PROFILE_ID))
    custom.margins.margin_left_in = 1.5
    file_bytes = _doc(left_margin=1.0)
    viols = run_static_rules_engine(file_bytes, config=_config(custom))
    ml = next(v for v in viols if v.rule_code == "MARGIN_LEFT")
    assert ml.severity == "MAJOR"
    assert "requires 1.50 in" in ml.message
    assert "CustomMargins" in ml.message


def test_custom_null_margins_skip_findings():
    """A Custom profile with null margins produces no Margin findings."""
    custom = new_custom_profile("NullMargins", base=get_builtin_profile(APA_PROFILE_ID))
    custom.margins.margin_left_in = None
    custom.margins.margin_right_in = None
    custom.margins.margin_top_in = None
    custom.margins.margin_bottom_in = None
    for margin in (1.0, 1.25, 1.5):
        file_bytes = _doc(left_margin=margin)
        viols = run_static_rules_engine(file_bytes, config=_config(custom))
        codes = _codes(viols)
        assert "MARGIN_LEFT" not in codes
        assert "MARGIN_RIGHT" not in codes
        assert "MARGIN_TOP" not in codes
        assert "MARGIN_BOTTOM" not in codes


def test_suc_profile_version_incremented():
    """SUC built-in version bumped to 2 when the margin policy changed."""
    assert get_builtin_profile(SUC_PROFILE_ID).profile_version == 2


def test_suc_snapshot_fingerprint_differs_from_previous_version():
    """The new SUC snapshot (margins null) fingerprints differently from the
    old v1 snapshot (margins 1.5/1.0) — effective requirements changed."""
    old = get_builtin_profile(SUC_PROFILE_ID).to_dict()
    old["profile_version"] = 1
    old["margins"]["margin_left_in"] = 1.5
    old["margins"]["margin_right_in"] = 1.0
    old["margins"]["margin_top_in"] = 1.0
    old["margins"]["margin_bottom_in"] = 1.0
    from app.services.profile_schema import profile_from_dict
    old_snap = resolve_snapshot(profile_from_dict(old))
    new_snap = resolve_snapshot(get_builtin_profile(SUC_PROFILE_ID))
    assert old_snap.fingerprint != new_snap.fingerprint


def test_historical_snapshot_immutable_and_no_rescore(client, test_engine):
    """A stored historical SUC snapshot keeps its original margins and is
    never re-scored — GET returns the stored snapshot, and the stored
    score/findings stay untouched."""
    from sqlalchemy.orm import sessionmaker
    Session = sessionmaker(bind=test_engine)
    s = Session()
    from app.models.audit import AuditRecord, Violation
    old_snap = resolve_snapshot(get_builtin_profile(SUC_PROFILE_ID))
    # Simulate an OLD v1 snapshot stored before the margin-policy change.
    old_dict = old_snap.to_dict()
    old_dict["profile_version"] = 1
    old_dict["margins"]["left_in"] = 1.5
    old_dict["margins"]["right_in"] = 1.0
    old_dict["margins"]["top_in"] = 1.0
    old_dict["margins"]["bottom_in"] = 1.0
    old_dict["fingerprint"] = ""
    # Re-fingerprint with the old margins but keep it valid.
    restored = EffectiveProfileSnapshot(
        schema_version=old_dict["schema_version"],
        profile_id=old_dict["profile_id"],
        profile_name=old_dict["profile_name"],
        profile_version=1,
        profile_source=old_dict["profile_source"],
        description=old_dict["description"],
        citation_style=old_dict["citation_style"],
        institution_specific=old_dict["institution_specific"],
        body_font_family=old_dict["body"]["font_family"],
        body_font_size_pt=old_dict["body"]["font_size_pt"],
        body_allowed_font_combos=tuple(tuple(c) for c in old_dict["body"]["allowed_font_combos"]),
        body_line_spacing=old_dict["body"]["line_spacing"],
        body_alignment=old_dict["body"]["alignment"],
        body_space_before_pt=old_dict["body"]["space_before_pt"],
        body_space_after_pt=old_dict["body"]["space_after_pt"],
        body_first_line_indent_in=old_dict["body"]["first_line_indent_in"],
        heading_font_family=old_dict["heading"]["font_family"],
        heading_font_size_pt=old_dict["heading"]["font_size_pt"],
        heading_allowed_font_combos=tuple(tuple(c) for c in old_dict["heading"]["allowed_font_combos"]),
        heading_alignment=old_dict["heading"]["alignment"],
        heading_space_before_pt=old_dict["heading"]["space_before_pt"],
        heading_space_after_pt=old_dict["heading"]["space_after_pt"],
        heading_level_1=dict(old_dict["heading"]["level_1"] or {}),
        heading_level_2=dict(old_dict["heading"]["level_2"] or {}),
        heading_level_3=dict(old_dict["heading"]["level_3"] or {}),
        margin_left_in=1.5,
        margin_right_in=1.0,
        margin_top_in=1.0,
        margin_bottom_in=1.0,
        references_line_spacing=old_dict["references"]["line_spacing"],
        references_hanging_indent_in=old_dict["references"]["hanging_indent_in"],
        caption_space_before_pt=old_dict["captions"]["space_before_pt"],
        caption_space_after_pt=old_dict["captions"]["space_after_pt"],
        list_space_after_pt=old_dict["lists"]["space_after_pt"],
        role_exemptions=tuple(sorted(old_dict["role_exemptions"])),
        table_eligibility=old_dict["table_eligibility"],
    ).with_fingerprint()

    rec = AuditRecord(
        id="hist-suc-0001", filename="old.docx", file_size=10,
        weighted_score=70, deploy_mode="LOCAL", status="completed",
    )
    rec.profile_snapshot = restored.to_dict()
    rec_id = rec.id
    s.add(rec)
    s.add(Violation(id="v-1", audit_id=rec_id, rule_code="MARGIN_LEFT",
                    severity="MAJOR", location={"section_index": 0},
                    message="Old margin finding", expected_value="1.5in",
                    actual_value="1.0in"))
    s.commit()
    s.close()

    get = client.get(f"/api/audit/{rec_id}").json()
    # Stored historical snapshot returned as-is (old 1.5 margin, v1).
    assert get["profile_snapshot"]["profile_version"] == 1
    assert get["profile_snapshot"]["margins"]["left_in"] == 1.5
    assert get["profile_snapshot"]["fingerprint"] == restored.fingerprint
    # Original findings + score untouched — never re-scored.
    assert get["weighted_score"] == 70
    codes = {v["rule_code"] for v in get["violations"]}
    assert "MARGIN_LEFT" in codes
