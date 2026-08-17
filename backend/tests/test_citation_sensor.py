"""Unit tests for the pure-Python APA citation sensor."""
import pytest
from docx import Document

from app.services.citation_sensor import (
    NARRATIVE_PATTERN,
    PARENTHETICAL_PATTERN,
    run_citation_sensor,
)
from app.services.document_parser import extract_paragraphs


def _build_and_scan(paragraphs, references=None):
    """Helper: build a docx, extract paragraphs, run sensor, return violations."""
    doc = Document()
    if references is not None:
        for p in references:
            doc.add_paragraph(p)
    else:
        for p in paragraphs:
            doc.add_paragraph(p)
    if references is not None:
        doc.add_paragraph("References")
        for r in references:
            doc.add_paragraph(r)
    para_meta = extract_paragraphs(doc)
    return run_citation_sensor(doc, para_meta)


# ---------------------------------------------------------------------------
# Regex pattern direct tests
# ---------------------------------------------------------------------------

def test_narrative_form_detected():
    text = "Smith (2020) found that…"
    matches = NARRATIVE_PATTERN.findall(text)
    assert len(matches) == 1
    assert matches[0][0] == "Smith"
    assert matches[0][1] == "2020"


def test_narrative_form_with_page_locator():
    text = "Smith (2020, p. 12) argued…"
    matches = NARRATIVE_PATTERN.findall(text)
    assert len(matches) == 1
    assert matches[0][0] == "Smith"
    assert matches[0][1] == "2020"


def test_parenthetical_single_author_detected():
    text = "As shown earlier (Smith, 2020)."
    matches = PARENTHETICAL_PATTERN.findall(text)
    assert len(matches) == 1
    assert matches[0][0] == "Smith"
    assert matches[0][1] == "2020"


def test_parenthetical_two_authors_with_ampersand():
    text = "Confirmed by (Smith & Jones, 2020)."
    matches = PARENTHETICAL_PATTERN.findall(text)
    assert len(matches) == 1
    assert matches[0][0] == "Smith"
    assert matches[0][1] == "2020"


def test_parenthetical_et_al_detected():
    text = "(Smith et al., 2020) found…"
    matches = PARENTHETICAL_PATTERN.findall(text)
    assert len(matches) == 1
    assert matches[0][0] == "Smith"
    assert matches[0][1] == "2020"


# ---------------------------------------------------------------------------
# Sensor integration with paragraphs
# ---------------------------------------------------------------------------

def test_present_in_references_not_flagged():
    paragraphs = [
        "Intro line.",
        "Smith (2020) wrote about it.",
        "References",
        "Smith, J. (2020). On things. Press.",
    ]
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert viols == []


def test_missing_from_references_flagged_major():
    paragraphs = [
        "Body line.",
        "Smith (2020) wrote about it.",
        "References",
        "Jones, A. (2019). Other things. Press.",
    ]
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert len(viols) == 1
    assert viols[0].severity == "MAJOR"
    assert viols[0].rule_code == "CITATION_MISMATCH"
    assert "Smith" in viols[0].message
    assert "2020" in viols[0].message


def test_no_references_header_flags_every_citation():
    doc = Document()
    doc.add_paragraph("Body line.")
    doc.add_paragraph("Smith (2020) wrote about it.")
    doc.add_paragraph("And (Jones, 2019) confirmed it.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    # Two orphan citations, no References header
    assert len(viols) == 2
    authors = {v.message.split("'")[1].split(" (")[0] for v in viols}
    assert authors == {"Smith", "Jones"}


def test_references_header_case_insensitive():
    doc = Document()
    doc.add_paragraph("Smith (2020) said…")
    doc.add_paragraph("REFERENCES")
    doc.add_paragraph("Smith, J. (2020). Title. Press.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert viols == []


def test_bibliography_alias_recognized():
    doc = Document()
    doc.add_paragraph("Smith (2020) said…")
    doc.add_paragraph("Bibliography")
    doc.add_paragraph("Smith, J. (2020). Title. Press.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert viols == []


def test_works_cited_alias_recognized():
    doc = Document()
    doc.add_paragraph("Smith (2020) said…")
    doc.add_paragraph("Works Cited")
    doc.add_paragraph("Smith, J. (2020). Title. Press.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert viols == []


def test_reference_list_alias_recognized():
    doc = Document()
    doc.add_paragraph("Smith (2020) said…")
    doc.add_paragraph("Reference List")
    doc.add_paragraph("Smith, J. (2020). Title. Press.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert viols == []


def test_multiple_paragraphs_all_scanned():
    doc = Document()
    doc.add_paragraph("Smith (2020) said…")
    doc.add_paragraph("Body middle line.")
    doc.add_paragraph("Brown (2021) replied…")
    doc.add_paragraph("More body.")
    doc.add_paragraph("Jones (2019) noted…")
    doc.add_paragraph("References")
    doc.add_paragraph("Smith, J. (2020). Title. Press.")
    doc.add_paragraph("Brown, A. (2021). Other. Press.")
    # Jones missing -> 1 violation
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert len(viols) == 1
    assert "Jones" in viols[0].message


def test_duplicate_citation_in_same_paragraph_deduped():
    doc = Document()
    doc.add_paragraph("Smith (2020) said one thing. Later Smith (2020) said another.")
    doc.add_paragraph("References")
    doc.add_paragraph("Smith, J. (2020). Title. Press.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert viols == []  # dedupe, not double-counted


def test_empty_paragraphs_returns_empty_list():
    doc = Document()
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert viols == []


def test_violation_location_has_paragraph_index():
    doc = Document()
    doc.add_paragraph("Body intro.")
    doc.add_paragraph("Smith (2020) said…")
    # No References
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert len(viols) == 1
    loc = viols[0].location
    assert "paragraph_index" in loc
    assert loc["paragraph_index"] == 1  # second paragraph


def test_parenthetical_multi_author_resolves_first_surname():
    doc = Document()
    doc.add_paragraph("(Garcia & Lee, 2018) showed…")
    # References has Garcia, not Lee
    doc.add_paragraph("References")
    doc.add_paragraph("Garcia, R. (2018). Title. Press.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    # Garcia present -> no violation
    assert viols == []


def test_year_with_letter_suffix_detected():
    doc = Document()
    doc.add_paragraph("Smith (2020a) wrote…")
    doc.add_paragraph("References")
    doc.add_paragraph("Smith, J. (2020a). Title. Press.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert viols == []


def test_message_format_matches_spec():
    doc = Document()
    doc.add_paragraph("Orphan (Garcia, 2018) text.")
    # No References
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert len(viols) == 1
    expected = "Citation 'Garcia (2018)' was found in text, but no matching entry was found in the References bibliography."
    assert viols[0].message == expected


# ---------------------------------------------------------------------------
# Multi-author narrative forms (Taylor & Green, Smith and Jones)
# ---------------------------------------------------------------------------

def test_narrative_multi_author_ampersand():
    text = "Taylor & Green (2018) found that…"
    matches = NARRATIVE_PATTERN.findall(text)
    assert len(matches) == 1
    assert matches[0][0] == "Taylor & Green"
    assert matches[0][1] == "2018"


def test_narrative_multi_author_word_and():
    text = "Smith and Jones (2019) argue that…"
    matches = NARRATIVE_PATTERN.findall(text)
    assert len(matches) == 1
    assert matches[0][0] == "Smith and Jones"
    assert matches[0][1] == "2019"


def test_narrative_year_suffix_still_detected():
    text = "Doe (2020a) shows that…"
    matches = NARRATIVE_PATTERN.findall(text)
    assert len(matches) == 1
    assert matches[0][0] == "Doe"
    assert matches[0][1] == "2020a"


def test_cross_ref_multi_author_uses_primary():
    """Multi-author narrative cite resolves to bibliography via primary surname."""
    doc = Document()
    doc.add_paragraph("Taylor & Green (2018) explored the topic.")
    doc.add_paragraph("References")
    doc.add_paragraph("Taylor, A. B. (2018). Title. Press.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert viols == []  # Taylor present, Green not required


def test_cross_ref_missing_multi_author_violation():
    """Missing primary author -> CITATION_MISMATCH at correct paragraph_index."""
    doc = Document()
    doc.add_paragraph("Body intro.")
    doc.add_paragraph("Taylor & Green (2018) wrote about it.")
    doc.add_paragraph("References")
    doc.add_paragraph("Jones, A. (2019). Other things. Press.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert len(viols) == 1
    v = viols[0]
    assert v.rule_code == "CITATION_MISMATCH"
    assert v.severity == "MAJOR"
    assert v.location["paragraph_index"] == 1


def test_violation_shape_contract():
    """Pin LayoutViolation shape so future refactors do not silently break it."""
    from app.services.layout_violation import LayoutViolation
    doc = Document()
    doc.add_paragraph("Orphan (Garcia, 2018) text.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert len(viols) == 1
    assert isinstance(viols[0], LayoutViolation)
    assert viols[0].rule_code == "CITATION_MISMATCH"
    assert viols[0].severity == "MAJOR"
    assert isinstance(viols[0].location, dict)
    assert "paragraph_index" in viols[0].location


# ---------------------------------------------------------------------------
# False-positive regression: numbered reference lists + appendix/manifest
# ---------------------------------------------------------------------------

def test_numbered_reference_entry_matches_narrative_citation():
    """Smith (2020) with a numbered reference '1. Smith, J. (2020)' must NOT flag.

    Regression: first-token index building dropped entries prefixed by list
    numbering, turning a valid citation into a false CITATION_MISMATCH.
    """
    doc = Document()
    doc.add_paragraph("Smith (2020) reported the first finding.")
    doc.add_paragraph("References")
    doc.add_paragraph("1. Smith, J. (2020). On things. Press.")
    doc.add_paragraph("2. Jones, A. (2019). Other things. Press.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert viols == []


def test_bulleted_reference_entry_matches():
    doc = Document()
    doc.add_paragraph("Smith (2020) reported the first finding.")
    doc.add_paragraph("References")
    doc.add_paragraph("- Smith, J. (2020). On things. Press.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert viols == []


def test_hyphenated_and_apostrophe_surnames_normalize():
    """Initials, punctuation, and case normalize safely for matching."""
    doc = Document()
    doc.add_paragraph("Smith-Jones (2020) reported it. O'Brien (2019) replied.")
    doc.add_paragraph("References")
    doc.add_paragraph("SMITH-JONES, A. (2020). On things. Press.")
    doc.add_paragraph("O'Brien, P. (2019). Other things. Press.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert viols == []


def test_garcia_and_lee_same_paragraph_separate_mismatches():
    """Two missing authors in one paragraph stay two distinct violations."""
    doc = Document()
    doc.add_paragraph("Smith (2020) reported the first finding.")
    doc.add_paragraph("However, Garcia (2018) and Lee (2021) disagreed.")
    doc.add_paragraph("References")
    doc.add_paragraph("Smith, J. (2020). On things. Press.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert len(viols) == 2
    authors = {v.message.split("'")[1].split(" (")[0] for v in viols}
    assert authors == {"Garcia", "Lee"}
    assert all(v.location["paragraph_index"] == 1 for v in viols)


def test_appendix_manifest_text_not_scanned_as_citations():
    """Paragraphs after the References header must not create findings."""
    doc = Document()
    doc.add_paragraph("Smith (2020) reported the first finding.")
    doc.add_paragraph("However, Garcia (2018) and Lee (2021) disagreed.")
    doc.add_paragraph("References")
    doc.add_paragraph("Smith, J. (2020). On things. Press.")
    doc.add_paragraph("Appendix A: Manifest")
    doc.add_paragraph("Manifest lists Smith (2020) and Garcia (2018) as verified sources.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    # Only the two body-paragraph mismatches; the manifest's Garcia mention
    # must NOT produce a third violation.
    assert len(viols) == 2
    assert all(v.location["paragraph_index"] == 1 for v in viols)
    assert {v.message.split("'")[1].split(" (")[0] for v in viols} == {"Garcia", "Lee"}


def test_fixture_shape_smith_valid_garcia_lee_mismatch():
    """Full fixture expectation: Smith resolves, Garcia + Lee flagged once each."""
    doc = Document()
    doc.add_paragraph("Smith (2020) reported the first finding.")
    doc.add_paragraph("However, Garcia (2018) and Lee (2021) disagreed.")
    doc.add_paragraph("References")
    doc.add_paragraph("1. Smith, J. (2020). On things. Press.")
    doc.add_paragraph("2. Jones, A. (2019). Other things. Press.")
    doc.add_paragraph("Appendix A: Manifest")
    doc.add_paragraph("Manifest lists Smith (2020) and Garcia (2018) as verified sources.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert len(viols) == 2
    assert {v.message.split("'")[1].split(" (")[0] for v in viols} == {"Garcia", "Lee"}


# ---------------------------------------------------------------------------
# References-heading boundary: numbered heading forms + appendix termination
# ---------------------------------------------------------------------------

def _doc_with_header(header_text, body_cite="Smith (2020) wrote about it.",
                     ref_entry="Smith, J. (2020). On things. Press."):
    doc = Document()
    doc.add_paragraph(body_cite)
    doc.add_paragraph(header_text)
    doc.add_paragraph(ref_entry)
    return doc


@pytest.mark.parametrize("header", [
    "References",
    "REFERENCES",
    "6 References",
    "6. References",
    "6.1 References",
    "6.1.2 References",
])
def test_references_heading_forms_recognized(header):
    """Bare, uppercase, and numbered heading forms all bound the section."""
    doc = _doc_with_header(header)
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert viols == []


@pytest.mark.parametrize("misleading", [
    "References are listed below",
    "6 References are listed below",
    "See References section",
    "Reference list items",
])
def test_misleading_references_body_text_not_a_header(misleading):
    """Prose mentioning 'References' must not be treated as the section start."""
    doc = _doc_with_header(misleading)
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    # No real header → reference index empty → Smith is an orphan
    assert len(viols) == 1
    assert "Smith" in viols[0].message


def test_numbered_heading_smith_resolves_garcia_lee_mismatch():
    """Fixture 2 shape: '6 References' Heading 1, Smith entry, appendices."""
    doc = Document()
    doc.add_paragraph("Smith (2020) reported the first finding.")
    doc.add_paragraph("However, Garcia (2018) and Lee (2021) disagreed.")
    doc.add_paragraph("6 References")
    doc.add_paragraph("Smith, J. (2020). On things. Press.")
    doc.add_paragraph("Appendix A: Expected Detection Manifest")
    doc.add_paragraph("Manifest lists Smith (2020) and Garcia (2018) as verified sources.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    authors = {v.message.split("'")[1].split(" (")[0] for v in viols}
    assert authors == {"Garcia", "Lee"}
    assert len(viols) == 2


def test_appendix_heading_terminates_reference_collection():
    """Reference entries after an Appendix heading are not indexed.

    Jones (2019) is cited in the body but its 'entry' lives after the
    Appendix heading — without termination it would be indexed and Jones
    would (wrongly) resolve. With termination Jones stays an orphan.
    """
    doc = Document()
    doc.add_paragraph("Smith (2020) and Jones (2019) wrote about it.")
    doc.add_paragraph("References")
    doc.add_paragraph("Smith, J. (2020). On things. Press.")
    doc.add_paragraph("Appendix A: Extra Material")
    doc.add_paragraph("Jones, A. (2019). Should not be indexed as a reference entry.")
    paras = extract_paragraphs(doc)
    viols = run_citation_sensor(doc, paras)
    assert len(viols) == 1
    assert "Jones" in viols[0].message
    assert "Smith" not in viols[0].message


# ---------------------------------------------------------------------------
# Source-faithful evidence preservation (parenthetical citation defect)
# ---------------------------------------------------------------------------

def _build_with_unrelated_references(body_paragraphs, reference_entries):
    """Doc with body text + a References section that does NOT include the
    body's authors — every body citation is an orphan (mismatch)."""
    doc = Document()
    for p in body_paragraphs:
        doc.add_paragraph(p)
    doc.add_paragraph("References")
    for r in reference_entries:
        doc.add_paragraph(r)
    paras = extract_paragraphs(doc)
    return run_citation_sensor(doc, paras)


def test_parenthetical_citation_preserves_exact_matched_span():
    """`(Garcia, 2018)` must survive as actual_value VERBATIM — the canonical
    narrative form lives only in the message (reference identity), never in
    the evidence."""
    viols = _build_with_unrelated_references(
        ["Results agree with (Garcia, 2018) elsewhere."],
        ["Other, X. (2020). Unrelated entry. Press."],
    )
    assert len(viols) == 1
    v = viols[0]
    assert v.rule_code == "CITATION_MISMATCH"
    # source-faithful evidence: exact span including parentheses and comma
    assert v.actual_value == "(Garcia, 2018)"
    # canonical identity stays in the message (author/year stable)
    assert "Garcia (2018)" in v.message
    assert v.expected_value == "Reference entry for Garcia"


def test_narrative_citation_evidence_remains_narrative():
    """`Garcia (2018)` stays narrative — never rewritten."""
    viols = _build_with_unrelated_references(
        ["Garcia (2018) argues that results matter."],
        ["Other, X. (2020). Unrelated entry. Press."],
    )
    assert len(viols) == 1
    assert viols[0].actual_value == "Garcia (2018)"
    assert "Garcia (2018)" in viols[0].message


def test_parenthetical_multi_author_preserves_span_and_primary_identity():
    """Multi-author parenthetical keeps the whole span; the primary author
    stays the canonical identity."""
    viols = _build_with_unrelated_references(
        ["Confirmed by (Smith & Jones, 2020a, p. 12) earlier."],
        ["Other, X. (2020). Unrelated entry. Press."],
    )
    assert len(viols) == 1
    v = viols[0]
    assert v.actual_value == "(Smith & Jones, 2020a, p. 12)"
    assert v.message.startswith("Citation 'Smith (2020a)'")
    assert "Smith" in v.expected_value


def test_garcia_and_lee_same_paragraph_stay_distinct_with_own_spans():
    """Two parenthetical citations in one paragraph keep SEPARATE findings
    with their OWN source-faithful spans."""
    viols = _build_and_scan(
        ["Garcia (2018) and Lee (2021) disagree with each other."],
    )
    assert len(viols) == 2
    spans = sorted(v.actual_value for v in viols)
    assert spans == ["Garcia (2018)", "Lee (2021)"]
    # canonical identity per finding
    assert {v.message.split("'")[1].split(" (")[0] for v in viols} == {"Garcia", "Lee"}
