"""bbox-fixture-1 generator: python-docx -> LibreOffice PDF.

Coverage: bordered tables (Table Grid), borderless table, merged cells,
spanning table, multiple figures per page, different sizes/alignments,
rotated image, header logo (repeated), decorative image, semantic/manual/
missing captions. No document text is logged by any consumer.
"""
import io
import struct
import zlib
import subprocess
import tempfile
import uuid
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches

OUT = Path(__file__).resolve().parent
SOFFICE = r"C:\Program Files\LibreOffice\program\soffice.com"


def make_png(width, height, rgb):
    """Minimal valid PNG (no deps)."""
    def chunk(typ, data):
        c = struct.pack(">I", len(data)) + typ + data
        return c + struct.pack(">I", zlib.crc32(typ + data) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    raw = b"".join(b"\x00" + bytes(rgb) * width for _ in range(height))
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


def add_image(doc, w_in, h_in, rgb, align=WD_ALIGN_PARAGRAPH.LEFT, rotate_deg=None):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run()
    run.add_picture(io.BytesIO(make_png(64, 48, rgb)), width=Inches(w_in), height=Inches(h_in))
    if rotate_deg is not None:
        inline = p.runs[0]._r.findall(qn("w:drawing"))[0]
        xfrm = inline.find(".//" + qn("a:xfrm"))
        if xfrm is not None:
            xfrm.set("rot", str(int(rotate_deg * 60000)))
    return p


def add_borders(table, single=True):
    """Explicit table borders (independent of style support)."""
    from docx.oxml import OxmlElement
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single" if single else "nil")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def build():
    doc = Document()
    tables_meta = []
    figures_meta = []
    docx_order = []

    def rec_table(name, table, caption=None, above=False, borderless=False):
        cells = [[table.cell(r, c).text for c in range(len(table.columns))] for r in range(len(table.rows))]
        tables_meta.append({"index": name, "cells": cells, "caption": caption, "above": above, "borderless": borderless})

    # header logo (repeated on every page)
    sec = doc.sections[0]
    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.add_run().add_picture(io.BytesIO(make_png(32, 24, (0, 90, 160))), width=Inches(0.25), height=Inches(0.19))
    docx_order.append({"kind": "paragraph", "index": 0})

    # ---- page 1: three tables ----
    doc.add_paragraph("BBox Fixture Title", style="Heading 1")

    cap = doc.add_paragraph("Table 1: Bordered results")
    cap.style = doc.styles["Caption"]  # semantic caption
    docx_order.append({"kind": "paragraph", "index": 1})
    t = doc.add_table(rows=3, cols=3)
    add_borders(t)
    for r in range(3):
        for c in range(3):
            t.cell(r, c).text = f"T1 R{r} C{c}"
    rec_table(0, t, caption="Table 1: Bordered results", above=True)
    docx_order.append({"kind": "table", "index": 0})
    doc.add_paragraph("Between-tables prose paragraph one.")
    doc.add_paragraph("Between-tables prose paragraph two.")
    docx_order.append({"kind": "paragraph", "index": 2})
    docx_order.append({"kind": "paragraph", "index": 3})

    # borderless table (explicit nil borders)
    doc.add_paragraph("Some prose before the borderless table.")
    docx_order.append({"kind": "paragraph", "index": 4})
    t2 = doc.add_table(rows=3, cols=2)
    add_borders(t2, single=False)
    for r in range(3):
        for c in range(2):
            t2.cell(r, c).text = f"T2 R{r} C{c}"
    rec_table(1, t2)
    docx_order.append({"kind": "table", "index": 1})
    doc.add_paragraph("Prose after the borderless table.")
    docx_order.append({"kind": "paragraph", "index": 5})

    # merged cells + manual caption
    cap3 = doc.add_paragraph("Table 3: Merged layout")
    docx_order.append({"kind": "paragraph", "index": 6})
    t3 = doc.add_table(rows=3, cols=3)
    add_borders(t3)
    t3.cell(0, 0).merge(t3.cell(0, 1))
    t3.cell(0, 0).text = "MERGED CELL"
    for r in range(1, 3):
        for c in range(3):
            t3.cell(r, c).text = f"T3 R{r} C{c}"
    t3.cell(2, 2).text = "T3 R2 C2"
    rec_table(2, t3, caption="Table 3: Merged layout", above=True)
    docx_order.append({"kind": "table", "index": 2})

    # ---- page 2: figures ----
    doc.add_page_break()
    doc.add_paragraph("Figures section", style="Heading 1")
    docx_order.append({"kind": "paragraph", "index": 7})
    add_image(doc, 3.0, 2.25, (200, 30, 30), align=WD_ALIGN_PARAGRAPH.CENTER)
    figures_meta.append({"image_index": 1, "decorative": False, "in_header_footer": False})
    docx_order.append({"kind": "figure", "index": 1})
    doc.add_paragraph("Figure 1 caption below the centered image.")
    docx_order.append({"kind": "paragraph", "index": 8})
    add_image(doc, 2.0, 1.5, (30, 160, 60), align=WD_ALIGN_PARAGRAPH.LEFT)
    figures_meta.append({"image_index": 2, "decorative": False, "in_header_footer": False})
    docx_order.append({"kind": "figure", "index": 2})
    doc.add_paragraph("Body text between figures.")
    docx_order.append({"kind": "paragraph", "index": 9})
    add_image(doc, 1.5, 1.125, (30, 60, 200), align=WD_ALIGN_PARAGRAPH.RIGHT, rotate_deg=45)
    figures_meta.append({"image_index": 3, "decorative": False, "in_header_footer": False})
    docx_order.append({"kind": "figure", "index": 3})
    doc.add_paragraph("Rotated image above this text.")
    docx_order.append({"kind": "paragraph", "index": 10})
    # decorative small image inline in body
    add_image(doc, 0.4, 0.3, (220, 220, 40), align=WD_ALIGN_PARAGRAPH.LEFT)
    figures_meta.append({"image_index": 4, "decorative": True, "in_header_footer": False})
    docx_order.append({"kind": "figure", "index": 4})
    doc.add_paragraph("Decorative inline icon paragraph.")
    docx_order.append({"kind": "paragraph", "index": 11})

    # ---- pages 3-4: spanning table ----
    doc.add_page_break()
    doc.add_paragraph("Appendix section", style="Heading 1")
    docx_order.append({"kind": "paragraph", "index": 12})
    doc.add_paragraph("Table 4: Long appendix")
    docx_order.append({"kind": "paragraph", "index": 13})
    t4 = doc.add_table(rows=45, cols=3)
    add_borders(t4)
    for r in range(45):
        for c in range(3):
            t4.cell(r, c).text = f"T4 R{r} C{c}"
    rec_table(3, t4, caption="Table 4: Long appendix", above=True)
    docx_order.append({"kind": "table", "index": 3})
    doc.add_paragraph("After the spanning table.")
    docx_order.append({"kind": "paragraph", "index": 14})

    buf = io.BytesIO()
    doc.save(buf)
    meta = {
        "docx_order": docx_order,
        "tables": tables_meta,
        "figures": figures_meta,
        "drawing_order": [1, 2, 3, 4],
        "rels_order": [1, 2, 3, 4],
        "header_image_indices": [0],
    }
    return buf.getvalue(), meta


def main():
    import json
    import shutil
    workdir = Path(tempfile.mkdtemp(prefix="bboxfix_"))
    try:
        docx_bytes, meta = build()
        docx_path = workdir / f"{uuid.uuid4().hex}.docx"
        docx_path.write_bytes(docx_bytes)
        profile = workdir / "profile"
        subprocess.run(
            [
                SOFFICE, "--headless", "--convert-to", "pdf:writer_pdf_Export",
                "--outdir", str(workdir), f"-env:UserInstallation={profile.as_uri()}",
                str(docx_path),
            ],
            check=True, timeout=120, capture_output=True,
        )
        pdf = next(workdir.glob("*.pdf"))
        out_pdf = OUT / "bbox-fixture-1.pdf"
        out_pdf.write_bytes(pdf.read_bytes())
        meta_path = OUT / "bbox-fixture-1-meta.json"
        meta_path.write_text(json.dumps(meta, indent=1))
        print(f"written {out_pdf.name} bytes={pdf.stat().st_size} meta={meta_path.name}")
    finally:
        shutil.rmtree(workdir, ignore_errors=True)


if __name__ == "__main__":
    main()
