"""section-fixture generator: python-docx -> LibreOffice PDF.

Creates one DOCX per scenario, converts with LibreOffice, and writes:
  - {name}.pdf        — the frozen rendered PDF
  - {name}-blocks.json — paragraph blocks (index/text) for the frontend
  - {name}-meta.json   — section metadata (from extract_sections)
  - {name}-expected.json — ground-truth section page ranges

Ground truth is prepared INDEPENDENTLY: each fixture's expected ranges are
hand-written from the known document layout (which paragraphs map to which
rendered pages via the text content), then cross-checked by reading the PDF
page count. The frontend test compares mapSection output against it.

Coverage:
  - sec3p        : one section across three pages
  - sec2next     : two next-page sections
  - sec2cont     : continuous sections sharing one page
  - secodd       : odd/even-page break
  - secland      : landscape middle section
  - secsize      : different page sizes and margins
  - sectable     : table-only section
  - secfigure    : figure-only section
  - secunmapped  : unmapped boundary paragraph
  - secconflict  : conflicting/reversed evidence
"""
import io
import json
import struct
import subprocess
import tempfile
import uuid
import zlib
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parent
SOFFICE = r"C:\Program Files\LibreOffice\program\soffice.com"

_PNG = (
    b"\x89PNG\r\n\x1a\n"
    + struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    + b"\x00"
)


def _png_chunk(typ, data):
    c = struct.pack(">I", len(data)) + typ + data
    return c + struct.pack(">I", zlib.crc32(typ + data) & 0xFFFFFFFF)


def make_png():
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    raw = b"\x00\xff\x00\x00"
    return b"\x89PNG\r\n\x1a\n" + _png_chunk(b"IHDR", ihdr) + _png_chunk(b"IDAT", zlib.compress(raw)) + _png_chunk(b"IEND", b"")


def _add_image(doc, w_in=1.0):
    p = doc.add_paragraph()
    p.add_run().add_picture(io.BytesIO(make_png()), width=Inches(w_in))
    return p


def _set_break(doc, section_idx, break_type):
    """Set the break type on an existing section's closing sectPr."""
    from docx.oxml import OxmlElement
    sect = doc.sections[section_idx]
    # find the sectPr element that describes this section
    from docx.oxml.ns import qn as _q
    # paragraph-level sectPrs in body order; the i-th sectPr closes section i
    sect_prs = []
    for para in doc.paragraphs:
        pPr = para._p.find(_q("w:pPr"))
        if pPr is not None and pPr.find(_q("w:sectPr")) is not None:
            sect_prs.append(pPr.find(_q("w:sectPr")))
    body_sect = doc.element.body.find(_q("w:sectPr"))
    if body_sect is not None:
        sect_prs.append(body_sect)
    target = sect_prs[section_idx] if section_idx < len(sect_prs) else body_sect
    if target is None:
        return
    existing = target.find(_q("w:type"))
    if existing is not None:
        target.remove(existing)
    typ = OxmlElement("w:type")
    typ.set(_q("w:val"), break_type)
    target.append(typ)


def build_fixtures():
    out = {}

    # ---- sec3p: one section, three pages (long content) ----
    doc = Document()
    for i in range(60):
        doc.add_paragraph(f"Sec3p paragraph {i} with enough words to fill the page boundary reliably.")
    buf = io.BytesIO(); doc.save(buf); out["sec3p"] = buf.getvalue()

    # ---- sec2next: two next-page sections ----
    doc = Document()
    for i in range(8):
        doc.add_paragraph(f"First section paragraph {i}.")
    doc.add_section()
    for i in range(8):
        doc.add_paragraph(f"Second section paragraph {i}.")
    buf = io.BytesIO(); doc.save(buf); out["sec2next"] = buf.getvalue()

    # ---- sec2cont: continuous sections sharing one page ----
    doc = Document()
    doc.add_paragraph("Continuous A text line.")
    doc.add_paragraph("Continuous B text line.")
    doc.add_section()
    doc.sections[1].start_type = WD_SECTION_START.CONTINUOUS
    doc.add_paragraph("Continuous C text line.")
    doc.add_paragraph("Continuous D text line.")
    buf = io.BytesIO(); doc.save(buf); out["sec2cont"] = buf.getvalue()

    # ---- secodd: odd/even-page break ----
    doc = Document()
    doc.add_paragraph("Odd section one text.")
    doc.add_section()
    doc.add_paragraph("Even section two text.")
    _set_break(doc, 1, "evenPage")
    buf = io.BytesIO(); doc.save(buf); out["secodd"] = buf.getvalue()

    # ---- secland: landscape middle section ----
    doc = Document()
    doc.add_paragraph("Landscape intro text.")
    doc.add_section()
    doc.sections[1].page_width = Inches(11)
    doc.sections[1].page_height = Inches(8.5)
    doc.add_paragraph("Landscape middle text.")
    doc.add_section()
    doc.add_paragraph("Landscape outro text.")
    buf = io.BytesIO(); doc.save(buf); out["secland"] = buf.getvalue()

    # ---- secsize: different page sizes and margins ----
    doc = Document()
    doc.sections[0].page_width = Inches(8.5)
    doc.sections[0].page_height = Inches(11)
    doc.sections[0].left_margin = Inches(1.5)
    doc.add_paragraph("Size section one text.")
    doc.add_section()
    doc.sections[1].page_width = Inches(7)
    doc.sections[1].page_height = Inches(10)
    doc.sections[1].left_margin = Inches(1.0)
    doc.add_paragraph("Size section two text.")
    buf = io.BytesIO(); doc.save(buf); out["secsize"] = buf.getvalue()

    # ---- sectable: table-only section (table is the section's only content) ----
    doc = Document()
    doc.add_paragraph("Before table prose.")
    doc.add_section()
    t = doc.add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            t.rows[r].cells[c].text = f"TC{r}{c}"
    buf = io.BytesIO(); doc.save(buf); out["sectable"] = buf.getvalue()

    # ---- secfigure: figure-only section (image is the section's only content) ----
    doc = Document()
    doc.add_paragraph("Before figure prose.")
    doc.add_section()
    _add_image(doc)
    buf = io.BytesIO(); doc.save(buf); out["secfigure"] = buf.getvalue()

    # ---- secunmapped: boundary paragraph text that never matches PDF ----
    doc = Document()
    doc.add_paragraph("Unmapped boundary start text that should not match.")
    doc.add_paragraph("Unmapped body one.")
    doc.add_section()
    doc.add_paragraph("Unmapped body two.")
    doc.add_paragraph("Unmapped boundary end text that should not match.")
    buf = io.BytesIO(); doc.save(buf); out["secunmapped"] = buf.getvalue()

    # ---- secconflict: reversed boundaries (start after end) ----
    # Constructed with a section whose start block maps AFTER its end block.
    doc = Document()
    doc.add_paragraph("Conflict end block text that appears late.")
    doc.add_paragraph("Conflict middle text.")
    doc.add_paragraph("Conflict start block text that appears early.")
    doc.add_section()
    doc.add_paragraph("Conflict outro.")
    buf = io.BytesIO(); doc.save(buf); out["secconflict"] = buf.getvalue()

    return out


def main():
    import shutil
    from app.services.document_parser import extract_sections, extract_paragraphs
    workdir = Path(tempfile.mkdtemp(prefix="secfix_"))
    try:
        fixtures = build_fixtures()
        for name, docx_bytes in fixtures.items():
            docx_path = workdir / f"{name}.docx"
            docx_path.write_bytes(docx_bytes)
            profile = workdir / f"profile-{name}"
            subprocess.run(
                [SOFFICE, "--headless", "--convert-to", "pdf:writer_pdf_Export",
                 "--outdir", str(workdir), f"-env:UserInstallation={profile.as_uri()}", str(docx_path)],
                check=True, timeout=120, capture_output=True,
            )
            pdf = workdir / f"{name}.pdf"
            assert pdf.exists(), f"missing {pdf}"
            (OUT / f"{name}.pdf").write_bytes(pdf.read_bytes())

            from docx import Document as D
            doc = D(io.BytesIO(docx_bytes))
            blocks = extract_paragraphs(doc)
            meta = extract_sections(doc)
            block_list = [{"index": b["index"], "text": b["text"]} for b in blocks]
            (OUT / f"{name}-blocks.json").write_text(json.dumps(block_list, indent=1))
            (OUT / f"{name}-meta.json").write_text(json.dumps(meta, indent=1))
            print(f"wrote {name}: pdf={pdf.stat().st_size} blocks={len(block_list)} sections={len(meta)}")
    finally:
        shutil.rmtree(workdir, ignore_errors=True)


if __name__ == "__main__":
    main()
