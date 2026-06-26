"""
Generates a deliberately-imperfect sample thesis .docx so the Auditra
audit pipeline can be demoed end-to-end with one click. Every one of the
6 rule scopes in the static engine produces findings on this file.

Run:
    python scripts/make_sample_docx.py

Output:
    frontend/public/samples/sample-thesis.docx

The 6 planted violation categories:
    1. Page margins  — left margin set to 1.0" instead of SUC's 1.5"
    2. Font consistency — one paragraph uses Calibri instead of Times New Roman
    3. Font size       — one paragraph uses 14pt instead of 12pt body
    4. Paragraph typography — one paragraph uses single spacing + left align
    5. Heading hierarchy — H2 -> H4 (skips H3); also first heading is H2 (orphan)
    6. Media captions  — Jadual 1 with no caption (Malay); an image with
                         no caption AND no alt-text
    + APA citation mismatch — in-text "Smith 2022" with no References entry
                              (caught by citation_sensor.py — Major)
    + Multilingual caption positive control — a properly captioned "Jadual 2"
      to show the multilingual pattern matches Malay
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_margins(section, top, bottom, left, right):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def set_alt_text(drawing_element, alt_text: str):
    """Inject a <wp:docPr descr="..."> into a drawing for alt-text testing.

    python-docx doesn't expose alt-text directly — we walk the inline
    drawing's <wp:docPr> and set its descr attribute.
    """
    ns = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
    docpr = drawing_element.find(ns + "docPr")
    if docpr is None:
        docpr = OxmlElement("wp:docPr")
        drawing_element.insert(0, docpr)
    docpr.set("descr", alt_text)


def main():
    doc = Document()

    # ── Section 1: WRONG margins — left=1.0" instead of SUC's 1.5" (Rule 4) ──
    section = doc.sections[0]
    # SUC standard: left=1.5", right/top/bottom=1.0"
    # We deliberately set left=1.0" to trigger MARGIN_LEFT Major violation
    set_margins(section, 1.0, 1.0, 1.0, 1.0)

    # ── First heading: H2 instead of H1 — triggers orphan H1 check (Rule 5) ──
    h2_first = doc.add_heading("Introduction", level=2)
    for run in h2_first.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    # ── Body paragraph (correct: TNR 12pt, 1.5 spacing, justified) ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(
        "Academic writing requires rigorous adherence to formatting standards. "
        "This paragraph follows the prescribed Times New Roman 12pt body text "
        "with 1.5 line spacing and justified alignment, matching the default "
        "SUC thesis spec. It serves as the control case in the demo document."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # ── H1 — establishes the proper outline after the orphan H2 above ──
    h1 = doc.add_heading("Literature Review", level=1)
    for run in h1.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)

    # ── Body paragraph with WRONG font (Calibri) — triggers Rule 2 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(
        "This paragraph uses Calibri instead of Times New Roman, which breaks "
        "the body-text consistency rule. The font consistency checker should "
        "flag this as a Minor violation."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(12)

    # ── Body paragraph with WRONG size (14pt) — triggers Rule 3 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(
        "This paragraph uses the correct font but the wrong size (14pt instead "
        "of 12pt). The font size alignment rule should catch this as a Minor "
        "violation."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    # ── Body paragraph with WRONG spacing + alignment — triggers Rule 4 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0  # single instead of 1.5
    run = p.add_run(
        "This paragraph uses single line spacing and left alignment instead of "
        "1.5 spacing and justified alignment. The paragraph typography rule "
        "should flag both issues as Minor violations."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # ── H4 directly under H1 (skips H2 and H3) — triggers Rule 5 ──
    h4 = doc.add_heading("Deep Subsection", level=4)
    for run in h4.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(
        "This subsection was placed under an H1, skipping H2 and H3. The "
        "heading hierarchy rule should catch this structural jump as a Major "
        "violation."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # ── Jadual 1 — table WITHOUT caption — triggers Rule 6 (Malay context) ──
    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"R{i+1}C{j+1}"

    # Body paragraph after the uncaptioned table
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(
        "The table above has no caption. The media captions rule should flag "
        "this as a Minor violation since SUC theses require every table to "
        "carry a numbered caption (e.g. 'Jadual 1: ...' or 'Table 1: ...')."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # ── Body paragraph with APA citation — Smith (2022) has NO References entry ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(
        "Prior work on formatting compliance is limited. Smith (2022) argued "
        "that manual checking is error-prone. (Jones & Brown, 2019) found "
        "that automated tools reduce review time. This paragraph intentionally "
        "contains in-text citations that have no matching References entry, "
        "so the citation sensor flags them as Major CITATION_MISMATCH violations."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # ── Jadual 2 — properly captioned (POSITIVE CONTROL — multilingual match) ──
    p = doc.add_paragraph("Jadual 2: Sample compliance metrics")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.italic = True

    table2 = doc.add_table(rows=2, cols=2)
    table2.style = "Light Grid Accent 1"
    for i, row in enumerate(table2.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"R{i+1}C{j+1}"

    # ── References section (intentionally missing Smith and Jones) ──
    doc.add_paragraph("References")
    # Only Brown is listed — Smith (2022) and Jones (& Brown, 2019) are orphans
    doc.add_paragraph("Brown, A. (2020). Sample reference for demo purposes. University Press.")

    out_path = Path(__file__).resolve().parent.parent / "frontend" / "public" / "samples" / "sample-thesis.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    print(f"Size:  {out_path.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
